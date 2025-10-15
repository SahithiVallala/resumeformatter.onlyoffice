"""
Test to verify no duplication and proper formatting
"""
import os
import sys
from docx import Document

# Simple test data
test_data = {
    'name': 'Test Candidate',
    'email': 'test@email.com',
    'phone': '555-1234',
    'experience': [
        {
            'company': 'Company A',
            'role': 'Software Engineer',
            'duration': '2020-2023',
            'details': ['Built applications', 'Led team']
        },
        {
            'company': 'Company B',
            'role': 'Developer',
            'duration': '2018-2020',
            'details': ['Developed features']
        }
    ],
    'education': [
        {
            'degree': 'Master of Science',
            'institution': 'University A',
            'year': '2018',
            'details': []
        }
    ],
    'skills': ['Python', 'Java'],
    'sections': {}
}

print("\n" + "="*70)
print("TESTING: No Duplication + Proper Formatting")
print("="*70)

from models.database import TemplateDB
from config import Config
from utils.word_formatter import format_word_document

# Get template
db = TemplateDB()
templates = db.get_all_templates()

if not templates:
    print("\n❌ No templates found!")
    sys.exit(1)

template = db.get_template(templates[0]['id'])
template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])

print(f"\nUsing template: {template['name']}")

template_analysis = {
    'template_path': template_path,
    'template_type': 'docx'
}

output_path = os.path.join(Config.OUTPUT_FOLDER, 'test_no_duplication.docx')

print(f"\nFormatting document...")
print(f"Output: {output_path}\n")

success = format_word_document(test_data, template_analysis, output_path)

if success:
    print("\n" + "="*70)
    print("✅ SUCCESS")
    print("="*70)
    
    # Verify output
    doc = Document(output_path)
    
    # Count how many times "Company A" appears
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    company_a_count = full_text.count('Company A')
    
    print(f"\n📊 Verification:")
    print(f"  • Total paragraphs: {len(doc.paragraphs)}")
    print(f"  • Total tables: {len(doc.tables)}")
    print(f"  • 'Company A' appears: {company_a_count} time(s)")
    
    if company_a_count == 1:
        print(f"\n✅ NO DUPLICATION - Company A appears exactly once!")
    else:
        print(f"\n⚠️  DUPLICATION DETECTED - Company A appears {company_a_count} times (should be 1)")
    
    print(f"\n📁 Open file to verify:")
    print(f"   {output_path}")
    
else:
    print("\n❌ FAILED - Check logs above")

print("\n")
