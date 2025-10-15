"""
Complete end-to-end test: Parse resume -> Format with template
"""
import os
import sys
from docx import Document

# Test data with correct structure (company, role, duration)
test_resume_data = {
    'name': 'John Smith',
    'email': 'john.smith@email.com',
    'phone': '(555) 987-6543',
    'address': '456 Tech Ave, San Francisco, CA',
    'linkedin': 'linkedin.com/in/johnsmith',
    'dob': '05/20/1985',
    'experience': [
        {
            'company': 'Infosys',
            'role': 'Senior Developer',
            'duration': '2021-2025',
            'title': 'Senior Developer - Infosys',
            'details': [
                'Led development of enterprise applications using Java and Spring Boot',
                'Managed team of 5 junior developers',
                'Implemented microservices architecture',
                'Reduced system latency by 40% through optimization'
            ]
        },
        {
            'company': 'Tech Solutions Inc',
            'role': 'Software Engineer',
            'duration': '2018-2021',
            'title': 'Software Engineer - Tech Solutions Inc',
            'details': [
                'Developed web applications using React and Node.js',
                'Integrated third-party APIs',
                'Participated in agile development process'
            ]
        }
    ],
    'education': [
        {
            'degree': 'Master of Science in Data Science',
            'institution': 'PS University',
            'year': '2020-2022',
            'details': [
                'Specialized in Machine Learning and Big Data Analytics',
                'GPA: 3.8/4.0'
            ]
        },
        {
            'degree': 'Bachelor of Technology in Computer Science',
            'institution': 'State Technical University',
            'year': '2014-2018',
            'details': [
                'Graduated with Honors'
            ]
        }
    ],
    'skills': ['Python', 'Java', 'React', 'Node.js', 'AWS', 'Docker', 'Kubernetes', 'SQL', 'MongoDB'],
    'sections': {}
}

print("\n" + "="*70)
print("COMPLETE WORKFLOW TEST")
print("="*70)

# Get template
from models.database import TemplateDB
from config import Config

db = TemplateDB()
templates = db.get_all_templates()

if not templates:
    print("\nERROR: No templates found!")
    print("Please add a template first")
    sys.exit(1)

template = db.get_template(templates[0]['id'])
template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])

print(f"\nTemplate: {template['name']}")
print(f"Template file: {template_path}")
print(f"Template exists: {os.path.exists(template_path)}")

# Show parsed data structure
print("\n" + "="*70)
print("PARSED RESUME DATA")
print("="*70)

print(f"\nCandidate: {test_resume_data['name']}")
print(f"Email: {test_resume_data['email']}")
print(f"Phone: {test_resume_data['phone']}")

print(f"\n--- EXPERIENCE ({len(test_resume_data['experience'])} entries) ---")
for idx, exp in enumerate(test_resume_data['experience'], 1):
    print(f"\nEntry {idx}:")
    print(f"  Company: {exp['company']}")
    print(f"  Role: {exp['role']}")
    print(f"  Duration: {exp['duration']}")
    print(f"  Details: {len(exp['details'])} items")

print(f"\n--- EDUCATION ({len(test_resume_data['education'])} entries) ---")
for idx, edu in enumerate(test_resume_data['education'], 1):
    print(f"\nEntry {idx}:")
    print(f"  Degree: {edu['degree']}")
    print(f"  Institution: {edu['institution']}")
    print(f"  Year: {edu['year']}")

# Format with Word formatter
print("\n" + "="*70)
print("FORMATTING WITH WORD TEMPLATE")
print("="*70)

from utils.word_formatter import format_word_document

template_analysis = {
    'template_path': template_path,
    'template_type': 'docx'
}

output_path = os.path.join(Config.OUTPUT_FOLDER, 'test_formatted_output.docx')

print(f"\nRunning formatter...")
print(f"Output: {output_path}\n")

success = format_word_document(test_resume_data, template_analysis, output_path)

print("\n" + "="*70)
if success:
    print("SUCCESS - Document formatted!")
    print("="*70)
    print(f"\nOutput file: {output_path}")
    print(f"File exists: {os.path.exists(output_path)}")
    
    # Open and verify output
    print("\n--- VERIFICATION ---")
    doc = Document(output_path)
    print(f"Output has {len(doc.paragraphs)} paragraphs")
    print(f"Output has {len(doc.tables)} tables")
    
    print("\nLooking for experience/education tables...")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
        if len(table.rows) > 0:
            first_cell_text = table.rows[0].cells[0].text[:60]
            print(f"  First cell: {first_cell_text}...")
    
    print(f"\nOPEN THE FILE TO SEE FORMATTED RESULT:")
    print(f"  {output_path}")
else:
    print("FAILED - Check logs above")
    print("="*70)

print("\n")
