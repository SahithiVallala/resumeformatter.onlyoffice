"""
Advanced Template Analyzer
Deeply analyzes templates to extract:
1. Visual elements (letterheads, logos, images)
2. Text placeholders and their positions
3. Section structure and formatting
4. Field mappings (name, email, phone, etc.)
"""

import pdfplumber
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
try:
    from PIL import Image
except ImportError:
    Image = None
import io

class TemplateAnalyzer:
    """Comprehensive template analysis"""
    
    def __init__(self, template_path):
        self.template_path = template_path
        self.file_type = self._get_file_type()
        self.analysis = {}
        
    def _get_file_type(self):
        ext = os.path.splitext(self.template_path)[1].lower()
        return ext.replace('.', '')
    
    def analyze(self):
        """Main analysis method"""
        print(f"\n{'='*70}")
        print(f"🔍 ANALYZING TEMPLATE: {os.path.basename(self.template_path)}")
        print(f"{'='*70}\n")
        
        if self.file_type == 'pdf':
            return self._analyze_pdf()
        elif self.file_type == 'docx':
            return self._analyze_docx()
        else:
            print(f"⚠️ Unsupported file type: {self.file_type}")
            return self._get_default_analysis()
    
    def _analyze_pdf(self):
        """Deep PDF template analysis"""
        try:
            with pdfplumber.open(self.template_path) as pdf:
                page = pdf.pages[0]
                
                # Extract all text with positions
                chars = page.chars
                words = page.extract_words()
                
                # Detect images/logos
                images = self._extract_pdf_images(page)
                
                # Detect placeholders
                placeholders = self._detect_placeholders(words)
                
                # Analyze layout zones
                zones = self._analyze_layout_zones(words, page.width, page.height)
                
                # Detect fields (name, email, phone, etc.)
                fields = self._detect_fields(words, placeholders)
                
                # Analyze formatting
                formatting = self._analyze_pdf_formatting(chars)
                
                analysis = {
                    'template_path': self.template_path,
                    'template_type': 'pdf',
                    'page': {
                        'width': page.width,
                        'height': page.height,
                    },
                    'images': images,
                    'placeholders': placeholders,
                    'fields': fields,
                    'zones': zones,
                    'formatting': formatting,
                    'has_letterhead': self._has_letterhead(zones, images),
                    'sections': self._detect_sections_advanced(words)
                }
                
                self._print_analysis_summary(analysis)
                return analysis
                
        except Exception as e:
            print(f"❌ Error analyzing PDF: {e}")
            import traceback
            traceback.print_exc()
            return self._get_default_analysis()
    
    def _analyze_docx(self):
        """Deep DOCX template analysis"""
        try:
            doc = Document(self.template_path)
            
            # Extract images
            images = self._extract_docx_images(doc)
            
            # Analyze paragraphs
            paragraphs_data = []
            placeholders = []
            fields = {}
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                para_data = {
                    'index': i,
                    'text': text,
                    'alignment': para.alignment,
                    'style': para.style.name if para.style else 'Normal',
                    'runs': []
                }
                
                # Analyze runs (formatted text segments)
                for run in para.runs:
                    run_data = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else 11,
                        'color': self._get_color(run.font.color)
                    }
                    para_data['runs'].append(run_data)
                
                paragraphs_data.append(para_data)
                
                # Detect placeholders
                detected = self._detect_text_placeholders(text)
                if detected:
                    placeholders.extend(detected)
                
                # Detect fields
                field = self._identify_field_type(text)
                if field:
                    fields[field] = {
                        'paragraph_index': i,
                        'original_text': text
                    }
            
            # Analyze tables
            tables_data = self._analyze_docx_tables(doc)
            
            # Detect sections
            sections = self._detect_docx_sections(paragraphs_data)
            
            analysis = {
                'template_path': self.template_path,
                'template_type': 'docx',
                'page': {
                    'width': 612,  # Letter size default
                    'height': 792,
                },
                'images': images,
                'placeholders': placeholders,
                'fields': fields,
                'paragraphs': paragraphs_data,
                'tables': tables_data,
                'sections': sections,
                'has_letterhead': len(images) > 0 or self._has_header_footer(doc),
                'formatting': self._get_docx_formatting(doc)
            }
            
            self._print_analysis_summary(analysis)
            return analysis
            
        except Exception as e:
            print(f"❌ Error analyzing DOCX: {e}")
            import traceback
            traceback.print_exc()
            return self._get_default_analysis()
    
    def _extract_pdf_images(self, page):
        """Extract images from PDF page"""
        images = []
        try:
            # Get images from page
            if hasattr(page, 'images'):
                for img in page.images:
                    images.append({
                        'x': img.get('x0', 0),
                        'y': img.get('top', 0),
                        'width': img.get('width', 0),
                        'height': img.get('height', 0),
                        'type': 'image'
                    })
        except:
            pass
        return images
    
    def _extract_docx_images(self, doc):
        """Extract images from DOCX"""
        images = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append({
                        'type': 'image',
                        'relationship_id': rel.rId
                    })
        except:
            pass
        return images
    
    def _detect_placeholders(self, words):
        """Detect placeholder text patterns"""
        placeholders = []
        placeholder_patterns = [
            r'\[([A-Z_]+)\]',  # [NAME], [EMAIL]
            r'\{([A-Za-z_]+)\}',  # {name}, {email}
            r'<([A-Za-z_]+)>',  # <name>, <email>
            r'YOUR\s+([A-Z]+)',  # YOUR NAME
            r'CANDIDATE\s+([A-Z]+)',  # CANDIDATE NAME
        ]
        
        for word in words:
            text = word['text']
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    placeholders.append({
                        'text': text,
                        'x': word['x0'],
                        'y': word['top'],
                        'type': matches[0].lower()
                    })
        
        return placeholders
    
    def _detect_text_placeholders(self, text):
        """Detect placeholders in text"""
        patterns = [
            (r'\[NAME\]', 'name'),
            (r'\[EMAIL\]', 'email'),
            (r'\[PHONE\]', 'phone'),
            (r'\[ADDRESS\]', 'address'),
            (r'YOUR NAME', 'name'),
            (r'your\.email@', 'email'),
            (r'\(123\) 456-7890', 'phone'),
            (r'linkedin\.com/in/username', 'linkedin'),
        ]
        
        found = []
        for pattern, field_type in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                found.append({
                    'pattern': pattern,
                    'type': field_type,
                    'text': text
                })
        return found
    
    def _detect_fields(self, words, placeholders):
        """Detect field types from context"""
        fields = {}
        
        # Common field indicators
        field_indicators = {
            'name': ['name', 'candidate', 'applicant'],
            'email': ['email', 'e-mail', '@'],
            'phone': ['phone', 'mobile', 'tel', 'contact'],
            'address': ['address', 'location', 'city'],
            'linkedin': ['linkedin', 'profile'],
            'dob': ['dob', 'date of birth', 'birth date'],
            'experience': ['experience', 'work history', 'employment'],
            'education': ['education', 'qualification', 'degree'],
            'skills': ['skills', 'competencies', 'expertise']
        }
        
        for word in words:
            text_lower = word['text'].lower()
            for field_type, indicators in field_indicators.items():
                if any(ind in text_lower for ind in indicators):
                    if field_type not in fields:
                        fields[field_type] = {
                            'x': word['x0'],
                            'y': word['top'],
                            'indicator_text': word['text']
                        }
        
        return fields
    
    def _identify_field_type(self, text):
        """Identify what type of field this text represents"""
        text_lower = text.lower()
        
        if any(x in text_lower for x in ['[name]', 'your name', 'candidate name']):
            return 'name'
        elif any(x in text_lower for x in ['[email]', 'email', '@']):
            return 'email'
        elif any(x in text_lower for x in ['[phone]', 'phone', 'mobile', '(123)']):
            return 'phone'
        elif any(x in text_lower for x in ['[address]', 'address', 'location']):
            return 'address'
        elif 'linkedin' in text_lower:
            return 'linkedin'
        elif any(x in text_lower for x in ['dob', 'date of birth']):
            return 'dob'
        
        return None
    
    def _analyze_layout_zones(self, words, width, height):
        """Divide page into zones (header, body, footer)"""
        if not words:
            return {'header': [], 'body': [], 'footer': []}
        
        # Define zones
        header_threshold = height * 0.15  # Top 15%
        footer_threshold = height * 0.85  # Bottom 15%
        
        zones = {'header': [], 'body': [], 'footer': []}
        
        for word in words:
            y = word['top']
            if y < header_threshold:
                zones['header'].append(word)
            elif y > footer_threshold:
                zones['footer'].append(word)
            else:
                zones['body'].append(word)
        
        return zones
    
    def _has_letterhead(self, zones, images):
        """Check if template has letterhead"""
        # Letterhead typically in header zone with images or special formatting
        return len(images) > 0 or len(zones.get('header', [])) > 0
    
    def _has_header_footer(self, doc):
        """Check if DOCX has header/footer"""
        try:
            for section in doc.sections:
                if section.header.paragraphs or section.footer.paragraphs:
                    return True
        except:
            pass
        return False
    
    def _analyze_pdf_formatting(self, chars):
        """Analyze PDF text formatting"""
        if not chars:
            return {}
        
        fonts = {}
        sizes = {}
        
        for char in chars:
            font = char.get('fontname', 'Unknown')
            size = char.get('size', 10)
            
            fonts[font] = fonts.get(font, 0) + 1
            sizes[size] = sizes.get(size, 0) + 1
        
        # Get most common
        common_font = max(fonts.items(), key=lambda x: x[1])[0] if fonts else 'Helvetica'
        common_size = max(sizes.items(), key=lambda x: x[1])[0] if sizes else 10
        
        return {
            'fonts': fonts,
            'sizes': sizes,
            'common_font': common_font,
            'common_size': common_size
        }
    
    def _get_docx_formatting(self, doc):
        """Get DOCX formatting info"""
        return {
            'styles': [style.name for style in doc.styles],
            'has_header': self._has_header_footer(doc)
        }
    
    def _get_color(self, color):
        """Extract color from font"""
        try:
            if color and color.rgb:
                return str(color.rgb)
        except:
            pass
        return None
    
    def _analyze_docx_tables(self, doc):
        """Analyze tables in DOCX"""
        tables_data = []
        for i, table in enumerate(doc.tables):
            table_data = {
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'cells': []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data['cells'].append(row_data)
            
            tables_data.append(table_data)
        
        return tables_data
    
    def _detect_sections_advanced(self, words):
        """Advanced section detection"""
        sections = []
        
        section_keywords = [
            'experience', 'employment', 'work history', 'professional background',
            'education', 'academic', 'degrees',
            'skills', 'competencies', 'technical', 'expertise',
            'summary', 'profile', 'objective', 'about',
            'projects', 'portfolio',
            'certifications', 'certificates', 'licenses',
            'awards', 'achievements', 'honors', 'qualifications'
        ]
        
        for word in words:
            text = word['text'].lower().strip()
            if any(keyword in text for keyword in section_keywords):
                sections.append({
                    'heading': word['text'],
                    'x': word['x0'],
                    'y': word['top'],
                    'font': word.get('fontname', 'Helvetica'),
                    'size': word.get('size', 11)
                })
        
        return sections
    
    def _detect_docx_sections(self, paragraphs_data):
        """Detect sections in DOCX"""
        sections = []
        
        section_keywords = [
            'experience', 'employment', 'work history', 'professional background',
            'education', 'academic', 'degrees',
            'skills', 'competencies', 'technical', 'expertise',
            'summary', 'profile', 'objective', 'about',
            'projects', 'portfolio',
            'certifications', 'certificates', 'licenses',
            'awards', 'achievements', 'honors'
        ]
        
        for para in paragraphs_data:
            text_lower = para['text'].lower()
            # Check if paragraph contains any section keyword
            if any(keyword in text_lower for keyword in section_keywords):
                # Also check if it looks like a heading (short text, possibly bold)
                if len(para['text']) < 100:  # Headings are usually short
                    sections.append({
                        'heading': para['text'],
                        'index': para['index'],
                        'style': para['style']
                    })
        
        return sections
    
    def _get_default_analysis(self):
        """Default analysis structure"""
        return {
            'template_path': self.template_path,
            'template_type': self.file_type,
            'page': {'width': 612, 'height': 792},
            'images': [],
            'placeholders': [],
            'fields': {},
            'zones': {'header': [], 'body': [], 'footer': []},
            'sections': [],
            'has_letterhead': False,
            'formatting': {}
        }
    
    def _print_analysis_summary(self, analysis):
        """Print analysis summary"""
        print(f"📄 Template Type: {analysis['template_type'].upper()}")
        print(f"📐 Page Size: {analysis['page']['width']} x {analysis['page']['height']}")
        print(f"🖼️  Images Found: {len(analysis['images'])}")
        print(f"📝 Placeholders: {len(analysis['placeholders'])}")
        print(f"🏷️  Fields Detected: {len(analysis['fields'])}")
        print(f"📑 Sections Found: {len(analysis['sections'])}")
        print(f"🎨 Has Letterhead: {'Yes' if analysis['has_letterhead'] else 'No'}")
        
        if analysis['fields']:
            print(f"\n🔍 Detected Fields:")
            for field_name in analysis['fields'].keys():
                print(f"   • {field_name.upper()}")
        
        if analysis['sections']:
            print(f"\n📚 Detected Sections:")
            for section in analysis['sections'][:5]:
                print(f"   • {section['heading']}")
        
        print(f"\n{'='*70}\n")


def analyze_template(template_path):
    """Main function to analyze template"""
    analyzer = TemplateAnalyzer(template_path)
    return analyzer.analyze()
