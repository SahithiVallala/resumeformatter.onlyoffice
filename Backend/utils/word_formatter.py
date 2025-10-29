"""
Enhanced Word Document Formatter
Handles both .doc and .docx templates
Preserves all formatting, images, headers, footers
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import os
import re
import shutil
import traceback
import json

# Try to import win32com for .doc support
try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False
    print("⚠️  win32com not available - .doc files will have limited support")

class WordFormatter:
    """Enhanced Word document formatting"""
    
    def __init__(self, resume_data, template_analysis, output_path):
        self.resume_data = resume_data
        self.template_analysis = template_analysis
        self.output_path = output_path
        self.template_path = template_analysis.get('template_path')
        self.template_type = template_analysis.get('template_type')
        
    def format(self):
        """Main formatting method"""
        print(f"\n{'='*70}")
        print(f"📝 WORD DOCUMENT FORMATTING")
        print(f"{'='*70}\n")
        
        print(f"📄 Template: {os.path.basename(self.template_path)}")
        print(f"👤 Candidate: {self.resume_data['name']}")
        print(f"📁 Output: {os.path.basename(self.output_path)}\n")
        
        try:
            # Handle .doc files
            if self.template_path.lower().endswith('.doc'):
                return self._format_doc_file()
            else:
                return self._format_docx_file()
                
        except Exception as e:
            print(f"❌ Error formatting Word document: {e}")
            traceback.print_exc()
            return False
    
    def _format_doc_file(self):
        """Handle .doc files (old Word format)"""
        print("📋 Processing .doc file (old Word format)...")
        
        if HAS_WIN32:
            # Convert .doc to .docx first
            print("✓ Converting .doc to .docx...")
            docx_path = self._convert_doc_to_docx(self.template_path)
            
            if docx_path:
                # Update template path temporarily
                original_path = self.template_path
                self.template_path = docx_path
                
                # Format the docx
                result = self._format_docx_file()
                
                # Cleanup
                try:
                    os.remove(docx_path)
                except:
                    pass
                
                self.template_path = original_path
                return result
            else:
                print("❌ Failed to convert .doc to .docx")
                return False
        else:
            print("⚠️  Cannot process .doc files without win32com")
            print("💡 Please convert template to .docx format or install pywin32")
            return False
    
    def _convert_doc_to_docx(self, doc_path):
        """Convert .doc to .docx using Word COM"""
        try:
            import pythoncom
            pythoncom.CoInitialize()  # Initialize COM for this thread
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Open .doc file
            doc = word.Documents.Open(os.path.abspath(doc_path))
            
            # Save as .docx
            docx_path = doc_path.replace('.doc', '_temp.docx')
            doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)  # 16 = docx format
            
            doc.Close()
            word.Quit()
            
            pythoncom.CoUninitialize()  # Clean up COM
            
            print(f"✓ Converted to: {docx_path}")
            return docx_path
            
        except Exception as e:
            print(f"❌ Conversion error: {e}")
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except:
                pass
            return None

    def _postprocess_with_word_com(self, docx_path):
        """Final pass using Word COM to replace placeholders that may live in shapes/text boxes.
        Only affects placeholders that still exist (i.e., were not handled by python-docx).
        """
        if not HAS_WIN32:
            return
        try:
            import pythoncom
            pythoncom.CoInitialize()  # Initialize COM for this thread
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))

            def find_replace_in_range(rng, find_text, replace_text, wildcard=False):
                find = rng.Find
                find.ClearFormatting()
                find.Replacement.ClearFormatting()
                find.Text = find_text
                find.Replacement.Text = replace_text
                find.Forward = True
                find.Wrap = 1  # wdFindContinue
                find.MatchCase = False
                find.MatchWholeWord = False
                find.MatchWildcards = bool(wildcard)
                find.Execute(Replace=2)  # wdReplaceAll

            # Build replacement strings from resume_data
            # SUMMARY
            summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {})) or []
            summary_text = (self.resume_data.get('summary') or '').strip()
            summary_replace = ''
            if summary_lines:
                summary_replace = '\r'.join(['• ' + s.strip().lstrip('•–—-*● ') for s in summary_lines if s.strip()])
            elif summary_text:
                summary_replace = summary_text

            # SKILLS - COM replacement (Note: COM has 255 char limit per field)
            skills_list = self.resume_data.get('skills', []) or []
            skill_lines = []
            for s in skills_list:  # Include all skills (will truncate if exceeds COM limit)
                skill_name = (s if isinstance(s, str) else s.get('name', '')).strip()
                if skill_name and len(skill_name) < 50:  # Skip very long skill names
                    skill_lines.append('• ' + skill_name)
            skills_replace = '\r'.join(skill_lines)
            # Limit to 255 chars
            if len(skills_replace) > 255:
                skills_replace = skills_replace[:252] + '...'

            # EDUCATION
            education = self.resume_data.get('education', []) or []
            if not education:
                sect = self._find_matching_resume_section('education', self.resume_data.get('sections', {})) or []
                if sect:
                    education = self._build_education_from_bullets(sect)
            edu_lines = []
            for edu in education:  # Include all entries (will truncate if exceeds COM limit)
                deg = (edu.get('degree') or '').strip()
                inst = (edu.get('institution') or '').strip()
                yr = self._clean_duration((edu.get('year') or '').strip())
                
                # Keep it short for COM
                if deg and len(deg) > 50:
                    deg = deg[:47] + '...'
                if inst and len(inst) > 40:
                    inst = inst[:37] + '...'
                
                parts = []
                if deg:
                    parts.append(deg)
                if inst:
                    parts.append(inst)
                if yr:
                    parts.append(yr)
                line = ' - '.join(parts[:-1]) if len(parts) > 1 else (parts[0] if parts else '')
                if yr and line:
                    line = f"{line} {yr}"
                if line and len(line) < 200:  # Safety check
                    edu_lines.append('• ' + line)
            education_replace = '\r'.join(edu_lines)
            
            # Limit total length to 255 chars (Word COM limit)
            if len(education_replace) > 255:
                education_replace = education_replace[:252] + '...'
            
            print(f"  📝 COM replacement strings prepared:")
            print(f"     - Education: {len(education_replace)} chars, {len(edu_lines)} lines")
            print(f"     - Skills: {len(skills_replace)} chars")
            print(f"     - Summary: {len(summary_replace)} chars")

            # Build candidate name replace string
            candidate_name = (self.resume_data.get('name') or '').strip()
            bracket_name = f"<{candidate_name}>" if candidate_name else ''

            # Story ranges include shapes/text frames and headers/footers
            story = doc.StoryRanges(1)  # wdMainTextStory = 1
            while story is not None:
                # Candidate name → bracketed name
                if bracket_name:
                    # Exact case and uppercase variants
                    find_replace_in_range(story, candidate_name, bracket_name, wildcard=False)
                    find_replace_in_range(story, candidate_name.upper(), bracket_name, wildcard=False)

                if summary_replace:
                    for pat in ["<summary>", "<professional summary>", "<profile>", "professional summary", "<summary*>"]:
                        find_replace_in_range(story, pat, summary_replace, wildcard='*' in pat)
                if skills_replace:
                    for pat in ["<skills>", "<technical skills>", "<list skills>", "<skills*>"]:
                        find_replace_in_range(story, pat, skills_replace, wildcard='*' in pat)
                if education_replace:
                    edu_pats = [
                        "<List candidate’s education background>",
                        "<List candidate’s education background>",
                        "<list candidate’s education background>",
                        "<list candidate’s education background>",
                        "list candidate’s education background",
                        "list candidate’s education background",
                        "education background",
                        "<education background>",
                        "<education>",
                        "<academic background>",
                        "<academic qualifications>",
                    ]
                    for pat in edu_pats:
                        find_replace_in_range(story, pat, education_replace, wildcard=False)
                story = story.NextStoryRange

            # Also traverse shapes explicitly (in case some shapes are not part of StoryRanges loop)
            def replace_in_shapes(shapes):
                for shp in shapes:
                    try:
                        if shp.TextFrame.HasText:
                            rng = shp.TextFrame.TextRange
                            # Candidate name
                            if bracket_name:
                                find_replace_in_range(rng, candidate_name, bracket_name)
                                find_replace_in_range(rng, candidate_name.upper(), bracket_name)
                            if summary_replace:
                                for pat in ["<summary>", "<professional summary>", "<profile>"]:
                                    find_replace_in_range(rng, pat, summary_replace)
                            if skills_replace:
                                for pat in ["<skills>", "<technical skills>", "<list skills>"]:
                                    find_replace_in_range(rng, pat, skills_replace)
                            if education_replace:
                                for pat in [
                                    "<List candidate’s education background>",
                                    "<List candidate’s education background>",
                                    "<list candidate’s education background>",
                                    "<list candidate’s education background>",
                                    "list candidate’s education background",
                                    "list candidate’s education background",
                                    "education background",
                                ]:
                                    find_replace_in_range(rng, pat, education_replace)
                    except Exception:
                        continue

            replace_in_shapes(doc.Shapes)
            for sec in doc.Sections:
                replace_in_shapes(sec.Headers(1).Shapes)
                replace_in_shapes(sec.Headers(2).Shapes)
                replace_in_shapes(sec.Headers(3).Shapes)
                replace_in_shapes(sec.Footers(1).Shapes)
                replace_in_shapes(sec.Footers(2).Shapes)
                replace_in_shapes(sec.Footers(3).Shapes)

            doc.Save()
            doc.Close(False)
            word.Quit()
            pythoncom.CoUninitialize()  # Clean up COM
            print("✓ COM post-processing complete (shapes/text boxes handled)")
        except Exception as e:
            print(f"⚠️  COM post-processing error: {e}")
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except:
                pass

    def _scan_primary_anchors(self, doc):
        """Scan the template once to locate primary anchors for SUMMARY, SKILLS, EMPLOYMENT, EDUCATION.
        If multiple EDUCATION headings exist and one is embedded immediately after EMPLOYMENT
        (likely sample content), pick the later EDUCATION heading as the primary.
        Returns (primary_anchors, all_anchors).
        """
        def is_heading_text(t):
            if not t:
                return False
            t = t.strip()
            return len(t) < 50

        keys = {
            'EMPLOYMENT': ['EMPLOYMENT HISTORY', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EMPLOYMENT', 'EXPERIENCE'],
            'EDUCATION': ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND', 'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS', 'TRAINING', 'ACADEMICS', 'EDUCATION/CERTIFICATES', 'EDUCATION / CERTIFICATES'],
            'SKILLS': ['SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES', 'EXPERTISE', 'ABILITIES'],
            'SUMMARY': ['SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE', 'CAREER SUMMARY', 'EXECUTIVE SUMMARY', 'OVERVIEW'],
            'PROJECTS': ['PROJECTS', 'PORTFOLIO', 'PERSONAL PROJECTS', 'KEY PROJECTS'],
            'AWARDS': ['AWARDS', 'ACHIEVEMENTS', 'HONORS', 'RECOGNITION', 'ACCOMPLISHMENTS'],
            'PUBLICATIONS': ['PUBLICATIONS', 'PAPERS', 'ARTICLES', 'RESEARCH'],
            'LANGUAGES': ['LANGUAGES', 'LANGUAGE SKILLS'],
            'REFERENCES': ['REFERENCES', 'RECOMMENDATIONS']
        }
        all_anchors = {k: [] for k in keys}
        for idx, p in enumerate(doc.paragraphs):
            txt = (p.text or '').strip().upper()
            if not is_heading_text(txt):
                continue
            for k, aliases in keys.items():
                if any(a == txt or txt.startswith(a) for a in aliases):
                    all_anchors[k].append(idx)
                    break

        primary = {k: (v[0] if v else None) for k, v in all_anchors.items()}

        # Heuristic: if an EDUCATION heading appears immediately (<= 8 paras) after EMPLOYMENT
        # and there exists another EDUCATION later, treat the first as embedded placeholder
        emp_idx = primary.get('EMPLOYMENT')
        edu_list = all_anchors.get('EDUCATION') or []
        if emp_idx is not None and len(edu_list) > 1:
            first_edu = edu_list[0]
            if 0 <= (first_edu - emp_idx) <= 8:
                primary['EDUCATION'] = edu_list[-1]

        print("\n🔎 Anchor scan:")
        for k in ['SUMMARY', 'SKILLS', 'EMPLOYMENT', 'EDUCATION', 'PROJECTS', 'AWARDS', 'PUBLICATIONS', 'LANGUAGES', 'REFERENCES']:
            primary_idx = primary.get(k)
            all_idx = all_anchors.get(k, [])
            if primary_idx is not None or all_idx:
                print(f"  - {k}: primary={primary_idx} all={all_idx}")

        return primary, all_anchors
    
    def _build_template_order_map(self, doc):
        """Build a map of template section order to respect original template structure"""
        print("\n📋 Building template section order map...")
        
        # Initialize tracking variables
        self._template_section_order = []
        self._template_section_positions = {}
        self._last_known_section_position = 0
        self._existing_template_sections = {}
        
        section_keywords = {
            'SUMMARY': ['SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE'],
            'EMPLOYMENT': ['EMPLOYMENT', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EXPERIENCE'],
            'EDUCATION': ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS', 'ACADEMICS', 'QUALIFICATIONS'],
            'SKILLS': ['SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES', 'EXPERTISE'],
            'PROJECTS': ['PROJECTS', 'PORTFOLIO'],
            'AWARDS': ['AWARDS', 'ACHIEVEMENTS', 'HONORS'],
            'LANGUAGES': ['LANGUAGES'],
            'REFERENCES': ['REFERENCES']
        }
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            if len(text) < 50 and len(text) > 0:  # Likely a heading
                for section_name, keywords in section_keywords.items():
                    if any(kw in text for kw in keywords):
                        if section_name not in self._template_section_positions:
                            self._template_section_positions[section_name] = para_idx
                            self._template_section_order.append(section_name)
                            self._existing_template_sections[section_name] = text
                            print(f"  ✓ Found {section_name} at paragraph {para_idx}: '{text}'")
                            self._last_known_section_position = para_idx
                        break
        
        print(f"  📊 Template order: {' → '.join(self._template_section_order)}")
        print(f"  📍 Last section position: {self._last_known_section_position}")
    
    def _add_dynamic_candidate_sections(self, doc):
        """
        Dynamically detect and add any sections from candidate resume that don't exist in template.
        This handles custom sections like hobbies, volunteer work, publications, etc.
        Sections are added AFTER all template sections in template's formatting style.
        """
        added_count = 0
        
        # Get all section names from candidate resume
        candidate_sections = self.resume_data.get('sections', {})
        
        # Standard sections that we already handle
        standard_sections = {
            'summary', 'profile', 'objective', 'professional_summary',
            'employment', 'experience', 'work_history', 'work_experience', 'professional_experience',
            'education', 'academic_background', 'certificates', 'certifications',
            'skills', 'technical_skills', 'core_competencies'
        }
        
        # Find sections in candidate resume that aren't in template
        dynamic_sections = {}
        for section_name, section_content in candidate_sections.items():
            section_lower = section_name.lower().replace('_', ' ').replace('-', ' ')
            
            # Skip if it's a standard section we already processed
            if any(std in section_lower for std in standard_sections):
                continue
            
            # Skip if empty
            if not section_content or (isinstance(section_content, list) and len(section_content) == 0):
                continue
            
            # Check if this section already exists in template
            section_exists = False
            for template_section in self._existing_template_sections.values():
                if section_lower in template_section.lower() or template_section.lower() in section_lower:
                    section_exists = True
                    break
            
            if not section_exists:
                dynamic_sections[section_name] = section_content
                print(f"  🔍 Found dynamic section: {section_name} ({len(section_content) if isinstance(section_content, list) else 1} items)")
        
        if not dynamic_sections:
            return 0
        
        # Find insertion point: after last template section
        insertion_point = self._last_known_section_position + 10
        if insertion_point >= len(doc.paragraphs):
            insertion_point = len(doc.paragraphs) - 1
        
        print(f"  📍 Will insert dynamic sections after paragraph {insertion_point}")
        
        # Add each dynamic section
        for section_name, content in dynamic_sections.items():
            try:
                # Format section name
                display_name = section_name.replace('_', ' ').replace('-', ' ').title()
                
                # Insert section heading
                if insertion_point < len(doc.paragraphs):
                    anchor_para = doc.paragraphs[insertion_point]
                    heading_para = self._insert_paragraph_after(anchor_para, display_name.upper())
                else:
                    heading_para = doc.add_paragraph(display_name.upper())
                
                # Format heading
                for run in heading_para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                heading_para.paragraph_format.space_before = Pt(6)
                heading_para.paragraph_format.space_after = Pt(3)
                
                # Insert content
                last_para = heading_para
                content_list = content if isinstance(content, list) else [content]
                
                for item in content_list:
                    item_text = str(item).strip()
                    if item_text:
                        content_para = self._insert_paragraph_after(last_para, f"• {item_text}")
                        content_para.paragraph_format.left_indent = Inches(0.25)
                        for run in content_para.runs:
                            run.font.size = Pt(10)
                        content_para.paragraph_format.space_after = Pt(2)
                        last_para = content_para
                
                added_count += 1
                insertion_point += len(content_list) + 2
                print(f"  ✅ Added {display_name} section with {len(content_list)} items")
                
            except Exception as e:
                print(f"  ⚠️  Error adding {section_name}: {e}")
                continue
        
        return added_count
    
    def _format_docx_file(self):
        """Format .docx file"""
        print("📋 Processing .docx file...")
        
        # Open template
        doc = Document(self.template_path)
        
        print(f"✓ Template loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        
        # Pre-scan anchors so we always insert into the correct template sections
        self._primary_anchors, self._all_anchors = self._scan_primary_anchors(doc)
        
        # Build template section order map
        self._build_template_order_map(doc)
        
        # Initialize section tracking flags
        self._summary_inserted = False
        self._experience_inserted = False
        self._skills_inserted = False
        self._education_inserted = False
        
        # DEBUG: Show what education data we have
        print(f"\n📊 Resume data check:")
        print(f"   - education list: {len(self.resume_data.get('education', []))} entries")
        print(f"   - sections.education: {len(self.resume_data.get('sections', {}).get('education', []))} lines")
        if self.resume_data.get('education'):
            for i, edu in enumerate(self.resume_data.get('education', [])[:3]):
                print(f"      {i+1}. {edu.get('degree', 'N/A')[:40]} | {edu.get('institution', 'N/A')[:30]} | {edu.get('year', 'N/A')}")
        
        # Ensure CAI CONTACT section is inserted with persistent data ONLY if template has it
        try:
            # Check if template has CAI CONTACT section
            has_cai_contact = False
            for p in doc.paragraphs[:20]:  # Check first 20 paragraphs
                if 'CAI CONTACT' in (p.text or '').upper():
                    has_cai_contact = True
                    break
            
            if has_cai_contact:
                print(f"  ✓ Template has CAI CONTACT section, will process it")
                self._ensure_cai_contact(doc)
            else:
                print(f"  ⏭️  Template does not have CAI CONTACT section, skipping")
        except Exception as e:
            print(f"❌ Error ensuring CAI contact: {e}")
        
        # Show what data we have from resume
        print(f"\n📊 Resume Data Available:")
        print(f"  • Name: {self.resume_data.get('name', 'NOT FOUND')}")
        print(f"  • Email: {self.resume_data.get('email', 'NOT FOUND')}")
        print(f"  • Phone: {self.resume_data.get('phone', 'NOT FOUND')}")
        print(f"  • Experience entries: {len(self.resume_data.get('experience', []))}")
        print(f"  • Education entries: {len(self.resume_data.get('education', []))}")
        print(f"  • Skills: {len(self.resume_data.get('skills', []))}")
        print(f"  • Sections: {list(self.resume_data.get('sections', {}).keys())}")
        
        # PHASE 1: Template Structure Analysis - Respect existing template layout
        print(f"\n🔍 PHASE 1: Template Structure Analysis...")
        
        # First, do a comprehensive scan of existing template sections and their positions
        self._existing_template_sections = self._scan_existing_template_sections(doc)
        self._candidate_sections = self._extract_all_candidate_sections()
        
        print(f"  📋 Template existing sections: {list(self._existing_template_sections.keys())}")
        print(f"  📋 Template section order: {[f'{k}(pos:{v})' for k,v in self._existing_template_sections.items()]}")
        print(f"  📋 Candidate sections: {list(self._candidate_sections.keys())}")
        
        # Mark sections that exist in template - DO NOT CREATE DUPLICATES
        self._mark_existing_template_sections()
        
        # Only identify truly missing sections (not in template at all)
        truly_missing_sections = self._identify_truly_missing_sections()
        if truly_missing_sections:
            print(f"  ⚠️  Truly missing sections (will add): {truly_missing_sections}")
            # Add the missing sections to the template
            self._add_missing_template_sections(doc, truly_missing_sections)
        else:
            print(f"  ✅ All essential sections exist in template - will fill existing sections only")
        
        # Ensure education completeness with enhanced extraction
        enhanced_education = self._ensure_education_completeness()
        if enhanced_education:
            self.resume_data['education'] = enhanced_education
            print(f"  🎓 Enhanced education extraction: {len(enhanced_education)} comprehensive entries")
        
        # Create comprehensive replacement map
        replacements = self._create_replacement_map()
        print(f"\n📝 Created {len(replacements)} replacement mappings")
        
        # CRITICAL: Process skills tables in-place based on table headers (respect template order)
        table_replaced = 0
        print(f"\n🔍 STEP 1: Scanning {len(doc.tables)} tables...")
        for table_idx, table in enumerate(doc.tables):
            # Check if this is a skills table
            if self._is_skills_table(table):
                print(f"  📊 Found skills table at index {table_idx}")
                skills_filled = self._fill_skills_table(table)
                print(f"  ✅ Filled {skills_filled} skill rows")
                table_replaced += skills_filled
                # Mark as inserted if we actually filled rows
                if skills_filled > 0:
                    self._skills_inserted = True
        
        print(f"\n✓ Processed {table_replaced} table entries")
        
        # STEP 2: Replace in all paragraphs
        replaced_count = 0
        print(f"\n🔍 STEP 2: Scanning {len(doc.paragraphs)} paragraphs for placeholders...")
        
        # Prepare bracketed name and compute name anchor index
        candidate_name = self.resume_data.get('name', '').strip()
        
        # CRITICAL: Remove role/title from name if present (e.g., "JOHN DOE BUSINESS ANALYST" → "JOHN DOE")
        # Role usually appears after multiple spaces or on same line
        if candidate_name:
            # Split by multiple spaces or newlines
            name_parts = re.split(r'\s{2,}|\n', candidate_name)
            if len(name_parts) > 1:
                # First part is usually the actual name
                candidate_name = name_parts[0].strip()
                print(f"  🔧 Cleaned name: '{candidate_name}' (removed role/title)")
            
            # Also check for common role keywords and remove everything after them
            role_keywords = ['BUSINESS ANALYST', 'SOFTWARE ENGINEER', 'PROJECT MANAGER', 'DATA SCIENTIST', 
                           'DEVELOPER', 'CONSULTANT', 'SPECIALIST', 'COORDINATOR', 'MANAGER', 'ANALYST',
                           'ENGINEER', 'ARCHITECT', 'DESIGNER', 'ADMINISTRATOR']
            for keyword in role_keywords:
                if keyword in candidate_name.upper():
                    # Remove the keyword and everything after it
                    idx = candidate_name.upper().index(keyword)
                    candidate_name = candidate_name[:idx].strip()
                    print(f"  🔧 Removed role keyword, final name: '{candidate_name}'")
                    break
        
        bracketed_name = f"<{candidate_name}>" if candidate_name else '<Candidate Name>'
        self._name_anchor_idx = None
        try:
            # Look for name or name placeholder in main content area (skip early CAI CONTACT section)
            name_placeholder_patterns = [
                r'<\s*[Cc]andidate[^>]*[Nn]ame[^>]*>',
                r'<\s*[Nn]ame\s*>',
                r'<\s*[Ff]ull\s*[Nn]ame\s*>',
                r'<\s*YOUR\s*NAME\s*>',
            ]
            for idx, p in enumerate(doc.paragraphs[:40]):
                t = (p.text or '').strip()
                # Skip very early paragraphs (likely CAI CONTACT) - increased from 5 to 10
                if idx < 10:
                    continue
                # Also skip if this paragraph is in CAI CONTACT section
                if 'CAI CONTACT' in t.upper():
                    continue
                # Check for actual name or bracketed name
                if (candidate_name and t == candidate_name) or (bracketed_name and t == bracketed_name):
                    self._name_anchor_idx = idx
                    print(f"  📍 Name anchor found at paragraph {idx}: '{t}'")
                    break
                # Check for name placeholder patterns
                for pat in name_placeholder_patterns:
                    if re.search(pat, t):
                        self._name_anchor_idx = idx
                        print(f"  📍 Name placeholder anchor found at paragraph {idx}: '{t[:50]}'")
                        break
                if self._name_anchor_idx is not None:
                    break
        except Exception as e:
            print(f"  ⚠️  Name anchor detection error: {e}")
            self._name_anchor_idx = None
        
        # PRE-PASS: Remove all SKILLS sections before EMPLOYMENT to prevent left-panel fills
        print(f"\n🔍 PRE-PASS: Removing any SKILLS sections in CAI CONTACT/left panel only...")
        try:
            emp_idx = self._primary_anchors.get('EMPLOYMENT')
            name_idx = getattr(self, '_name_anchor_idx', None)
            # Limit removal strictly to very early left-panel region (CAI CONTACT), not the main content
            left_boundary = name_idx if name_idx is not None else 5
            skills_removed = 0
            # Determine scan limit: before EMPLOYMENT but never beyond the left boundary
            scan_limit = min(emp_idx, left_boundary) if emp_idx is not None else left_boundary
            if scan_limit and scan_limit > 0:
                for idx in range(scan_limit - 1, -1, -1):
                    if idx >= len(doc.paragraphs):
                        continue
                    para = doc.paragraphs[idx]
                    t = (para.text or '').strip().upper()
                    if t in ('SKILLS', 'TECHNICAL SKILLS') or ('SKILLS' in t and len(t) < 30):
                        print(f"  🗑️  Removing SKILLS heading at para {idx} (CAI CONTACT area)")
                        # Clear content after this heading until next section or for ~20 lines
                        j = idx + 1
                        cleared = 0
                        while j < len(doc.paragraphs) and cleared < 20:
                            para_j = doc.paragraphs[j]
                            txt = (para_j.text or '').strip().upper()
                            # Stop at next major section
                            if len(txt) < 50 and any(h in txt for h in ['EMPLOYMENT', 'WORK HISTORY', 'EDUCATION', 'SUMMARY', 'CAI CONTACT', 'CERTIFICATIONS']):
                                break
                            try:
                                for r in para_j.runs:
                                    r.text = ''
                            except Exception:
                                pass
                            j += 1
                            cleared += 1
                        # Delete the SKILLS heading itself
                        try:
                            self._delete_paragraph(para)
                            skills_removed += 1
                        except Exception:
                            pass
            print(f"  ✅ Removed {skills_removed} SKILLS section(s) from CAI CONTACT area")
        except Exception as e:
            print(f"  ⚠️  Pre-pass error: {e}")
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            # Check each replacement
            for key, value in replacements.items():
                if self._text_contains(paragraph.text, key):
                    print(f"  📍 Found '{key}' in paragraph {para_idx}: '{paragraph.text[:50]}...'")
                    count = self._replace_in_paragraph(paragraph, key, value)
                    if count > 0:
                        print(f"  ✅ Replaced with: '{value[:50]}...'")
                    else:
                        print(f"  ⚠️  Found but couldn't replace (might be in multiple runs)")

            # Regex-driven fallback for angle bracket placeholders with variations
            # Candidate name generic patterns - very flexible to catch all variations
            name_patterns = [
                r"<\s*Candidate['']?s?\s+full\s+name\s*>",
                r"<\s*Candidate['']?s?\s+name\s*>",
                r"<\s*Name\s*>",
                r"<\s*Full\s+Name\s*>",
                r"<\s*YOUR\s+NAME\s*>",
            ]
            for pat in name_patterns:
                if re.search(pat, paragraph.text, re.IGNORECASE):
                    before = paragraph.text
                    # Replace with angle-bracketed name format: <CANDIDATE NAME>
                    self._regex_replace_paragraph(paragraph, pat, bracketed_name)
                    if paragraph.text != before:
                        print(f"  ✅ Regex replaced candidate name in paragraph {para_idx} with {bracketed_name}")
                        replaced_count += 1

            # Generic catch-all: any <...> containing both 'candidate' and 'name' (any order)
            generic_name_pat = r"<[^>]*(?:candidate[^>]*name|name[^>]*candidate)[^>]*>"
            if re.search(generic_name_pat, paragraph.text, re.IGNORECASE):
                before = paragraph.text
                self._regex_replace_paragraph(paragraph, generic_name_pat, bracketed_name)
                if paragraph.text != before:
                    print(f"  ✅ Generic regex replaced candidate name in paragraph {para_idx} with {bracketed_name}")
                    replaced_count += 1

            # CRITICAL: Check if this is an EMPLOYMENT HISTORY section heading
            if not self._experience_inserted:
                para_upper = paragraph.text.strip().upper()
                is_emp_heading = any(h in para_upper for h in [
                    'EMPLOYMENT HISTORY', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE',
                    'WORK EXPERIENCE', 'CAREER HISTORY', 'EMPLOYMENT'
                ])
                
                # NOTE: We don't gate by primary anchor here because paragraph indices shift after insertions
                # Instead, we rely on the _experience_inserted flag to prevent duplicates
                if is_emp_heading and len(paragraph.text.strip()) < 50:
                    print(f"  💼 Found EMPLOYMENT HISTORY heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    
                    # CRITICAL: Preserve and format the heading text
                    original_heading = paragraph.text.strip()
                    if not any(h in original_heading.upper() for h in ['EMPLOYMENT', 'EXPERIENCE', 'WORK']):
                        original_heading = 'Employment History'  # Fallback
                    
                    # Clear and reset heading with proper formatting (BOLD + UNDERLINE + CAPS)
                    paragraph.clear()
                    run = paragraph.add_run(original_heading.upper())
                    run.bold = True
                    run.underline = True
                    run.font.size = Pt(12)
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.space_after = Pt(6)
                    
                    experience_data = self.resume_data.get('experience', [])
                    if not experience_data or len(experience_data) == 0:
                        print(f"     ⚠️  No experience data to insert, skipping")
                    else:
                        print(f"     → Will insert {len(experience_data)} experience entries after heading")
                        
                        # CRITICAL: Clear ALL content between EMPLOYMENT heading and next section
                        # This prevents old template content, sample names, and placeholders from remaining
                        paras_to_clear = []
                        stop_headings = [
                            'EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND',
                            'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS',
                            'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS', 'ACADEMICS',
                            'SKILLS', 'TECHNICAL SKILLS',
                            'SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE',
                            'PROJECTS', 'AWARDS', 'REFERENCES'
                        ]
                        
                        # Scan ahead and collect paragraphs to clear
                        for check_idx in range(para_idx + 1, min(para_idx + 150, len(doc.paragraphs))):
                            check_para = doc.paragraphs[check_idx]
                            check_text = check_para.text.strip().upper()
                            check_text_full = check_para.text.strip()
                            
                            # CRITICAL: Skip if this looks like the heading itself being re-checked
                            if check_idx == para_idx:
                                continue
                            
                            # Stop at next major section heading
                            if any(h in check_text for h in stop_headings) and len(check_text) < 50:
                                # Check if this is an education-related heading (EDUCATION, CERTIFICATES, etc.)
                                is_edu_related = any(edu_term in check_text for edu_term in [
                                    'EDUCATION', 'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS',
                                    'ACADEMIC', 'QUALIFICATIONS', 'ACADEMICS'
                                ])
                                
                                if is_edu_related:
                                    primary_edu = self._primary_anchors.get('EDUCATION')
                                    # Only stop at education section, don't clear it
                                    if primary_edu is not None and check_idx == primary_edu:
                                        print(f"     → Stopped at primary EDUCATION section at {check_idx}")
                                        break
                                    else:
                                        # Could be the actual education section even if not marked as primary
                                        # Stop here to be safe, don't clear it
                                        print(f"     → Stopped at EDUCATION-related section at {check_idx}: {check_text[:40]}")
                                        break
                                else:
                                    print(f"     → Stopped at section: {check_text[:30]}")
                                    break
                            
                            # AGGRESSIVE: Clear paragraphs that look like sample data
                            # Check for sample names (like "ADIKA MAUL")
                            if re.search(r'^[A-Z][A-Z\s]{5,30}$', check_text_full.strip()):
                                print(f"     → Clearing sample name: {check_text_full[:40]}")
                                paras_to_clear.append(check_para)
                                continue
                            
                            # Check for contact info patterns
                            if re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', check_text_full) and '@' in check_text_full:
                                print(f"     → Clearing contact info: {check_text_full[:40]}")
                                paras_to_clear.append(check_para)
                                continue
                            
                            # Check for "EXPERIENCE" heading (not same as EMPLOYMENT HISTORY)
                            if check_text.strip() == 'EXPERIENCE' and len(check_text) < 15:
                                print(f"     → Clearing duplicate EXPERIENCE heading at {check_idx}")
                                paras_to_clear.append(check_para)
                                continue
                            
                            # Skip empty paragraphs (don't clear them, they're spacing)
                            if not check_text_full.strip():
                                continue
                            
                            # Clear this paragraph
                            paras_to_clear.append(check_para)
                        
                        print(f"     → Clearing {len(paras_to_clear)} paragraphs in EMPLOYMENT section")
                        for p in paras_to_clear:
                            try:
                                p.clear()
                            except:
                                try:
                                    for run in p.runs:
                                        run.text = ''
                                except:
                                    pass
                        
                        # Look for instructional text in the NEXT paragraph (after clearing)
                        next_para = None
                        is_instruction = False
                        
                        if para_idx + 1 < len(doc.paragraphs):
                            next_para = doc.paragraphs[para_idx + 1]
                            next_text = next_para.text.strip().lower()
                            
                            # Check if next paragraph is instructional text
                            is_instruction = any(phrase in next_text for phrase in [
                                'please list', 'list the candidate', 'please provide',
                                'insert employment', 'add employment', 'employment details'
                            ])
                        
                        if is_instruction:
                            print(f"     → Found instructional text: '{next_para.text[:60]}'")
                            
                            experience_data = self.resume_data.get('experience', [])
                            if experience_data and len(experience_data) > 0:
                                print(f"     → Will insert {len(experience_data)} experience entries after heading")
                                
                                # Clear the instructional paragraph
                                for run in next_para.runs:
                                    run.text = ''
                                
                                # CRITICAL: Clear any existing employment content after instructional text
                                paras_to_clear = []
                                for check_idx in range(para_idx + 2, min(para_idx + 50, len(doc.paragraphs))):
                                    check_para = doc.paragraphs[check_idx]
                                    check_text = check_para.text.strip().upper()
                                    
                                    # Stop if we hit another section heading
                                    if any(h in check_text for h in ['EMPLOYMENT', 'WORK HISTORY', 'SKILLS', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(check_text) < 50:
                                        print(f"     → Stopped clearing at section: {check_text[:30]}")
                                        break
                                    
                                    # Clear this paragraph (it's old employment content)
                                    paras_to_clear.append(check_para)
                                
                                print(f"     → Clearing {len(paras_to_clear)} old content paragraphs")
                                for p in paras_to_clear:
                                    try:
                                        for run in p.runs:
                                            run.text = ''
                                    except:
                                        pass
                                
                                # Insert employment blocks after the cleared paragraph
                                last_element = next_para
                                inserted_count = 0
                                for idx, exp in enumerate(experience_data):  # Insert ALL employment entries
                                    print(f"        → Inserting job {idx+1}/{len(experience_data)}: {exp.get('company', 'N/A')[:25]} | {exp.get('role', 'N/A')[:25]}")
                                    block = self._insert_experience_block(doc, last_element, exp)
                                    if block:
                                        last_element = block
                                        inserted_count += 1
                                        print(f"           ✓ Inserted successfully (total: {inserted_count})")
                                    else:
                                        print(f"           ✗ Failed to insert")
                                print(f"     ✅ Successfully inserted {inserted_count} employment entries")
                                
                                self._experience_inserted = True
                                # Remember tail paragraph to place SKILLS after EMPLOYMENT
                                self._employment_tail_para = last_element
                                print(f"  ✅ Inserted employment data after EMPLOYMENT HISTORY heading")
                                replaced_count += 1
                        else:
                            # No instructional text - there's existing employment content
                            # Clear all content after heading until next section
                            print(f"     → No instructional text found, clearing existing employment content")
                            
                            paras_to_clear = []
                            for check_idx in range(para_idx + 1, min(para_idx + 50, len(doc.paragraphs))):
                                check_para = doc.paragraphs[check_idx]
                                check_text = check_para.text.strip().upper()
                                
                                # Stop if we hit another section heading or end of document
                                # CRITICAL: Include EDUCATION to prevent deleting it!
                                if any(h in check_text for h in ['EMPLOYMENT', 'WORK HISTORY', 'EDUCATION', 'SKILLS', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(check_text) < 50:
                                    print(f"     → Stopped clearing at section: {check_text[:30]}")
                                    break
                                
                                # Clear this paragraph (it's old employment content)
                                paras_to_clear.append(check_para)
                            
                            print(f"     → Clearing {len(paras_to_clear)} old employment paragraphs")
                            for p in paras_to_clear:
                                try:
                                    for run in p.runs:
                                        run.text = ''
                                except:
                                    pass
                            
                            # Insert employment blocks after the heading
                            last_element = paragraph
                            inserted_count = 0
                            for idx, exp in enumerate(experience_data):  # Insert ALL employment entries
                                print(f"        → Inserting job {idx+1}/{len(experience_data)}: {exp.get('company', 'N/A')[:25]} | {exp.get('role', 'N/A')[:25]}")
                                block = self._insert_experience_block(doc, last_element, exp)
                                if block:
                                    last_element = block
                                    inserted_count += 1
                                    print(f"           ✓ Inserted successfully (total: {inserted_count})")
                                else:
                                    print(f"           ✗ Failed to insert")
                            print(f"     ✅ Successfully inserted {inserted_count} employment entries")
                            
                            self._experience_inserted = True
                            self._employment_tail_para = last_element
                            print(f"  ✅ Inserted employment data after EMPLOYMENT HISTORY heading (no instruction text)")
                            replaced_count += 1
            
            # Employment placeholder generic patterns (very flexible)
            if not self._experience_inserted:
                # Avoid inserting EXP content inside table cells (e.g., SKILLS table headers/columns)
                try:
                    if self._paragraph_in_table(paragraph):
                        pass
                    else:
                        emp_patterns = [
                            r"<[^>]*list[^>]*candidate['']?s?[^>]*employment[^>]*history[^>]*>",
                            r"<[^>]*employment[^>]*history[^>]*>",
                            r"<[^>]*work[^>]*history[^>]*>",
                            r"<[^>]*professional[^>]*experience[^>]*>",
                            r"<[^>]*career[^>]*(history|experience)[^>]*>",
                            r"<[^>]*history[^>]*(employ|employer|work|career)[^>]*>",
                            r"<[^>]*list[^>]*employment[^>]*history[^>]*>",
                        ]
                        for emp_pat in emp_patterns:
                            if re.search(emp_pat, paragraph.text, re.IGNORECASE):
                                print(f"  💼 Found employment placeholder in paragraph {para_idx}: '{paragraph.text[:60]}'")
                                
                                # Use structured experience data (not sections)
                                experience_data = self.resume_data.get('experience', [])
                                
                                if experience_data and len(experience_data) > 0:
                                    print(f"     → Will replace with {len(experience_data)} experience entries")
                                    
                                    # Clear the placeholder paragraph
                                    self._regex_replace_paragraph(paragraph, emp_pat, '')
                                    
                                    # Insert properly formatted experience blocks
                                    last_element = paragraph
                                    inserted_count = 0
                                    for exp in experience_data:  # Insert ALL employment entries
                                        block = self._insert_experience_block(doc, last_element, exp)
                                        if block:
                                            last_element = block
                                            inserted_count += 1
                                    print(f"     ✅ Successfully inserted {inserted_count} employment entries")
                                    
                                    # CRITICAL: Set flag to prevent duplicate insertion in _add_sections_content
                                    self._experience_inserted = True
                                    # Store tail paragraph for SKILLS placement
                                    self._employment_tail_para = last_element
                                    
                                    print(f"  ✅ Replaced employment placeholder with structured blocks")
                                    replaced_count += 1
                                    break
                                else:
                                    # Fallback: try to use sections data
                                    content = self._find_matching_resume_section('experience', self.resume_data.get('sections', {}))
                                    if content:
                                        bullets = []
                                        for item in content:  # Insert ALL items
                                            if item.strip():
                                                bullets.append('• ' + item.strip().lstrip('•').strip())
                                        self._regex_replace_paragraph(paragraph, emp_pat, '\n'.join(bullets))
                                        # Set flag even in fallback to prevent duplication
                                        self._experience_inserted = True
                                        print(f"  ✅ Regex replaced experience placeholder (fallback)")
                                        replaced_count += 1
                                        break
                except Exception:
                    pass

            # CRITICAL: Check if this is a SUMMARY section heading (only before EMPLOYMENT)
            if not self._summary_inserted:
                para_upper = paragraph.text.strip().upper()
                is_summary_heading = any(h in para_upper for h in [
                    'SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE',
                    'CAREER SUMMARY', 'EXECUTIVE SUMMARY'
                ])
                
                emp_anchor_idx = self._primary_anchors.get('EMPLOYMENT')
                name_idx = getattr(self, '_name_anchor_idx', None)
                # Accept summary headings only if they are near the name (within 10 paras after name) and before employment
                is_position_ok = False
                if name_idx is not None:
                    if emp_anchor_idx is not None:
                        is_position_ok = (para_idx <= name_idx + 10) and (para_idx < emp_anchor_idx)
                    else:
                        is_position_ok = (para_idx <= name_idx + 10)
                else:
                    # Fallback: very early paragraphs only (before employment)
                    is_position_ok = (emp_anchor_idx is None and para_idx < 15) or (emp_anchor_idx is not None and para_idx < min(emp_anchor_idx, 15))
                
                if is_summary_heading and len(paragraph.text.strip()) < 50 and is_position_ok:
                    print(f"  📝 Found SUMMARY heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    
                    summary_text = (self.resume_data.get('summary') or '').strip()
                    summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {}))
                    
                    if summary_text or summary_lines:
                        print(f"     → Will insert summary content after heading")
                        
                        # CRITICAL: Clear ALL content between SUMMARY heading and next section
                        paras_to_clear = []
                        stop_headings = ['EMPLOYMENT', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 
                                        'EDUCATION', 'SKILLS', 'TECHNICAL SKILLS', 'CERTIFICATIONS']
                        
                        for check_idx in range(para_idx + 1, min(para_idx + 30, len(doc.paragraphs))):
                            check_para = doc.paragraphs[check_idx]
                            check_text = check_para.text.strip().upper()
                            
                            # Stop at next major section
                            if any(h in check_text for h in stop_headings) and len(check_text) < 50:
                                print(f"     → Stopped clearing at section: {check_text[:30]}")
                                break
                            
                            paras_to_clear.append(check_para)
                        
                        print(f"     → Clearing {len(paras_to_clear)} paragraphs in SUMMARY section")
                        for p in paras_to_clear:
                            try:
                                for run in p.runs:
                                    run.text = ''
                            except:
                                pass
                        
                        # Insert summary content
                        # Convert summary_text to lines if needed
                        if not summary_lines and summary_text:
                            lines = [line.strip() for line in summary_text.split('\n') if line.strip()]
                            if len(lines) == 1:
                                sentences = re.split(r'\.\s+(?=[A-Z])', summary_text)
                                lines = [s.strip() + ('.' if not s.strip().endswith('.') else '') for s in sentences if s.strip()]
                            summary_lines = lines
                        
                        # Always insert as bullet points (no indentation for summary)
                        if summary_lines:
                            last_para = paragraph
                            for line in summary_lines:
                                txt = (line or '').strip()
                                if not txt:
                                    continue
                                bullet_para = self._insert_paragraph_after(last_para, '')
                                if bullet_para:
                                    # No left indent for summary bullets
                                    run = bullet_para.add_run('• ' + txt.lstrip('•–—-*● '))
                                    run.font.size = Pt(10)
                                    bullet_para.paragraph_format.space_after = Pt(2)
                                    last_para = bullet_para
                        
                        self._summary_inserted = True
                        print(f"  ✅ Inserted summary after SUMMARY heading")
                        replaced_count += 1
            
            # Summary placeholder patterns - flexible (we clear them and insert after name later)
            if not self._summary_inserted:
                summary_patterns = [
                r"<[^>]*summary[^>]*>",
                r"<[^>]*professional[^>]*summary[^>]*>",
                r"<[^>]*profile[^>]*>",
            ]
            for sum_pat in summary_patterns:
                if re.search(sum_pat, paragraph.text, re.IGNORECASE):
                    print(f"  📝 Found summary placeholder in paragraph {para_idx} — clearing and deferring insertion after name")
                    self._regex_replace_paragraph(paragraph, sum_pat, '')
                    replaced_count += 1
                    break

            # Skills placeholder patterns - flexible
            skills_patterns = [
                r"<[^>]*skills[^>]*>",
                r"<[^>]*technical[^>]*skills[^>]*>",
                r"<[^>]*list[^>]*skills[^>]*>",
            ]
            for skl_pat in skills_patterns:
                if re.search(skl_pat, paragraph.text, re.IGNORECASE):
                    skills_list = self.resume_data.get('skills', [])
                    # Always clear the placeholder, but only insert content AFTER employment is inserted
                    print(f"  🧰 Found skills placeholder in paragraph {para_idx} — clearing; will insert after EMPLOYMENT")
                    self._regex_replace_paragraph(paragraph, skl_pat, '')
                    if skills_list and self._experience_inserted:
                        self._insert_skills_bullets(doc, paragraph, skills_list)
                        self._skills_inserted = True
                    replaced_count += 1
                    break

            # SKILLS section heading (respect template order; do not force after EMPLOYMENT)
            if not self._skills_inserted:
                para_upper = paragraph.text.strip().upper()
                is_skills_heading = any(h in para_upper for h in [
                    'SKILLS', 'TECHNICAL SKILLS'
                ])
                # NOTE: We don't gate by primary anchor here because paragraph indices shift after insertions
                # Instead, we rely on the _skills_inserted flag to prevent duplicates
                if is_skills_heading and len(paragraph.text.strip()) < 50:
                    print(f"  🧰 Found SKILLS heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    skills_list = self.resume_data.get('skills', [])
                    if skills_list:
                        paras_to_clear = []
                        stop_headings = ['EMPLOYMENT', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EDUCATION', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']
                        for check_idx in range(para_idx + 1, min(para_idx + 30, len(doc.paragraphs))):
                            check_para = doc.paragraphs[check_idx]
                            check_text = check_para.text.strip().upper()
                            if any(h in check_text for h in stop_headings) and len(check_text) < 50:
                                break
                            paras_to_clear.append(check_para)
                        for p in paras_to_clear:
                            try:
                                for r in p.runs:
                                    r.text = ''
                            except:
                                pass
                        self._insert_skills_bullets(doc, paragraph, skills_list)
                        self._skills_inserted = True
                        replaced_count += 1

            # CRITICAL: Check if this is an EDUCATION section heading
            if not self._education_inserted:
                para_upper = paragraph.text.strip().upper()
                is_edu_heading = any(h in para_upper for h in [
                    'EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND',
                    'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND',
                    'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS', 'ACADEMICS',
                    'EDUCATION/CERTIFICATES', 'EDUCATION / CERTIFICATES'
                ])
                # NOTE: We don't gate by primary anchor here because paragraph indices shift after insertions
                # Instead, we rely on the _education_inserted flag to prevent duplicates
                
                # CRITICAL: Skip placeholder text like "<List candidate's education background>"
                is_placeholder = '<' in para_upper and '>' in para_upper and 'CANDIDATE' in para_upper
                
                # Process EDUCATION headings (skip only if in first 3 paragraphs AND very short)
                # This allows education in various template layouts
                skip_cai = para_idx < 3 and len(paragraph.text.strip()) < 15
                if is_edu_heading and not is_placeholder and not skip_cai and not self._education_inserted:
                    print(f"  🎓 Found EDUCATION heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    print(f"     _education_inserted flag: {self._education_inserted}")
                    
                    # CRITICAL: Preserve and format the heading text  
                    original_heading = paragraph.text.strip().upper()
                    # Keep original heading text (EDUCATION, CERTIFICATES, etc.)
                    if not any(h in original_heading for h in ['EDUCATION', 'CERTIFICATE', 'CREDENTIAL', 'ACADEMIC', 'QUALIFICATION']):
                        original_heading = 'EDUCATION'  # Fallback
                    
                    # Ensure heading is properly formatted (BOLD + UNDERLINE)
                    for run in paragraph.runs:
                        run.bold = True
                        run.underline = True
                        run.font.size = Pt(11)
                    
                    education_data = self.resume_data.get('education', [])
                    print(f"     Initial education_data from resume_data: {len(education_data) if education_data else 0} entries")
                    
                    # ALWAYS get education data from sections if not in structured format
                    if not education_data:
                        section_lines = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                        if section_lines:
                            education_data = self._build_education_from_bullets(section_lines)
                            print(f"     🔄 Built {len(education_data)} education entries from resume sections")
                    
                    # Also check resume_data sections directly
                    if not education_data:
                        sections = self.resume_data.get('sections', {})
                        if 'education' in sections:
                            edu_text = sections['education']
                            if isinstance(edu_text, str):
                                lines = [l.strip() for l in edu_text.split('\n') if l.strip()]
                                education_data = self._build_education_from_bullets(lines)
                                print(f"     🔄 Built {len(education_data)} education entries from sections.education text")
                    
                    print(f"     📊 Final education_data count: {len(education_data) if education_data else 0}")
                    
                    if not education_data or len(education_data) == 0:
                        print(f"     ⚠️  No education data available - WILL STILL INSERT HEADING and mark as processed")
                        # Keep the heading, just don't add content
                        self._education_inserted = True  # Mark as processed even if no data
                        replaced_count += 1  # Count this as a processed section
                        continue  # Skip to next paragraph
                    else:
                        print(f"     → Will insert {len(education_data)} education entries after heading")
                        
                        # Look for instructional text in the NEXT paragraph
                        next_para = None
                        is_instruction = False
                        
                        if para_idx + 1 < len(doc.paragraphs):
                            next_para = doc.paragraphs[para_idx + 1]
                            next_text = next_para.text.strip().lower()
                            
                            # Check if next paragraph is instructional text
                            is_instruction = any(phrase in next_text for phrase in [
                                'please list', 'list the candidate', 'please provide',
                                'insert education', 'add education', 'education details',
                                'educational background', 'academic background'
                            ])
                        
                        if is_instruction:
                            print(f"     → Found instructional text: '{next_para.text[:60]}'")
                            
                            education_data = self.resume_data.get('education', [])
                            if education_data and len(education_data) > 0:
                                print(f"     → Will insert {len(education_data)} education entries after heading")
                                
                                # Clear the instructional paragraph
                                for run in next_para.runs:
                                    run.text = ''
                                
                                # CRITICAL: Clear any existing education content after instructional text
                                paras_to_clear = []
                                for check_idx in range(para_idx + 2, min(para_idx + 50, len(doc.paragraphs))):
                                    check_para = doc.paragraphs[check_idx]
                                    check_text = check_para.text.strip().upper()
                                    
                                    # Stop if we hit another section heading
                                    if any(h in check_text for h in ['EMPLOYMENT', 'WORK HISTORY', 'SKILLS', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(check_text) < 50:
                                        print(f"     → Stopped clearing at section: {check_text[:30]}")
                                        break
                                    
                                    # Clear this paragraph (it's old education content)
                                    paras_to_clear.append(check_para)
                                
                                print(f"     → Clearing {len(paras_to_clear)} old content paragraphs")
                                for p in paras_to_clear:
                                    try:
                                        for run in p.runs:
                                            run.text = ''
                                    except:
                                        pass
                                
                                # Insert education blocks after the cleared paragraph
                                last_element = next_para
                                simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                                if simple_entries and len(simple_entries) == len(education_data):
                                    self._insert_education_bullets(doc, next_para, education_data)  # ALL entries
                                else:
                                    for edu in education_data:  # Insert ALL education entries
                                        block = self._insert_education_block(doc, last_element, edu)
                                        if block:
                                            last_element = block
                                
                                self._education_inserted = True
                                print(f"  ✅ Inserted education data after EDUCATION heading")
                                replaced_count += 1
                        else:
                            # No instructional text - there's existing education content
                            # Clear all content after heading until next section
                            print(f"     → No instructional text found, clearing existing education content")
                            
                            paras_to_clear = []
                            # CRITICAL: Only clear template placeholder text, not actual content
                            # Limit scan range to max 10 paragraphs to prevent clearing employment entries
                            for check_idx in range(para_idx + 1, min(para_idx + 10, len(doc.paragraphs))):
                                check_para = doc.paragraphs[check_idx]
                                check_text = check_para.text.strip()
                                check_text_upper = check_text.upper()
                                
                                # CRITICAL: Stop immediately if we hit SKILLS or any other section
                                if any(h in check_text_upper for h in ['SKILLS', 'EMPLOYMENT', 'WORK HISTORY', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(check_text) < 50:
                                    print(f"     → Stopped clearing at section: {check_text[:30]}")
                                    break
                                
                                # CRITICAL: Only clear if it looks like template placeholder text
                                # Don't clear if it has bullets, dates, or looks like real content
                                is_placeholder = ('<' in check_text and '>' in check_text) or len(check_text) < 5
                                has_bullet = check_text.startswith('•') or check_text.startswith('-') or check_text.startswith('*')
                                has_date = any(char.isdigit() for char in check_text)
                                
                                if is_placeholder:
                                    # Clear obvious placeholders
                                    paras_to_clear.append(check_para)
                                elif not has_bullet and not has_date and len(check_text) < 80:
                                    # Clear short non-content lines (likely template text)
                                    paras_to_clear.append(check_para)
                                else:
                                    # This looks like real content - stop clearing
                                    print(f"     → Stopped clearing at content: {check_text[:40]}")
                                    break
                            
                            print(f"     → Clearing {len(paras_to_clear)} old education paragraphs")
                            for p in paras_to_clear:
                                try:
                                    for run in p.runs:
                                        run.text = ''
                                except:
                                    pass
                            
                            # Insert education blocks after the heading
                            last_element = paragraph
                            simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                            if simple_entries and len(simple_entries) == len(education_data):
                                self._insert_education_bullets(doc, paragraph, education_data)  # ALL entries
                            else:
                                for edu in education_data:  # Insert ALL education entries
                                    block = self._insert_education_block(doc, last_element, edu)
                                    if block:
                                        last_element = block
                            
                            self._education_inserted = True
                            print(f"  ✅ Inserted education data after EDUCATION heading (no instruction text)")
                            replaced_count += 1
            
            # Education placeholder generic patterns - very flexible matching
            if not self._education_inserted:
                edu_patterns = [
                r"<[^>]*list[^>]*candidate['’]?s?[^>]*education[^>]*background[^>]*>",
                r"<[^>]*education[^>]*background[^>]*>",
                r"<[^>]*education[^>]*history[^>]*>",
                r"<[^>]*candidate['’]?s?[^>]*education[^>]*>",
                r"<[^>]*educational[^>]*background[^>]*>",
                r"<[^>]*academic[^>]*background[^>]*>",
                r"<[^>]*academic[^>]*qualifications[^>]*>",
                r"<[^>]*qualifications[^>]*>",
                r"<[^>]*(education|academic)[^>]*>",
                r"\blist\s*candidate(?:['’]s)?\s*education\s*background\b",
                ]
                for edu_pat in edu_patterns:
                    if re.search(edu_pat, paragraph.text, re.IGNORECASE):
                        # If we detected a primary EDUCATION anchor, avoid replacing placeholders that
                        # are far away from the anchor (prevents inserting inside EMPLOYMENT region).
                        if self._primary_anchors.get('EDUCATION') is not None:
                            anchor_idx = self._primary_anchors.get('EDUCATION')
                            # Only allow placeholder replacement near or after the anchor
                            if para_idx < anchor_idx - 2:
                                print(f"  ⏭️  Skipping education placeholder at {para_idx} (before primary anchor {anchor_idx})")
                                continue
                        print(f"  🎓 Found education placeholder in paragraph {para_idx}: '{paragraph.text[:60]}'")
                        
                        # Use structured education data (not sections)
                        education_data = self.resume_data.get('education', [])
                        
                        if education_data and len(education_data) > 0:
                            print(f"     → Will replace with {len(education_data)} education entries")
                            
                            # CRITICAL: Create EDUCATION heading instead of clearing placeholder
                            paragraph.clear()
                            run = paragraph.add_run('EDUCATION')
                            run.bold = True
                            run.underline = True  # UNDERLINE
                            run.font.size = Pt(12)
                            run.font.all_caps = True  # CAPITAL
                            paragraph.paragraph_format.space_before = Pt(12)
                            paragraph.paragraph_format.space_after = Pt(6)
                            print(f"     ✅ Created EDUCATION heading: BOLD, UNDERLINED, CAPITAL")
                            
                            # If entries are simple (no institution/details), insert as bullets to match layout
                            simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                            if simple_entries and len(simple_entries) == len(education_data):
                                self._insert_education_bullets(doc, paragraph, education_data)  # ALL entries
                            else:
                                # Insert properly formatted education blocks
                                last_element = paragraph
                                for edu in education_data:  # Insert ALL education entries
                                    block = self._insert_education_block(doc, last_element, edu)
                                    if block:
                                        last_element = block
                        
                            # CRITICAL: Set flag to prevent duplicate insertion in _add_sections_content
                            self._education_inserted = True
                            
                            print(f"  ✅ Replaced education placeholder with EDUCATION heading + structured blocks")
                            replaced_count += 1
                            break
                        else:
                            # Fallback: try to use sections data
                            content = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                            if content:
                                bullets = []
                                for item in content:  # Insert ALL items
                                    if item.strip():
                                        bullets.append('• ' + item.strip().lstrip('•').strip())
                                self._regex_replace_paragraph(paragraph, edu_pat, '\n'.join(bullets))
                                # Set flag even in fallback to prevent duplication
                                self._education_inserted = True
                                print(f"  ✅ Regex replaced education placeholder (fallback)")
                                replaced_count += 1
                                break
        
        print(f"\n Replaced {replaced_count} placeholders in paragraphs")
        
        # Fallback creation: SUMMARY (only if template has no SUMMARY heading)
        if not self._summary_inserted:
            if self._primary_anchors.get('SUMMARY') is not None:
                print(f"  SUMMARY heading exists in template; skipping fallback creation")
            else:
                summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {})) or []
                summary_text = (self.resume_data.get('summary') or '').strip()
                if summary_lines or summary_text:
                    print(f"  Creating SUMMARY section after candidate name")
                    
                    # Use the pre-computed name anchor (avoids CAI CONTACT section)
                    anchor_idx = getattr(self, '_name_anchor_idx', None)
                    
                    # Fallback: search for name placeholder patterns if anchor wasn't found
                    if anchor_idx is None:
                        print(f"  Name anchor not found, searching for placeholder...")
                        name_placeholder_patterns = [
                            r'<\s*[Cc]andidate[^>]*[Nn]ame[^>]*>',
                            r'<\s*[Nn]ame\s*>',
                            r'<\s*[Ff]ull\s*[Nn]ame\s*>',
                            r'<\s*YOUR\s*NAME\s*>',
                            r'<[^>]*LAWSON[^>]*>',
                            r'<[^>]*PAULA[^>]*>',
                        ]
                        for idx, p in enumerate(doc.paragraphs):
                            if idx < 10:  # Skip CAI CONTACT area - increased from 5 to 10
                                continue
                            # Also skip if in CAI CONTACT section
                            if idx < 20 and 'CAI CONTACT' in (p.text or '').upper():
                                continue
                            t = (p.text or '').strip()
                            for pat in name_placeholder_patterns:
                                if re.search(pat, t):
                                    anchor_idx = idx
                                    print(f"  Found name placeholder at paragraph {idx}: '{t}'")
                                    break
                            if anchor_idx is not None:
                                break
                
                    # Strategy: Place SUMMARY right after the candidate name placeholder
                    if anchor_idx is not None and anchor_idx >= 10 and anchor_idx < len(doc.paragraphs):
                        # Name anchor found in main content area (not CAI CONTACT) - increased from 5 to 10
                        anchor_para = doc.paragraphs[anchor_idx]
                        print(f"  Inserting SUMMARY after candidate name at paragraph {anchor_idx}")
                    else:
                        # Fallback: use paragraph before EMPLOYMENT if name not found
                        emp_idx = self._primary_anchors.get('EMPLOYMENT')
                        if emp_idx is not None and emp_idx > 0:
                            anchor_para = doc.paragraphs[emp_idx - 1]
                            print(f"  Fallback: Inserting SUMMARY before EMPLOYMENT at paragraph {emp_idx - 1}")
                        else:
                            # Skip SUMMARY insertion if no safe anchor found
                            print(f"  No safe anchor found for SUMMARY - skipping insertion")
                            self._summary_inserted = True  # Mark as inserted to prevent retry
                            anchor_para = None
                
                    if anchor_para is None:
                        # Don't insert if no safe location found
                        pass
                    else:
                        # Insert blank line for spacing
                        blank = self._insert_paragraph_after(anchor_para, '')
                        if blank is None:
                            blank = anchor_para
                        
                        # Insert SUMMARY heading with underline
                        heading = self._insert_paragraph_after(blank, 'SUMMARY')
                        if heading is None:
                            heading = anchor_para
                        for r in heading.runs:
                            r.bold = True
                            r.underline = True
                            r.font.size = Pt(11)
                        
                        # Insert summary content below the heading
                        # Convert summary_text to lines if needed
                        if not summary_lines and summary_text:
                            lines = [line.strip() for line in summary_text.split('\n') if line.strip()]
                            if len(lines) == 1:
                                sentences = re.split(r'\.\s+(?=[A-Z])', summary_text)
                                lines = [s.strip() + ('.' if not s.strip().endswith('.') else '') for s in sentences if s.strip()]
                            summary_lines = lines
                        
                        # Always insert as bullets (no indentation for summary)
                        if summary_lines:
                            last_para = heading
                            for line in summary_lines:
                                txt = (line or '').strip()
                                if not txt:
                                    continue
                                bullet_para = self._insert_paragraph_after(last_para, '')
                                if bullet_para:
                                    # No left indent for summary bullets
                                    run = bullet_para.add_run('• ' + txt.lstrip('•–—-*● '))
                                    run.font.size = Pt(10)
                                    bullet_para.paragraph_format.space_after = Pt(2)
                                    last_para = bullet_para
                            print(f"  Inserted SUMMARY heading + bullets")
                        
                        self._summary_inserted = True

        # Fallback creation: SKILLS (place AFTER EMPLOYMENT)
        if not self._skills_inserted:
            skills_list = self.resume_data.get('skills', []) or []
            if skills_list:
                # Anchor: after the last EMPLOYMENT paragraph if available; else after EMPLOYMENT heading; else end
                if hasattr(self, '_employment_tail_para') and self._employment_tail_para is not None:
                    anchor_para = self._employment_tail_para
                elif self._primary_anchors.get('EMPLOYMENT') is not None:
                    anchor_para = doc.paragraphs[self._primary_anchors.get('EMPLOYMENT')]
                else:
                    anchor_para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph('')
                # Clean up any stray SKILLS headings that occur before anchor
                try:
                    anchor_idx = None
                    for idx, p in enumerate(doc.paragraphs):
                        if p is anchor_para:
                            anchor_idx = idx
                            break
                    if anchor_idx is not None:
                        for idx in range(anchor_idx - 1, -1, -1):
                            t = (doc.paragraphs[idx].text or '').strip().upper()
                            if t in ('SKILLS', 'TECHNICAL SKILLS'):
                                self._delete_paragraph(doc.paragraphs[idx])
                except Exception:
                    pass
                heading = self._insert_paragraph_after(anchor_para, 'SKILLS')
                if heading is None:
                    heading = anchor_para
                for r in heading.runs:
                    r.bold = True
                    r.underline = True
                    r.font.size = Pt(11)
                self._insert_skills_bullets(doc, heading, skills_list)
                self._skills_inserted = True

        # CRITICAL: Final cleanup - remove any orphaned bullets that appear after section headings
        # This handles cases where resume parsing left stray content
        print(f"\n🧹 Final cleanup: Removing orphaned content...")
        sections_found = {}
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip().upper()
            
            # Track section positions
            if any(h in para_text for h in ['EMPLOYMENT HISTORY', 'EDUCATION', 'SKILLS', 'SUMMARY']) and len(para_text) < 50:
                sections_found[para_text[:20]] = para_idx
        
        # If we inserted employment and education, clear any bullets that appear after education
        if self._experience_inserted and self._education_inserted and 'EDUCATION' in str(sections_found):
            print(f"     → Checking for orphaned bullets after EDUCATION section...")
            # This is handled by the section clearing logic above
        
        # STEP 3: Process remaining table content (non-skills tables)
        print(f"\n🔍 STEP 3: Processing non-skills table content...")
        other_table_replaced = 0
        
        for table_idx, table in enumerate(doc.tables):
            # Skip skills tables (already processed in STEP 1)
            if self._is_skills_table(table):
                print(f"  ⏭️  Skipping already-processed skills table at index {table_idx}")
                continue
            else:
                # Regular placeholder replacement in non-skills tables
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # 1) Simple replacements
                            for key, value in replacements.items():
                                if self._text_contains(paragraph.text, key):
                                    other_table_replaced += self._replace_in_paragraph(paragraph, key, value)

                            # 1.5) EDUCATION heading inside table (check if heading OR placeholder exists)
                            if not self._education_inserted:
                                para_text = (paragraph.text or '').strip()
                                heading_text = para_text.upper()
                                
                                # Check if this cell contains EDUCATION heading or placeholder
                                has_heading = any(h in heading_text for h in ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND', 'ACADEMICS', 'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS', 'EDUCATION/CERTIFICATES', 'EDUCATION / CERTIFICATES'])
                                has_placeholder = bool(re.search(r"<[^>]*list[^>]*candidate['']?s?[^>]*education[^>]*background[^>]*>", para_text, re.IGNORECASE))
                                
                                if has_heading or has_placeholder:
                                    print(f"  🎓 Found EDUCATION in TABLE cell (heading={has_heading}, placeholder={has_placeholder})")
                                    print(f"     Cell text: '{para_text[:80]}'")
                                    
                                    education_data = self.resume_data.get('education', [])
                                    print(f"     → resume_data.education has {len(education_data)} entries")
                                    if not education_data:
                                        sect = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                                        print(f"     → sections.education has {len(sect) if sect else 0} lines")
                                        if sect:
                                            education_data = self._build_education_from_bullets(sect)
                                            print(f"     → Built {len(education_data)} education entries from sections")
                                    
                                    if education_data:
                                        print(f"     → Will insert {len(education_data)} education entries")
                                        # Clear the entire paragraph text (heading + placeholder)
                                        for run in paragraph.runs:
                                            run.text = ''
                                        # Rewrite just the heading
                                        if paragraph.runs:
                                            paragraph.runs[0].text = 'EDUCATION'
                                            paragraph.runs[0].bold = True
                                            paragraph.runs[0].font.size = Pt(11)
                                        
                                        # Clear any following raw content within the cell
                                        self._delete_following_bullets(paragraph, max_scan=80)
                                        self._delete_next_table(paragraph)
                                        
                                        # Insert content - ALL education entries
                                        simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                                        if simple_entries and len(simple_entries) == len(education_data):
                                            self._insert_education_bullets(doc, paragraph, education_data)
                                            print(f"     ✅ Inserted {len(education_data)} education bullets")
                                        else:
                                            last_element = paragraph
                                            for edu in education_data:
                                                blk = self._insert_education_block(doc, last_element, edu)
                                                if blk:
                                                    last_element = blk
                                            print(f"     ✅ Inserted {len(education_data)} education blocks")
                                        
                                        self._education_inserted = True
                                        other_table_replaced += 1
                                    else:
                                        print(f"     ⚠️  No education data available to insert")
                                        # Still clear the placeholder even if no data
                                        for run in paragraph.runs:
                                            run.text = ''
                                        if paragraph.runs:
                                            paragraph.runs[0].text = 'EDUCATION'
                                            paragraph.runs[0].bold = True
                                        self._education_inserted = True

                            # 2) SUMMARY placeholder inside table
                            if not self._summary_inserted:
                                summary_patterns = [r"<[^>]*summary[^>]*>", r"<[^>]*professional[^>]*summary[^>]*>", r"<[^>]*profile[^>]*>"]
                                for sum_pat in summary_patterns:
                                    if re.search(sum_pat, paragraph.text, re.IGNORECASE):
                                        summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {}))
                                        summary_text = (self.resume_data.get('summary') or '').strip()
                                        if summary_lines or summary_text:
                                            print(f"  📝 Found summary placeholder in TABLE cell")
                                            self._regex_replace_paragraph(paragraph, sum_pat, '')
                                            if summary_lines:
                                                self._insert_skills_bullets(doc, paragraph, summary_lines)
                                            else:
                                                sp = self._insert_paragraph_after(paragraph, summary_text)
                                                if sp:
                                                    for run in sp.runs:
                                                        run.font.size = Pt(10)
                                            self._summary_inserted = True
                                            table_replaced += 1
                                            break

                            # 3) SKILLS placeholder inside table
                            if not self._skills_inserted:
                                skills_patterns = [r"<[^>]*skills[^>]*>", r"<[^>]*technical[^>]*skills[^>]*>", r"<[^>]*list[^>]*skills[^>]*>"]
                                for skl_pat in skills_patterns:
                                    if re.search(skl_pat, paragraph.text, re.IGNORECASE):
                                        skills_list = self.resume_data.get('skills', [])
                                        if skills_list:
                                            print(f"  🧰 Found skills placeholder in TABLE cell")
                                            self._regex_replace_paragraph(paragraph, skl_pat, '')
                                            self._insert_skills_bullets(doc, paragraph, skills_list)
                                            self._skills_inserted = True
                                            other_table_replaced += 1
                                        else:
                                            self._regex_replace_paragraph(paragraph, skl_pat, '')
                                        break

                            # 4) EDUCATION placeholder inside table
                            if not self._education_inserted:
                                edu_patterns = [
                                    r"<[^>]*list[^>]*candidate['’]?s?[^>]*education[^>]*background[^>]*>",
                                    r"<[^>]*education[^>]*background[^>]*>",
                                    r"<[^>]*education[^>]*history[^>]*>",
                                    r"<[^>]*candidate['’]?s?[^>]*education[^>]*>",
                                    r"<[^>]*educational[^>]*background[^>]*>",
                                    r"<[^>]*academic[^>]*background[^>]*>",
                                    r"<[^>]*academic[^>]*qualifications[^>]*>",
                                    r"<[^>]*qualifications[^>]*>",
                                    r"<[^>]*(education|academic)[^>]*>",
                                ]
                                for edu_pat in edu_patterns:
                                    if re.search(edu_pat, paragraph.text, re.IGNORECASE):
                                        print(f"  🎓 Found education placeholder in TABLE cell")
                                        education_data = self.resume_data.get('education', [])
                                        if not education_data:
                                            sect = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                                            if sect:
                                                education_data = self._build_education_from_bullets(sect)
                                        if education_data:
                                            self._regex_replace_paragraph(paragraph, edu_pat, '')
                                            simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                                            if simple_entries and len(simple_entries) == len(education_data):
                                                self._insert_education_bullets(doc, paragraph, education_data)
                                            else:
                                                last_element = paragraph
                                                for edu in education_data:  # Insert ALL entries
                                                    blk = self._insert_education_block(doc, last_element, edu)
                                                    if blk:
                                                        last_element = blk
                                            self._education_inserted = True
                                            other_table_replaced += 1
                                            break
        
        print(f"✓ Processed {other_table_replaced} non-skills table entries")
        print(f"\n{'='*70}")
        print(f"📊 FORMATTING SUMMARY")
        print(f"{'='*70}")
        print(f"  • Skills table entries: {table_replaced}")
        print(f"  • Paragraph replacements: {replaced_count}")
        print(f"  • Other table entries: {other_table_replaced}")
        
        # Replace in headers/footers
        header_footer_replaced = 0
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                for key, value in replacements.items():
                    if self._text_contains(paragraph.text, key):
                        header_footer_replaced += self._replace_in_paragraph(paragraph, key, value)
            
            # Footer
            for paragraph in section.footer.paragraphs:
                for key, value in replacements.items():
                    if self._text_contains(paragraph.text, key):
                        header_footer_replaced += self._replace_in_paragraph(paragraph, key, value)
        
        if header_footer_replaced > 0:
            print(f"✓ Replaced {header_footer_replaced} placeholders in headers/footers")
        
        # CLEANUP: Remove excessive empty paragraphs and fix spacing
        print(f"\n🧹 Cleaning up empty paragraphs and fixing spacing...")
        self._cleanup_empty_paragraphs(doc)
        
        # Add sections content
        sections_added = self._add_sections_content(doc)
        print(f"✓ Added {sections_added} sections")
        
        # Final sweep: clear leftover instructional phrases globally
        self._clear_instruction_phrases(doc)
        
        # PHASE 2: Add only truly missing candidate sections (no duplicates)
        try:
            print(f"\n🔍 PHASE 2: Adding truly uncovered candidate sections...")
            
            # Get sections that are actually covered by existing template structure
            template_covered_sections = self._get_existing_template_covered_sections()
            
            # Find candidate sections not covered by existing template structure
            uncovered_sections = self._get_truly_uncovered_candidate_sections(template_covered_sections)
            
            if uncovered_sections:
                print(f"  📋 Truly uncovered sections to add: {list(uncovered_sections.keys())}")
                additional_sections_added = self._insert_non_duplicate_additional_sections(doc, uncovered_sections)
                print(f"  ✅ Added {additional_sections_added} additional sections without duplicates")
            else:
                print(f"  ✓ All candidate sections covered by existing template structure")
                
            # Final content verification
            self._verify_complete_content_preservation()
            
        except Exception as e:
            print(f"⚠️  Additional sections insertion error: {e}")

        # Add any missing sections from candidate resume that don't exist in template
        print(f"\n🔍 Checking for additional candidate sections...")
        try:
            added_sections = self._add_dynamic_candidate_sections(doc)
            if added_sections > 0:
                print(f"  ✅ Added {added_sections} dynamic sections from candidate resume")
            else:
                print(f"  ℹ️  No additional sections needed")
        except Exception as e:
            print(f"  ⚠️  Error adding dynamic sections: {e}")
        
        # Save output
        output_docx = self.output_path.replace('.pdf', '.docx')
        # Enforce justified alignment across all paragraphs before saving
        try:
            for para in doc.paragraphs:
                if para is not None and para.text is not None:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except Exception:
            pass
        
        # CRITICAL: Verify EDUCATION section is still present before saving
        if self._education_inserted:
            education_found = False
            for para in doc.paragraphs:
                if 'EDUCATION' in para.text.upper() and len(para.text.strip()) < 50:
                    education_found = True
                    print(f"✅ EDUCATION section verified in document before save: '{para.text}'")
                    break
            if not education_found:
                print(f"⚠️  WARNING: EDUCATION section was marked as inserted but not found in document!")
                print(f"   This indicates the section was deleted during processing.")
        
        doc.save(output_docx)
        
        # DISABLED: COM post-processing was corrupting already-inserted content
        # The main document is already fully processed by python-docx
        # COM would only be needed for text boxes/shapes, which we don't use in this template
        print("ℹ️  Skipping COM post-processing (not needed - content already inserted directly)")

        print(f"\n✅ Successfully created formatted document!")
        print(f"📁 Saved to: {output_docx}\n")
        
        # Optionally convert to PDF
        if self.output_path.endswith('.pdf'):
            print("📄 Converting to PDF...")
            if self._convert_to_pdf(output_docx, self.output_path):
                print(f"✓ PDF created: {self.output_path}")
                # Keep both docx and pdf
            else:
                print("⚠️  PDF conversion failed, keeping .docx file")
        
        return True

    # Helper: insert a new paragraph directly after a given paragraph
    def _insert_paragraph_after(self, paragraph, text):
        try:
            new_p = OxmlElement('w:p')
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            new_para.add_run(text)
            # Ensure justified alignment for any new paragraph inserted
            try:
                new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            except Exception:
                pass
            return new_para
        except Exception:
            # Fallback: append to document if direct insert fails
            p = paragraph._parent.add_paragraph(text)
            try:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            except Exception:
                pass
            return p
    
    def _add_right_tab(self, paragraph, pos_twips=9360):
        """Add a right-aligned tab stop to a paragraph at the given twips position (1 inch = 1440 twips)."""
        try:
            pPr = paragraph._p.get_or_add_pPr()
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = OxmlElement('w:tabs')
                pPr.append(tabs)
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), str(pos_twips))
            tabs.append(tab)
        except Exception:
            # If this fails, the text will still render; right text just won't align via tab stop
            pass

    def _delete_paragraph(self, paragraph):
        """Safely delete a paragraph from the document body."""
        try:
            p = paragraph._element
            parent = p.getparent()
            parent.remove(p)
        except Exception:
            pass

    def _insert_skills_bullets(self, doc, after_paragraph, skills_list):
        """Insert a simple bullet list of skills after the given paragraph.
        Accepts list of strings or objects with 'name' field.
        Returns the last inserted paragraph.
        """
        try:
            last = after_paragraph
            count = 0
            # Insert ALL skills from candidate resume
            for skill in (skills_list or []):
                name = skill if isinstance(skill, str) else (skill.get('name', '') if isinstance(skill, dict) else str(skill))
                name = (name or '').strip()
                if not name:
                    continue
                p = self._insert_paragraph_after(last, '')
                if p is None:
                    break
                p.paragraph_format.left_indent = Inches(0.25)
                try:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                except Exception:
                    pass
                run = p.add_run('• ' + name.lstrip('•–—-*● '))
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                last = p
                count += 1
            if count > 0:
                print(f"    ✓ Inserted {count} skill bullets")
            return last
        except Exception as e:
            print(f"    ⚠️  Error inserting skills bullets: {e}")
            return after_paragraph

    def _insert_education_bullets(self, doc, after_paragraph, education_list):
        """Insert education as bullets like 'High School Graduation 1986.'
        Used when entries have no institution/details to match desired layout.
        """
        try:
            last = after_paragraph
            count = 0
            for edu in (education_list or []):  # Insert ALL education entries
                degree = (edu.get('degree') or '').strip()
                year = self._clean_duration((edu.get('year') or '').strip())
                inst = (edu.get('institution') or '').strip()
                details = edu.get('details') or []
                # Only bullet-render when institution and details are empty (simple case)
                if inst or (details and len(details) > 0):
                    continue
                text_parts = []
                if degree:
                    text_parts.append(degree)
                if year:
                    # Append year after a space to match example
                    text_parts.append(year)
                line = ' '.join(text_parts).strip()
                if not line:
                    continue
                p = self._insert_paragraph_after(last, '')
                if p is None:
                    break
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run('• ' + line)
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                last = p
                count += 1
            if count > 0:
                print(f"    ✓ Inserted {count} education bullets (simple mode)")
            return last
        except Exception as e:
            print(f"    ⚠️  Error inserting education bullets: {e}")
            return after_paragraph

    # ===== CAI CONTACT PERSISTENCE AND INSERTION =====
    def _cai_store_path(self):
        """Return a stable file path to store CAI contact details."""
        home = os.path.expanduser("~")
        return os.path.join(home, ".resume_formatter_cai_contact.json")

    def _load_cai_contact(self, proposed=None, edit=False):
        """Load CAI contact from disk. If edit=True and proposed provided, overwrite and save.
        Structure: {"name": str, "phone": str, "email": str}
        """
        path = self._cai_store_path()
        stored = {}
        try:
            if os.path.exists(path):
                with open(path, 'r', encoding='utf-8') as f:
                    stored = json.load(f) or {}
        except Exception:
            stored = {}

        if edit and isinstance(proposed, dict) and any(proposed.get(k) for k in ("name", "phone", "email")):
            data = {
                "name": (proposed.get("name") or stored.get("name") or ""),
                "phone": (proposed.get("phone") or stored.get("phone") or ""),
                "email": (proposed.get("email") or stored.get("email") or ""),
            }
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
            return data

        # No edit: fall back to stored, else proposed, else empty
        if stored:
            return stored
        if isinstance(proposed, dict):
            return {
                "name": proposed.get("name", ""),
                "phone": proposed.get("phone", ""),
                "email": proposed.get("email", ""),
            }
        return {"name": "", "phone": "", "email": ""}

    def _ensure_cai_contact(self, doc):
        """Ensure the CAI CONTACT section exists and is filled from persistent storage.
        Will not change stored values unless an explicit edit flag is provided via
        resume_data['edit_cai_contact'] or template_analysis['edit_cai_contact'].
        """
        edit_flag = bool(self.resume_data.get('edit_cai_contact') or (self.template_analysis or {}).get('edit_cai_contact'))
        proposed = self.resume_data.get('cai_contact', {})
        cai = self._load_cai_contact(proposed=proposed, edit=edit_flag)

        # Find existing CAI CONTACT heading
        heading_idx = None
        for idx, p in enumerate(doc.paragraphs):
            if (p.text or '').strip().upper() == 'CAI CONTACT':
                heading_idx = idx
                break

        if heading_idx is None:
            # Do NOT create CAI CONTACT unless explicitly requested via edit flag
            if not edit_flag:
                print("  ⏭️  No 'CAI CONTACT' heading in template; skipping CAI contact insertion")
                return
            # Explicitly requested: create near the top
            anchor = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph("")
            heading = self._insert_paragraph_after(anchor, 'CAI CONTACT')
            if heading is None:
                heading = doc.add_paragraph('CAI CONTACT')
            for r in heading.runs:
                r.bold = True
                r.font.size = Pt(11)
            # Write lines under heading
            self._write_cai_contact_block(heading, cai)
        else:
            # CAI CONTACT heading exists in template
            heading = doc.paragraphs[heading_idx]
            
            if not edit_flag:
                # No edit requested - leave template CAI CONTACT section completely unchanged
                print("  ⏭️  CAI CONTACT exists in template but edit_flag not set; leaving unchanged")
                return
            
            # Edit flag is set - ADD new contact info BELOW existing template content
            print("  ✏️  Edit CAI Contact enabled - adding new contact info below template defaults")
            
            # Find the last paragraph of the CAI CONTACT section
            # Scan only within CAI CONTACT - stop at empty line or candidate name placeholder
            last_cai_para = heading
            for j in range(1, 15):  # Scan up to 15 paragraphs after heading
                k = heading_idx + j
                if k >= len(doc.paragraphs):
                    break
                txt = (doc.paragraphs[k].text or '').strip()
                upper = txt.upper()
                
                # Stop at candidate name placeholder (indicates start of main content)
                if '<' in txt and '>' in txt and any(word in upper for word in ['NAME', 'CANDIDATE', 'PAULA', 'LAWSON']):
                    print(f"  🛑 Stopped at candidate name placeholder at para {k}")
                    break
                
                # Stop at next major section headings
                if any(kw in upper for kw in ['EMPLOYMENT HISTORY', 'WORK EXPERIENCE', 'EDUCATION', 'SUMMARY', 'SKILLS']) and len(txt) < 50:
                    print(f"  🛑 Stopped at section heading '{txt}' at para {k}")
                    break
                
                # Stop at multiple consecutive empty lines (indicates section break)
                if not txt:
                    # Check if next line is also empty or a major section
                    if k + 1 < len(doc.paragraphs):
                        next_txt = (doc.paragraphs[k + 1].text or '').strip()
                        if not next_txt or any(kw in next_txt.upper() for kw in ['EMPLOYMENT', 'SUMMARY', 'EDUCATION']):
                            print(f"  🛑 Stopped at empty line break at para {k}")
                            break
                
                # This is still part of CAI CONTACT section
                last_cai_para = doc.paragraphs[k]
            
            # Insert new contact block after the last CAI CONTACT paragraph
            print(f"  📍 Inserting edited CAI contact after paragraph {heading_idx + (doc.paragraphs.index(last_cai_para) - heading_idx if last_cai_para in doc.paragraphs else 0)}")
            self._write_cai_contact_block(last_cai_para, cai)

    def _write_cai_contact_block(self, heading_para, cai):
        """Write CAI contact lines under the given heading paragraph."""
        name = (cai.get('name') or '').strip()
        phone = (cai.get('phone') or '').strip()
        email = (cai.get('email') or '').strip()

        p_name = self._insert_paragraph_after(heading_para, name or '')
        if p_name is not None:
            for r in p_name.runs:
                r.bold = True
                r.font.size = Pt(10)

        if phone:
            p_phone = self._insert_paragraph_after(p_name or heading_para, f"Phone:  {phone}")
            if p_phone is not None:
                for r in p_phone.runs:
                    r.font.size = Pt(10)

        if email:
            p_email = self._insert_paragraph_after(p_phone or p_name or heading_para, f"Email:  {email}")
            if p_email is not None:
                for r in p_email.runs:
                    r.font.size = Pt(10)
    
    def _insert_experience_block(self, doc, after_paragraph, exp_data):
        """Insert a structured 2-column experience block"""
        try:
            # Get parsed data from resume parser
            company = exp_data.get('company', '')
            role = exp_data.get('role', '')
            duration = exp_data.get('duration', '')
            details = exp_data.get('details', [])
            
            # Fallback: if company/role not parsed, try to extract from title
            if not company and not role:
                title = exp_data.get('title', '')
                company, role = self._parse_company_role(title)
            
            # CRITICAL: Remove date fragments from company/role
            # Sometimes dates like "City – 08/ 06/" end up in company field
            import re
            if company:
                # Remove patterns like "City – 08/ 06/" or "– 04/" etc
                company = re.sub(r'\s*[–-]\s*\d{2}/\s*\d{2}/?\s*.*?$', '', company)
                company = re.sub(r'\s*City\s*[–-]\s*\d{2}/.*?$', '', company)
                company = company.strip(' ,–-')
            
            if role:
                role = re.sub(r'\s*[–-]\s*\d{2}/\s*\d{2}/?/\s*.*?$', '', role)
                role = re.sub(r'\s*City\s*[–-]\s*\d{2}/.*?$', '', role)
                role = role.strip(' ,–-')
            
            # Clean up duration format
            duration_clean = self._clean_duration(duration)
            
            # Build header line using a right-aligned tab instead of a table
            header_para = self._insert_paragraph_after(after_paragraph, '')
            # Right tab at ~6.5" (letter page width minus 1" margins), 1 inch = 1440 twips
            self._add_right_tab(header_para, pos_twips=9360)
            try:
                header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            except Exception:
                pass

            # CRITICAL: Company name ALWAYS goes on first line if available
            # Role goes on second line if both exist
            # If only one exists, it goes on first line
            
            if company and role:
                # CASE 1: Both company and role exist
                # Line 1: Company (bold) + dates
                # Truncate if too long to prevent date wrapping
                display_company = company
                if len(display_company) > 70:
                    display_company = display_company[:67] + '...'
                
                left_run = header_para.add_run(display_company)
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = True  # Make years bold
                    dur_run.font.size = Pt(10)
                header_para.paragraph_format.space_after = Pt(0)
                header_para.paragraph_format.keep_together = True
                
                # Line 2: Role (bold)
                role_para = self._insert_paragraph_after(header_para, '')
                role_run = role_para.add_run(role)
                role_run.bold = True
                role_run.font.size = Pt(10)
                role_para.paragraph_format.space_after = Pt(0)
                try:
                    role_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                except Exception:
                    pass
                last_para = role_para
                
            elif company:
                # CASE 2: Only company exists
                # Truncate if too long to prevent date wrapping
                display_company = company
                if len(display_company) > 70:
                    display_company = display_company[:67] + '...'
                
                left_run = header_para.add_run(display_company)
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = True  # Make years bold
                    dur_run.font.size = Pt(10)
                header_para.paragraph_format.space_after = Pt(0)
                header_para.paragraph_format.keep_together = True
                last_para = header_para
                
            elif role:
                # CASE 3: Only role exists
                # Truncate if too long to prevent date wrapping
                display_role = role
                if len(display_role) > 70:
                    display_role = display_role[:67] + '...'
                
                left_run = header_para.add_run(display_role)
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = True  # Make years bold
                    dur_run.font.size = Pt(10)
                header_para.paragraph_format.space_after = Pt(0)
                header_para.paragraph_format.keep_together = True
                last_para = header_para
                
            else:
                # CASE 4: Neither exists (fallback)
                left_run = header_para.add_run('Experience')
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = True  # Make years bold
                    dur_run.font.size = Pt(10)
                header_para.paragraph_format.space_after = Pt(0)
                last_para = header_para

            # Add details as individual bullet paragraphs
            if details:
                # Don't limit bullets - include ALL details from resume
                for detail in details:
                    txt = (detail or '').strip()
                    if not txt:
                        continue
                    p = self._insert_paragraph_after(last_para, '')
                    p.paragraph_format.left_indent = Inches(0.25)
                    try:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    except Exception:
                        pass
                    run = p.add_run('• ' + txt.lstrip('•–—-*● '))
                    run.font.size = Pt(10)
                    last_para = p
            
            return last_para
            
        except Exception as e:
            print(f"  ⚠️  Error inserting experience block: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _insert_education_block(self, doc, after_paragraph, edu_data):
        """Insert a structured 2-column education block"""
        try:
            # Get parsed data from resume parser
            degree = edu_data.get('degree', '').strip()
            institution = edu_data.get('institution', '').strip()
            year = edu_data.get('year', '').strip()
            details = edu_data.get('details', [])
            
            # DEBUG: Show what we received
            print(f"      📚 Education data: degree='{degree[:50] if degree else 'EMPTY'}', institution='{institution[:30] if institution else 'EMPTY'}', year='{year or 'EMPTY'}'")
            
            # CRITICAL: If we have neither degree nor institution, skip this entry
            if not degree and not institution:
                print(f"      ⚠️  Skipping empty education entry")
                return after_paragraph
            
            # Fallback: if institution not parsed, try to extract from degree or details
            if not institution:
                institution = self._extract_institution(degree, details)
                if institution:
                    print(f"      ✓ Extracted institution from degree: '{institution[:30]}'")
            
            # Clean up year format
            year_clean = self._clean_duration(year)
            
            # Build header line as paragraph with right-aligned tab (no tables)
            header_para = self._insert_paragraph_after(after_paragraph, '')
            if header_para is None:
                print(f"      ❌ Failed to insert paragraph")
                return after_paragraph
                
            self._add_right_tab(header_para, pos_twips=9360)
            try:
                header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            except Exception:
                pass
            
            # Parse degree to separate degree type from field
            # Handle both formats:
            # 1. "Master of Science : Leadership" (colon separator)
            # 2. "Master of Science in Data Science" (in separator)
            # 3. "Bachelor of Technology in Computer Science" (in separator)
            
            degree_type = degree
            field_and_institution = institution
            field = ''
            
            if ':' in degree:
                # Format: "Master of Science : Leadership"
                parts = degree.split(':', 1)
                degree_type = parts[0].strip()  # "Master of Science"
                field = parts[1].strip() if len(parts) > 1 else ''  # "Leadership"
                print(f"      ✂️  Split at colon: LEFT='{degree_type}' | Field='{field}'")
            
            elif ' in ' in degree.lower():
                # Format: "Master of Science in Data Science"
                # Find the position of " in " (case-insensitive)
                lower_degree = degree.lower()
                in_pos = lower_degree.find(' in ')
                if in_pos > 0:
                    degree_type = degree[:in_pos].strip()  # "Master of Science"
                    field = degree[in_pos + 4:].strip()     # "Data Science"
                    print(f"      ✂️  Split at 'in': LEFT='{degree_type}' | Field='{field}'")
            
            else:
                print(f"      ℹ️  No split: Using full degree as LEFT='{degree_type}'")
            
            # Combine field with institution
            if field and institution:
                field_and_institution = f"{field} - {institution}"
            elif field:
                field_and_institution = field
            
            print(f"      📐 Format: LEFT='{degree_type or '(no degree)'}' | RIGHT='{year_clean}'")
            print(f"      📐 Second line: '{field_and_institution or '(none)'}'")
            
            # CRITICAL FIX: Truncate degree_type if too long to prevent date wrapping
            # Max ~70 chars to ensure tab stop works properly (leaves room for date on right)
            display_degree = degree_type or institution or 'Education'
            if len(display_degree) > 70:
                display_degree = display_degree[:67] + '...'
                print(f"      ✂️  Truncated degree to: '{display_degree}'")
            
            # Degree type on the left (bold), year on the right
            deg_run = header_para.add_run(display_degree)
            deg_run.bold = True
            deg_run.font.size = Pt(10)
            if year_clean:
                header_para.add_run('\t')
                yr_run = header_para.add_run(year_clean)
                yr_run.bold = False
                yr_run.font.size = Pt(10)
            header_para.paragraph_format.space_after = Pt(0)
            # Prevent line from breaking
            header_para.paragraph_format.keep_together = True

            # Field + Institution on the next line (normal)
            last_para = header_para
            if field_and_institution:
                fi_para = self._insert_paragraph_after(header_para, field_and_institution)
                if fi_para is not None:
                    for run in fi_para.runs:
                        run.font.size = Pt(10)
                    fi_para.paragraph_format.space_after = Pt(2)
                    try:
                        fi_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    except Exception:
                        pass
                    last_para = fi_para
                    print(f"      ✅ Inserted field/institution line")
                else:
                    print(f"      ⚠️  Failed to insert field/institution paragraph")

            # Add details as bullet paragraphs - include ALL details from resume
            if details:
                detail_count = 0
                for detail in details:  # Don't limit or optimize - preserve ALL content
                    txt = (detail or '').strip()
                    if not txt or txt.lower() == (institution or '').lower():
                        continue
                    p = self._insert_paragraph_after(last_para, '')
                    if p is not None:
                        p.paragraph_format.left_indent = Inches(0.25)
                        try:
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        except Exception:
                            pass
                        run = p.add_run('• ' + txt.lstrip('•–—-*● '))
                        run.font.size = Pt(10)
                        p.paragraph_format.space_after = Pt(2)
                        last_para = p
                        detail_count += 1
                if detail_count > 0:
                    print(f"      ✅ Inserted {detail_count} detail bullets")
            
            # Add a blank line after each education entry for spacing
            
            # Add a blank line after each education entry for spacing
            blank = self._insert_paragraph_after(last_para, '')
            if blank is not None:
                blank.paragraph_format.space_after = Pt(6)
                try:
                    blank.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                except Exception:
                    pass
                last_para = blank
            
            print(f"      ✅ Education block inserted successfully")
            return last_para
            
        except Exception as e:
            print(f"  ⚠️  Error inserting education block: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _clean_duration(self, duration):
        """Normalize duration.
        Prefer 'Mon YYYY-Mon YYYY' (e.g., 'Nov 2011-Sept 2025').
        Fallback to 'YYYY-YYYY' or single 'YYYY'.
        """
        if not duration:
            return ''

        t = (duration or '').strip()
        if not t:
            return ''

        # Normalize connectors
        t = re.sub(r'[–—]', '-', t)
        t = re.sub(r'\s*(to|–|—|-)\s*', '-', t, flags=re.IGNORECASE)

        present = bool(re.search(r'\b(current|present)\b', t, re.IGNORECASE))

        # Month map with 'Sept' spelling
        month_map = {
            'january': 'Jan', 'jan': 'Jan',
            'february': 'Feb', 'feb': 'Feb',
            'march': 'Mar', 'mar': 'Mar',
            'april': 'Apr', 'apr': 'Apr',
            'may': 'May',
            'june': 'Jun', 'jun': 'Jun',
            'july': 'Jul', 'jul': 'Jul',
            'august': 'Aug', 'aug': 'Aug',
            'september': 'Sept', 'sept': 'Sept', 'sep': 'Sept',
            'october': 'Oct', 'oct': 'Oct',
            'november': 'Nov', 'nov': 'Nov',
            'december': 'Dec', 'dec': 'Dec',
        }

        def abbr(m):
            return month_map.get(m.lower(), m[:3].title())

        # Find month-year tokens
        my = [m.groups() for m in re.finditer(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s*(?:,\s*)?((?:19|20)\d{2})\b', t, flags=re.IGNORECASE)]
        if my:
            start_m, start_y = my[0]
            start_m = abbr(start_m)
            if len(my) >= 2:
                end_m, end_y = my[-1]
                end_m = abbr(end_m)
                return f"{start_m} {start_y}-{end_m} {end_y}"
            # Single month-year; attach Present or end year if available
            if present:
                return f"{start_m} {start_y}-Present"
            end_years = re.findall(r'\b((?:19|20)\d{2})\b', t)
            if len(end_years) >= 2:
                return f"{start_m} {end_years[0]}-{end_years[-1]}"
            return f"{start_m} {start_y}"

        # Fallback to years-only
        years = re.findall(r'\b(?:19|20)\d{2}\b', t)
        if len(years) >= 2:
            return f"{years[0]}-{years[-1]}"
        if len(years) == 1:
            return years[0]
        return ''
    
    def _parse_company_role(self, title):
        """Parse company and role from title line"""
        # Common patterns: "Company Name - Role" or "Role at Company" or "Role, Company"
        if ' - ' in title:
            parts = title.split(' - ', 1)
            return parts[0].strip(), parts[1].strip()
        elif ' at ' in title.lower():
            parts = re.split(r'\s+at\s+', title, flags=re.IGNORECASE)
            return parts[1].strip() if len(parts) > 1 else '', parts[0].strip()
        elif ', ' in title:
            parts = title.split(', ', 1)
            return parts[1].strip(), parts[0].strip()
        else:
            # Assume entire line is company or role
            return title.strip(), ''
    
    def _extract_institution(self, degree, details):
        """Extract institution name from degree line or details"""
        # Check if degree line contains institution (common pattern: "Degree, Institution")
        if ', ' in degree:
            parts = degree.split(', ', 1)
            return parts[1].strip()
        
        # Look in details for institution keywords
        institution_keywords = ['university', 'college', 'institute', 'school', 'academy']
        for detail in details:
            if any(kw in detail.lower() for kw in institution_keywords):
                return detail.strip()
        
        return ''
    
    def _remove_cell_borders(self, cell):
        """Remove all borders from a table cell"""
        try:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
        except:
            pass
    
    def _insert_table_after(self, doc, anchor, rows=1, cols=2):
        """Create a table and position it immediately after the given anchor (Paragraph or Table)."""
        try:
            # Simply add table to document - python-docx will handle placement
            table = doc.add_table(rows=rows, cols=cols)
            
            # Try to move it after anchor, but don't fail if it doesn't work
            try:
                tbl = table._element
                if hasattr(anchor, '_element'):
                    anchor_elm = anchor._element
                elif hasattr(anchor, '_tbl'):
                    anchor_elm = anchor._tbl
                else:
                    anchor_elm = anchor
                
                # Only move if we have a valid anchor
                if anchor_elm is not None and hasattr(anchor_elm, 'addnext'):
                    anchor_elm.addnext(tbl)
            except:
                # If moving fails, table will just be at end of document
                pass

            return table

        except Exception as e:
            print(f"  ⚠️  Error creating table: {e}")
            return None

    def _cleanup_duplicate_bullets_after_section(self, doc, section_heading_para, next_section_name):
        """
        AGGRESSIVE cleanup: Scan entire document after inserting formatted content
        and delete ANY remaining bullet points between this section and next section.
        This ensures NO duplication of raw content.
        """
        try:
            print(f"    🧹 AGGRESSIVE cleanup: Removing ALL raw content until '{next_section_name}'...")
            
            # Find the section heading paragraph index
            heading_idx = None
            for idx, para in enumerate(doc.paragraphs):
                if para._element == section_heading_para._element:
                    heading_idx = idx
                    break
            
            if heading_idx is None:
                return
            
            # Now scan from heading to next section and delete EVERYTHING except tables and section headings
            deleted = 0
            paras_to_delete = []
            
            # List of all section keywords to preserve
            section_keywords = ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 
                              'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY', 
                              'PROFESSIONAL EXPERIENCE', 'CAREER HISTORY', 'QUALIFICATIONS',
                              'ACHIEVEMENTS', 'AWARDS', 'LANGUAGES']
            
            for idx in range(heading_idx + 1, len(doc.paragraphs)):
                para = doc.paragraphs[idx]
                text = para.text.strip().upper()
                
                # Stop at next section
                if next_section_name in text and len(text) < 50:
                    print(f"       Stopped at next section: {text[:40]}")
                    break
                
                # Skip if paragraph is empty
                if not text:
                    continue
                
                # PRESERVE section headings (don't delete them!)
                is_section_heading = any(keyword in text for keyword in section_keywords) and len(text) < 50
                if is_section_heading:
                    print(f"       Preserved section heading: '{text[:40]}'")
                    continue
                
                # DELETE this paragraph (it's duplicate raw content)
                paras_to_delete.append(para)
                deleted += 1
                if deleted <= 5:
                    print(f"       Removing duplicate: '{text[:60]}'")
            
            # Actually delete the paragraphs
            for para in paras_to_delete:
                p_element = para._element
                p_element.getparent().remove(p_element)
            
            print(f"    🧹 Cleanup complete: Removed {deleted} duplicate paragraphs")
            
        except Exception as e:
            print(f"    ⚠️  Cleanup error: {e}")
    
    def _delete_following_bullets(self, paragraph, max_scan=200):
        """Delete ALL content after a heading until next section - includes TABLES and paragraphs."""
        try:
            body = paragraph._element.getparent()
            node = paragraph._element.getnext()
            deleted_paras = 0
            deleted_tables = 0
            scanned = 0
            
            print(f"    🔍 Starting deletion scan (max {max_scan} items)...")
            
            while node is not None and scanned < max_scan:
                scanned += 1
                next_node = node.getnext()
                
                # DELETE TABLES (raw content might be in tables)
                if node.tag.endswith('tbl'):
                    print(f"       Deleting table #{deleted_tables + 1}")
                    body.remove(node)
                    deleted_tables += 1
                    node = next_node
                    continue
                
                # DELETE PARAGRAPHS
                if node.tag.endswith('p'):
                    # Extract plain text
                    text_nodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    text = ''.join([t.text for t in text_nodes if t is not None and t.text is not None])
                    txt = (text or '').strip()
                    norm = txt.upper()
                    
                    # Stop ONLY at next section heading (not at tables or bullets)
                    section_keywords = ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 
                                      'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY', 
                                      'PROFESSIONAL EXPERIENCE', 'CAREER HISTORY', 'QUALIFICATIONS']
                    
                    if any(k in norm for k in section_keywords) and len(txt) < 50:
                        # This looks like a next section heading, stop deleting
                        print(f"       Stopped at next section: '{txt[:40]}'")
                        break
                    
                    # DELETE THIS PARAGRAPH (raw content)
                    if deleted_paras < 5:  # Log first 5 deletions
                        print(f"       Deleting para: '{txt[:60]}'...")
                    body.remove(node)
                    deleted_paras += 1
                
                node = next_node
            
            print(f"    🗑️  DELETED: {deleted_paras} paragraphs + {deleted_tables} tables (scanned {scanned} items)")
            
            if deleted_paras == 0 and deleted_tables == 0:
                print(f"    ⚠️  WARNING: Nothing was deleted! Content might still be there.")
                
        except Exception as e:
            print(f"    ⚠️  Error deleting content: {e}")
            import traceback
            traceback.print_exc()
    
    def _collect_bullets_after_heading(self, paragraph, max_scan=50):
        """Collect consecutive bullet-like paragraphs immediately after a heading/placeholder."""
        bullets = []
        try:
            node = paragraph._element.getnext()
            scanned = 0
            while node is not None and scanned < max_scan:
                scanned += 1
                if node.tag.endswith('tbl'):
                    # Collect all paragraph texts from the table as bullets
                    paras = node.xpath('.//w:p', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    for p in paras:
                        tnodes = p.xpath('.//w:t', namespaces=p.nsmap) if hasattr(p, 'xpath') else []
                        text = ''.join([t.text for t in tnodes if t is not None and t.text is not None]).strip()
                        if text:
                            bullets.append(text.lstrip(' •–—-*●'))
                    break
                if node.tag.endswith('p'):
                    text_nodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    text = ''.join([t.text for t in text_nodes if t is not None and t.text is not None])
                    txt = (text or '').strip()
                    norm = txt.upper()
                    if any(k in norm for k in ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY']):
                        break
                    if txt.startswith(('•', '-', '–', '—', '*', '●')) or re.match(r'^\d+[\).\-\s]', txt):
                        bullets.append(txt.lstrip(' •–—-*●'))
                    else:
                        break
                node = node.getnext()
        except Exception:
            pass
        return bullets

    def _delete_next_table(self, paragraph):
        """Delete the immediate next table after a heading/placeholder (used when raw content is a table)."""
        try:
            node = paragraph._element.getnext()
            deleted = 0
            # Delete multiple tables if they exist
            while node is not None and node.tag.endswith('tbl'):
                next_node = node.getnext()
                parent = node.getparent()
                parent.remove(node)
                deleted += 1
                node = next_node
            if deleted > 0:
                print(f"    🗑️  Deleted {deleted} old table(s)")
        except Exception as e:
            print(f"    ⚠️  Error deleting tables: {e}")
    
    def _paragraph_in_table(self, paragraph):
        try:
            node = paragraph._element
            while node is not None:
                if str(getattr(node, 'tag', '')).endswith('tbl'):
                    return True
                node = node.getparent()
        except Exception:
            pass
        return False
    
    def _remove_instructional_until_table(self, paragraph, max_scan=40):
        try:
            node = paragraph._element.getnext()
            scanned = 0
            while node is not None and scanned < max_scan:
                scanned += 1
                if node.tag.endswith('tbl'):
                    break
                if node.tag.endswith('p'):
                    tnodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    txt = ''.join([t.text for t in tnodes if t is not None and t.text is not None]).strip()
                    upper = txt.upper()
                    if any(h in upper for h in ['EMPLOYMENT', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EDUCATION', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(txt) < 50:
                        break
                    is_instr = bool(re.search(r'\bplease\b', txt, re.IGNORECASE)) or \
                              bool(re.search(r'(use this table|add or delete rows|respond with the years|list the candidate|required/desired)', txt, re.IGNORECASE))
                    if is_instr or not txt:
                        parent = node.getparent()
                        parent.remove(node)
                        node = paragraph._element.getnext()
                        continue
                    else:
                        break
                node = node.getnext()
        except Exception:
            pass

    def _clear_instruction_phrases(self, doc):
        """Remove all instructional text from the template"""
        try:
            phrases = [
                'PLEASE USE THIS TABLE TO LIST THE SKILLS',
                "PLEASE LIST THE CANDIDATE'S RELEVANT EMPLOYMENT HISTORY",
                'ADD OR DELETE ROWS AS NECESSARY',
                'INSERT NAME AND CONTACT INFORMATION FOR THE CAI CONTRACT MANAGER',
                'FOR EASE OF REFERENCE, THE CONTRACT MANAGERS',
                "CONTRACT MANAGERS' CONTACT INFORMATION APPEARS BELOW",
                'LISTED ON THE VECTORVMS REQUIREMENT',
                'SHANNON SWENSON',
                'SHANNON.SWENSON@CAI.IO',
                "LIST CANDIDATE'S EDUCATION BACKGROUND",
                'LIST EDUCATION BACKGROUND',
                'EDUCATION BACKGROUND'
            ]
            removed_count = 0
            paragraphs_to_clear = []
            
            for p in doc.paragraphs:
                t = (p.text or '').strip().upper()
                # Also check for angle bracket patterns
                if '<' in t and '>' in t:
                    # Remove angle brackets for comparison
                    clean_t = t.replace('<', '').replace('>', '')
                    if any(ph.replace('<', '').replace('>', '') in clean_t for ph in phrases):
                        paragraphs_to_clear.append(p)
                        continue
                    # CRITICAL: Only remove education placeholders if we haven't inserted education yet
                    # This prevents removing the actual education content we just added
                    if not self._education_inserted:
                        if 'CANDIDATE' in clean_t and 'EDUCATION' in clean_t:
                            paragraphs_to_clear.append(p)
                            continue
                
                if any(ph in t for ph in phrases):
                    paragraphs_to_clear.append(p)
            
            # Clear the paragraphs
            for p in paragraphs_to_clear:
                for r in p.runs:
                    r.text = ''
                removed_count += 1
            
            if removed_count > 0:
                print(f"  ✓ Removed {removed_count} instructional text paragraphs")
        except Exception as e:
            print(f"  ⚠️  Error clearing instructions: {e}")

    def _build_experience_from_bullets(self, bullets):
        """Best-effort convert raw bullet lines into structured exp list when parser is empty."""
        exps = []
        i = 0
        while i < len(bullets):
            line = bullets[i]
            # Case: role + dates on this line
            if re.search(r'(?:19|20)\d{2}', line):
                duration = self._clean_duration(line)
                role = re.sub(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-zA-Z]*\s+(?:19|20)\d{2}\b', '', line, flags=re.IGNORECASE)
                role = re.sub(r'\b(?:19|20)\d{2}\b', '', role)
                role = re.sub(r'\b(to|–|—|-)\b', '', role, flags=re.IGNORECASE).strip(' ,;:-')
                company = ''
                j = i + 1
                if i + 1 < len(bullets):
                    maybe_company = bullets[i+1].strip()
                    def _prob_company(text):
                        if not text:
                            return False
                        if len(text.split()) > 6:
                            return False
                        if text.endswith('.'):
                            return False
                        if re.search(r'\b(designed|managed|implemented|maintained|developed|led|created|configured|administered|tested|deployed|collaborated|supported|provided|responsible|oversaw|ownership|monitor|monitored|engineered)\b', text, re.IGNORECASE):
                            return False
                        if text[0].islower():
                            return False
                        return True
                    if _prob_company(maybe_company):
                        company = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', maybe_company, flags=re.IGNORECASE).strip()
                        j = i + 2
                details = []
                while j < len(bullets):
                    if re.search(r'(?:19|20)\d{2}', bullets[j]):
                        break
                    details.append(bullets[j])
                    j += 1
                exps.append({'company': company, 'role': role, 'duration': duration, 'details': details})
                i = j
            else:
                i += 1
        return exps
    
    def _build_education_from_bullets(self, bullets):
        """Convert raw education bullets into degree/institution/year list when parser is empty."""
        edus = []
        i = 0
        while i < len(bullets):
            line = bullets[i]
            degree = ''
            institution = ''
            year = ''
            # Try to split by institution keyword
            m = re.search(r'(university|college|school|institute|academy)\b.*', line, flags=re.IGNORECASE)
            if m:
                degree = line[:m.start()].strip(' ,;:-')
                institution = line[m.start():].strip()
                year = self._clean_duration(line)
            else:
                # If next line is year, treat current as degree+institution
                if i + 1 < len(bullets) and re.search(r'(?:19|20)\d{2}', bullets[i+1]):
                    degree = line
                    year = self._clean_duration(bullets[i+1])
                    i += 1
                else:
                    degree = line
                    year = self._clean_duration(line)
            # Cleanup
            degree = degree.strip()
            institution = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', institution, flags=re.IGNORECASE).strip()
            edus.append({'degree': degree, 'institution': institution, 'year': year, 'details': []})
            i += 1
        return edus

    def _optimize_details(self, details, max_bullets=12, max_words=22, max_chars=160):
        """Shorten and normalize bullet points while preserving meaning.
        - trims bullets, normalizes acronyms, removes duplicate entries
        - caps each bullet by words/chars
        """
        cleaned = []
        seen = set()
        for d in details:
            if not d or not isinstance(d, str):
                continue
            t = d.strip()
            if not t:
                continue
            # Remove leading bullet chars
            t = t.lstrip('•–—-*● \t-').strip()
            # Normalize common wording
            repl = [
                (r'\bas well as\b', 'and'),
                (r'\bin order to\b', 'to'),
                (r'\bkey performance indicators\s*\(([^\)]+)\)', r'\1'),
                (r'\bkey performance indicators\b', 'KPIs'),
                (r'\bquickbooks\b', 'QuickBooks'),
                (r'\bums?\s*worldship\b', 'UPS WorldShip'),
            ]
            for pattern, repl_to in repl:
                t = re.sub(pattern, repl_to, t, flags=re.IGNORECASE)

            t = self._shorten_text(t, max_words=max_words, max_chars=max_chars)
            t = self._normalize_acronyms(t)

            key = t.lower()
            if key and key not in seen:
                seen.add(key)
                cleaned.append(t.rstrip())

            if len(cleaned) >= max_bullets:
                break
        return cleaned

    def _normalize_acronyms(self, text):
        """Normalize common acronyms casing."""
        mapping = {
            'kpi': 'KPI', 'kpis': 'KPIs',
            'lms': 'LMS',
            'qa': 'QA',
        }
        def repl(m):
            return mapping.get(m.group(0).lower(), m.group(0))
        return re.sub(r'\b(kpis?|lms|qa)\b', repl, text, flags=re.IGNORECASE)

    def _shorten_text(self, text, max_words=22, max_chars=160):
        """Heuristic shortening: prefer cutting at clause boundaries, then word limit."""
        t = re.sub(r'\s+', ' ', text).strip()
        # Prefer to cut at clause markers if too long
        if len(t) > max_chars:
            for marker in ['; ', '. ', ' which ', ' that ', ' ensuring ', ' including ', ' while ', ' whereas ', ' whereby ']:
                idx = t.lower().find(marker)
                if 0 < idx <= max_chars:
                    t = t[:idx].rstrip('.; ,')
                    break
        # Enforce word cap
        words = t.split()
        if len(words) > max_words:
            t = ' '.join(words[:max_words]).rstrip(',;')
        # Ensure terminal period for readability
        if t and t[-1] not in '.!?':
            t = t + '.'
        return t
    
    def _create_replacement_map(self):
        """Create comprehensive replacement map"""
        replacements = {}
        
        # Personal information - Multiple formats
        # NOTE: Be specific to avoid replacing CAI contact manager info
        if self.resume_data.get('name'):
            display_name = f"<{self.resume_data['name']}>"
            replacements['[NAME]'] = display_name
            replacements['[CANDIDATE NAME]'] = display_name
            replacements['<CANDIDATE NAME>'] = display_name
            replacements["<Candidate's full name>"] = display_name
            replacements['<Candidate Name>'] = display_name
            replacements['<Name>'] = display_name
            replacements['Your Name'] = display_name
            replacements['CANDIDATE NAME'] = display_name
            # DO NOT replace "Insert name" as it might be in CAI contact section
        
        if self.resume_data.get('email'):
            # ONLY replace explicit placeholders, NOT actual email addresses
            replacements['[EMAIL]'] = self.resume_data['email']
            replacements['[Email]'] = self.resume_data['email']
            replacements['<EMAIL>'] = self.resume_data['email']
            replacements['<Email>'] = self.resume_data['email']
            replacements['<Candidate Email>'] = self.resume_data['email']
            # DO NOT replace example emails or "Email:" labels to avoid changing CAI contact info
        
        if self.resume_data.get('phone'):
            # ONLY replace explicit placeholders, NOT actual phone numbers
            replacements['[PHONE]'] = self.resume_data['phone']
            replacements['[Phone]'] = self.resume_data['phone']
            replacements['<PHONE>'] = self.resume_data['phone']
            replacements['<Phone>'] = self.resume_data['phone']
            replacements['<Candidate Phone>'] = self.resume_data['phone']
            # DO NOT replace example numbers or "Phone:" labels to avoid changing CAI contact info
        
        if self.resume_data.get('address'):
            replacements['[ADDRESS]'] = self.resume_data['address']
            replacements['[Address]'] = self.resume_data['address']
            replacements['<ADDRESS>'] = self.resume_data['address']
            replacements['<Address>'] = self.resume_data['address']
            replacements['Your Address'] = self.resume_data['address']
        
        if self.resume_data.get('linkedin'):
            replacements['[LINKEDIN]'] = self.resume_data['linkedin']
            replacements['[LinkedIn]'] = self.resume_data['linkedin']
            replacements['<LINKEDIN>'] = self.resume_data['linkedin']
            replacements['<LinkedIn>'] = self.resume_data['linkedin']
            replacements['linkedin.com/in/username'] = self.resume_data['linkedin']
        
        if self.resume_data.get('dob'):
            replacements['[DOB]'] = self.resume_data['dob']
            replacements['[Date of Birth]'] = self.resume_data['dob']
            replacements['<DOB>'] = self.resume_data['dob']
        
        return replacements
    
    def _text_contains(self, text, search_term):
        """Case-insensitive text search"""
        return search_term.lower() in text.lower()
    
    def _replace_in_paragraph(self, paragraph, search_term, replacement):
        """Replace text in paragraph while preserving formatting AND alignment, remove highlighting"""
        replaced = 0
        
        # PRESERVE ALIGNMENT: Store original alignment before modification
        original_alignment = paragraph.alignment
        
        # First try: Replace in individual runs
        for run in paragraph.runs:
            if self._text_contains(run.text, search_term):
                # Case-insensitive replacement
                pattern = re.compile(re.escape(search_term), re.IGNORECASE)
                run.text = pattern.sub(replacement, run.text)
                # Remove highlighting
                try:
                    run.font.highlight_color = None
                except:
                    pass
                replaced += 1
        
        # Second try: If not found in individual runs, text might be split
        # Combine all runs and check
        if replaced == 0 and self._text_contains(paragraph.text, search_term):
            # Text is split across runs - need to handle differently
            full_text = paragraph.text
            pattern = re.compile(re.escape(search_term), re.IGNORECASE)
            new_text = pattern.sub(replacement, full_text)
            
            if new_text != full_text:
                # Clear all runs and add new text
                for run in paragraph.runs:
                    run.text = ''
                
                # Add replacement text to first run
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    # Remove highlighting
                    try:
                        paragraph.runs[0].font.highlight_color = None
                    except:
                        pass
                    replaced += 1
                else:
                    # No runs, add new run
                    new_run = paragraph.add_run(new_text)
                    # Remove highlighting
                    try:
                        new_run.font.highlight_color = None
                    except:
                        pass
                    replaced += 1
        
        # RESTORE ALIGNMENT: Apply original alignment after replacement
        if original_alignment is not None:
            paragraph.alignment = original_alignment
        
        return replaced

    def _regex_replace_paragraph(self, paragraph, pattern, replacement):
        """Regex-based replacement across runs: rebuilds paragraph text, removes highlighting, preserves alignment."""
        try:
            # PRESERVE ALIGNMENT: Store original alignment before modification
            original_alignment = paragraph.alignment
            
            full_text = paragraph.text or ''
            new_text = re.sub(pattern, replacement, full_text, flags=re.IGNORECASE)
            if new_text != full_text:
                # clear runs and set new_text
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    # CRITICAL: Remove yellow highlighting from name
                    try:
                        from docx.enum.text import WD_COLOR_INDEX
                        paragraph.runs[0].font.highlight_color = None
                    except:
                        pass
                else:
                    new_run = paragraph.add_run(new_text)
                    # Remove highlighting from new run
                    try:
                        from docx.enum.text import WD_COLOR_INDEX
                        new_run.font.highlight_color = None
                    except:
                        pass
                
                # RESTORE ALIGNMENT: Apply original alignment after replacement
                if original_alignment is not None:
                    paragraph.alignment = original_alignment
        except Exception:
            pass
    
    def _add_sections_content(self, doc):
        """Add resume sections to document and replace placeholders - SIMPLIFIED to prevent duplication"""
        sections_added = 0
        
        # Flags are initialized in _format_docx_file()
        # This method will check them to prevent duplicate insertion
        
        print(f"\n🔍 Scanning document for sections (SUMMARY, EXPERIENCE, EDUCATION, SKILLS)...")
        print(f"  📊 Section status: Summary={self._summary_inserted}, Experience={self._experience_inserted}, Education={self._education_inserted}")
        
        # SINGLE PASS: Look for headings only (ignore placeholders to avoid duplication)
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.upper().strip()
            
            # SUMMARY SECTION
            if not self._summary_inserted and any(marker in para_text for marker in ['SUMMARY', 'OBJECTIVE', 'PROFILE', 'PROFESSIONAL SUMMARY']):
                summary = (self.resume_data.get('summary') or '').strip()
                summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {}))
                
                # If summary_lines is empty but summary text exists, split it into bullet points
                if not summary_lines and summary:
                    # Split by common delimiters (newlines, periods followed by capital letters, bullet points)
                    # First try splitting by newlines
                    lines = [line.strip() for line in summary.split('\n') if line.strip()]
                    # If only one line, try splitting by sentences
                    if len(lines) == 1:
                        # Split by period followed by space and capital letter
                        sentences = re.split(r'\.\s+(?=[A-Z])', summary)
                        lines = [s.strip() + ('.' if not s.strip().endswith('.') else '') for s in sentences if s.strip()]
                    summary_lines = lines
                
                if summary or summary_lines:
                    print(f"  ✓ Found SUMMARY at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    
                    # Clear the heading paragraph (keep only the heading text)
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = 'SUMMARY'
                        paragraph.runs[0].bold = True
                        paragraph.runs[0].font.size = Pt(12)
                    
                    # Delete any following content before inserting new
                    self._delete_following_bullets(paragraph, max_scan=20)
                    
                    # Always use bullet-style summary (no indentation for summary)
                    if summary_lines:
                        last_para = paragraph
                        for line in summary_lines:
                            txt = (line or '').strip()
                            if not txt:
                                continue
                            bullet_para = self._insert_paragraph_after(last_para, '')
                            if bullet_para:
                                # No left indent for summary bullets
                                run = bullet_para.add_run('• ' + txt.lstrip('•–—-*● '))
                                run.font.size = Pt(10)
                                bullet_para.paragraph_format.space_after = Pt(2)
                                last_para = bullet_para
                    
                    self._summary_inserted = True
                    sections_added += 1
                    print(f"    → Inserted summary ({len(summary_lines)} bullet points)")
                    continue
                else:
                    # No summary data: remove heading and any following content until next section
                    print(f"  ⚠️  SUMMARY heading found but no data; removing section")
                    self._delete_following_bullets(paragraph, max_scan=80)
                    self._delete_next_table(paragraph)
                    self._delete_paragraph(paragraph)
                    continue
            
            # EXPERIENCE SECTION - Check if it hasn't been inserted yet
            if (not self._experience_inserted \
                and any(marker in para_text for marker in ['EMPLOYMENT HISTORY', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'WORK HISTORY', 'CAREER HISTORY']) \
                and len(paragraph.text.strip()) < 50 \
                and not self._paragraph_in_table(paragraph)):
                experiences = self.resume_data.get('experience', [])
                # Fallback: build structured experiences from the raw bullets beneath the heading
                if not experiences:
                    raw_bullets = self._collect_bullets_after_heading(paragraph, max_scan=120)
                    if raw_bullets:
                        experiences = self._build_experience_from_bullets(raw_bullets)
                        print(f"  🔄 Built {len(experiences)} experience entries from raw bullets")
                if experiences:
                    print(f"  ✓ Found EXPERIENCE at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    
                    # STEP 1: Clear the heading paragraph (keep only the heading text)
                    original_heading = paragraph.text.strip()
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = 'EMPLOYMENT HISTORY'
                        paragraph.runs[0].bold = True
                        paragraph.runs[0].font.size = Pt(12)
                    
                    # STEP 2: Delete old template content only if experiences were built from fallback
                    # (Don't delete if we have structured experience from resume parser)
                    if not self.resume_data.get('experience'):
                        # Only clear template placeholders/raw bullets when using fallback
                        self._delete_following_bullets(paragraph, max_scan=800)
                        self._delete_next_table(paragraph)
                    
                    # STEP 3: Insert clean structured blocks - ALL entries
                    last_element = paragraph
                    for exp in experiences:
                        table = self._insert_experience_block(doc, last_element, exp)
                        if table:
                            last_element = table
                    
                    # STEP 4: Skip aggressive cleanup to preserve newly inserted bullets
                    # self._cleanup_duplicate_bullets_after_section(doc, paragraph, 'EDUCATION')
                    
                    self._experience_inserted = True
                    sections_added += 1
                    print(f"    → Inserted {len(experiences)} experience entries")
                    continue
            
            # SKILLS SECTION
            if not self._skills_inserted and any(marker in para_text for marker in ['SKILLS', 'TECHNICAL SKILLS', 'COMPETENCIES', 'EXPERTISE']):
                skills = self.resume_data.get('skills', [])
                if skills and len(skills) > 0:
                    print(f"  ✓ Found SKILLS at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    
                    # Ensure heading formatting (BOLD, UNDERLINED, CAPITAL)
                    paragraph.clear()
                    run = paragraph.add_run('SKILLS')
                    run.bold = True
                    run.underline = True  # UNDERLINE
                    run.font.size = Pt(12)
                    run.font.all_caps = True  # CAPITAL
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.space_after = Pt(6)
                    print(f"    ✅ Formatted SKILLS heading: BOLD, UNDERLINED, CAPITAL")

                    # Determine if a table follows this heading
                    node = paragraph._element.getnext()
                    has_table_next = bool(node is not None and str(getattr(node, 'tag', '')).endswith('tbl'))

                    if not has_table_next:
                        # No table available → insert bullets
                        print("    → No skills table after heading; inserting bullet list")
                        # Remove any raw content under heading first
                        self._delete_following_bullets(paragraph, max_scan=50)
                        self._insert_skills_bullets(doc, paragraph, skills)
                    else:
                        print("    → Skills table detected; will fill during table pass")
                        # Remove instructional text between heading and table
                        self._remove_instructional_until_table(paragraph, max_scan=20)
                    
                    self._skills_inserted = True
                    sections_added += 1
                    continue
                else:
                    # No skills: remove section heading and trailing content
                    print(f"  ⚠️  SKILLS heading found but resume has no skills; removing section")
                    self._delete_following_bullets(paragraph, max_scan=50)
                    self._delete_next_table(paragraph)
                    self._delete_paragraph(paragraph)
                    continue
            
            # EDUCATION SECTION - Check multiple variations and handle fallback
            # CRITICAL: Skip placeholder text like "<List candidate's education background>"
            is_placeholder = '<' in para_text and '>' in para_text and 'CANDIDATE' in para_text
            if not self._education_inserted and not is_placeholder and any(marker in para_text for marker in ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND', 'ACADEMICS', 'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS', 'EDUCATION/CERTIFICATES', 'EDUCATION / CERTIFICATES']):
                education = self.resume_data.get('education', [])
                print(f"  🔍 Found potential EDUCATION heading at paragraph {para_idx}: '{paragraph.text[:50]}'")
                print(f"     Initial education count from resume_data: {len(education) if education else 0}")
                
                # Fallback priority 1: derive from resume sections (not from template content)
                if not education:
                    section_lines = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                    if section_lines:
                        education = self._build_education_from_bullets(section_lines)
                        print(f"  🔄 Built {len(education)} education entries from resume sections")
                
                # Also check sections dict directly
                if not education:
                    sections = self.resume_data.get('sections', {})
                    if 'education' in sections:
                        edu_text = sections['education']
                        if isinstance(edu_text, str):
                            lines = [l.strip() for l in edu_text.split('\n') if l.strip()]
                            education = self._build_education_from_bullets(lines)
                            print(f"  🔄 Built {len(education)} education entries from sections.education")
                        elif isinstance(edu_text, list):
                            education = self._build_education_from_bullets(edu_text)
                            print(f"  🔄 Built {len(education)} education entries from sections.education list")
                
                # Fallback priority 2: derive from any raw bullets directly under the heading in template (rare)
                if not education:
                    raw_bullets = self._collect_bullets_after_heading(paragraph, max_scan=80)
                    if raw_bullets:
                        education = self._build_education_from_bullets(raw_bullets)
                        print(f"  🔄 Built {len(education)} education entries from raw bullets under heading")
                
                print(f"     📊 Total education entries to insert: {len(education) if education else 0}")
                
                if education:
                    print(f"  ✓ Found EDUCATION at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    print(f"  📚 Have {len(education)} education entries to insert")
                    
                    # DEBUG: Show education data
                    for i, edu in enumerate(education[:3]):
                        print(f"      {i+1}. Degree: '{edu.get('degree', '')[:40]}', Institution: '{edu.get('institution', '')[:30]}', Year: '{edu.get('year', '')}'")
                    
                    # STEP 1: Clear and reset heading with proper formatting (BOLD, UNDERLINED, CAPITAL)
                    paragraph.clear()
                    run = paragraph.add_run('EDUCATION')
                    run.bold = True
                    run.underline = True  # UNDERLINE
                    run.font.size = Pt(12)
                    run.font.all_caps = True  # CAPITAL
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.space_after = Pt(6)
                    print(f"    ✅ Created EDUCATION heading: BOLD, UNDERLINED, CAPITAL")
                    
                    # STEP 2: Delete ALL following content (tables + paragraphs)
                    self._delete_following_bullets(paragraph)
                    self._delete_next_table(paragraph)
                    
                    # STEP 3: Insert clean structured blocks (or bullets for simple entries) - ALL entries
                    simple_entries = [e for e in education if not (e.get('institution') or (e.get('details') or []))]
                    inserted_count = 0
                    if simple_entries and len(simple_entries) == len(education):
                        self._insert_education_bullets(doc, paragraph, education)
                        inserted_count = len(education)
                    else:
                        last_element = paragraph
                        for edu in education:
                            block = self._insert_education_block(doc, last_element, edu)
                            if block:
                                last_element = block
                                inserted_count += 1
                    
                    # STEP 4: Skip aggressive cleanup to preserve newly inserted bullets
                    # self._cleanup_duplicate_bullets_after_section(doc, paragraph, 'SKILLS')
                    
                    self._education_inserted = True
                    sections_added += 1
                    print(f"    ✅ Successfully inserted {inserted_count} education entries and marked as complete")
                    continue
                else:
                    # No education data but heading exists - REMOVE heading and DON'T mark as processed
                    # This allows education to be added later when data is available
                    print(f"  ⚠️  EDUCATION heading found but no data available - removing heading to add later")
                    self._delete_following_bullets(paragraph, max_scan=50)
                    self._delete_next_table(paragraph)
                    self._delete_paragraph(paragraph)
                    # DO NOT set self._education_inserted = True here!
                    continue
        
        # FINAL SAFETY NET: If education is STILL not inserted but we have education data, add it at the end
        if not self._education_inserted:
            education = self.resume_data.get('education', [])
            if not education:
                section_lines = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                if section_lines:
                    education = self._build_education_from_bullets(section_lines)
            
            if not education:
                sections = self.resume_data.get('sections', {})
                if 'education' in sections:
                    edu_text = sections['education']
                    if isinstance(edu_text, str):
                        lines = [l.strip() for l in edu_text.split('\n') if l.strip()]
                        education = self._build_education_from_bullets(lines)
                    elif isinstance(edu_text, list):
                        education = self._build_education_from_bullets(edu_text)
            
            if education:
                print(f"\n🚨 SAFETY NET: Education not inserted yet but data exists! Adding after Employment...")
                print(f"   Have {len(education)} education entries to insert")
                
                # Find Employment History section and insert after it
                employment_idx = None
                for idx, para in enumerate(doc.paragraphs):
                    text = para.text.strip().upper()
                    if any(kw in text for kw in ['EMPLOYMENT HISTORY', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'EMPLOYMENT']):
                        if len(text) < 50:
                            employment_idx = idx
                            print(f"   📍 Found Employment at paragraph {idx}")
                            break
                
                # Find insertion point after employment content
                if employment_idx is not None:
                    # Scan forward to find end of employment section
                    insertion_idx = employment_idx + 20  # Default
                    for j in range(employment_idx + 1, min(employment_idx + 100, len(doc.paragraphs))):
                        next_text = doc.paragraphs[j].text.strip().upper()
                        if any(kw in next_text for kw in ['SKILLS', 'SUMMARY', 'PROJECTS', 'CERTIFICATIONS']) and len(next_text) < 50:
                            insertion_idx = j
                            print(f"   📍 Will insert EDUCATION at paragraph {j}")
                            break
                    
                    anchor_para = doc.paragraphs[insertion_idx] if insertion_idx < len(doc.paragraphs) else doc.paragraphs[-1]
                else:
                    # No employment found, use end of document
                    anchor_para = doc.paragraphs[-1]
                    print(f"   📍 No Employment found, inserting at end")
                
                # Add EDUCATION heading with proper formatting (BOLD, UNDERLINED, CAPITAL)
                heading_para = self._insert_paragraph_after(anchor_para, '')
                if heading_para:
                    heading_para.clear()
                    run = heading_para.add_run('EDUCATION')
                    run.bold = True
                    run.underline = True  # UNDERLINE
                    run.font.size = Pt(12)
                    run.font.all_caps = True  # CAPITAL
                    heading_para.paragraph_format.space_before = Pt(12)
                    heading_para.paragraph_format.space_after = Pt(6)
                    print(f"   ✅ Created EDUCATION heading: BOLD, UNDERLINED, CAPITAL")
                    
                    # Insert education entries
                    last_element = heading_para
                    for edu in education:
                        block = self._insert_education_block(doc, last_element, edu)
                        if block:
                            last_element = block
                    
                    self._education_inserted = True
                    sections_added += 1
                    print(f"   ✅ SAFETY NET: Successfully added {len(education)} education entries")
            else:
                print(f"\n⚠️  Education not inserted and no education data found - marking as processed")
                self._education_inserted = True
        
        print(f"\n✅ Section insertion complete. Summary: {self._summary_inserted}, Experience: {self._experience_inserted}, Education: {self._education_inserted}")
        
        # NEW: Add missing sections from candidate resume (Skills, Certificates, etc.)
        sections_added += self._add_missing_sections(doc)
        
        return sections_added
    
    def _add_missing_sections(self, doc):
        """Add any missing sections from candidate resume that aren't in template.
        Adds them in order after Education: Skills → Certificates → Projects → Languages
        """
        added_count = 0
        
        print(f"\n🔍 Checking for missing sections to add...")
        
        # Find anchor point: after Education or after Employment
        anchor_para = None
        anchor_idx = None
        
        # Try to find Education section end
        for idx, p in enumerate(doc.paragraphs):
            if 'EDUCATION' in (p.text or '').upper() and len(p.text.strip()) < 50:
                # Found education heading, scan forward to find end of section
                for j in range(idx + 1, min(idx + 50, len(doc.paragraphs))):
                    next_p = doc.paragraphs[j]
                    next_text = (next_p.text or '').strip().upper()
                    # Stop at next major section
                    if any(h in next_text for h in ['SKILLS', 'CERTIFICATES', 'PROJECTS', 'LANGUAGES', 'REFERENCES']) and len(next_text) < 50:
                        anchor_para = doc.paragraphs[j - 1]
                        anchor_idx = j - 1
                        break
                if anchor_para:
                    break
        
        # Fallback: use Employment History end
        if not anchor_para and hasattr(self, '_employment_tail_para'):
            anchor_para = self._employment_tail_para
            print(f"  Using Employment History tail as anchor")
        
        # Fallback: use last paragraph
        if not anchor_para and doc.paragraphs:
            anchor_para = doc.paragraphs[-1]
            print(f"  Using last paragraph as anchor")
        
        if not anchor_para:
            print(f"  ⚠️  No anchor found, skipping missing sections")
            return 0
        
        print(f"  📍 Anchor point found, will add missing sections after it")
        
        # Check and add SKILLS if missing
        if not self._skills_inserted:
            skills = self.resume_data.get('skills', [])
            if skills:
                print(f"  ➕ Adding missing SKILLS section ({len(skills)} skills)")
                # Add blank line
                blank = self._insert_paragraph_after(anchor_para, '')
                if blank:
                    anchor_para = blank
                # Add heading
                heading = self._insert_paragraph_after(anchor_para, 'SKILLS')
                if heading:
                    for r in heading.runs:
                        r.bold = True
                        r.underline = True
                        r.font.size = Pt(11)
                    heading.paragraph_format.space_before = Pt(12)
                    heading.paragraph_format.space_after = Pt(6)
                    # Add skills as bullets
                    self._insert_skills_bullets(doc, heading, skills)
                    anchor_para = heading
                    self._skills_inserted = True
                    added_count += 1
        
        # Check and add CERTIFICATES/CERTIFICATIONS if present in resume
        certificates = []
        sections = self.resume_data.get('sections', {})
        for key in ['certificates', 'certifications', 'certification', 'certificate']:
            if key in sections:
                cert_data = sections[key]
                if isinstance(cert_data, str):
                    certificates = [l.strip() for l in cert_data.split('\n') if l.strip()]
                elif isinstance(cert_data, list):
                    certificates = [str(c).strip() for c in cert_data if str(c).strip()]
                if certificates:
                    break
        
        if certificates:
            print(f"  ➕ Adding CERTIFICATES section ({len(certificates)} certificates)")
            # Add blank line
            blank = self._insert_paragraph_after(anchor_para, '')
            if blank:
                anchor_para = blank
            # Add heading
            heading = self._insert_paragraph_after(anchor_para, 'CERTIFICATIONS')
            if heading:
                for r in heading.runs:
                    r.bold = True
                    r.underline = True
                    r.font.size = Pt(11)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
                # Add certificates as bullets
                last_para = heading
                for cert in certificates:
                    bullet_para = self._insert_paragraph_after(last_para, '')
                    if bullet_para:
                        run = bullet_para.add_run('• ' + cert.lstrip('•–—-*● '))
                        run.font.size = Pt(10)
                        bullet_para.paragraph_format.space_after = Pt(2)
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
                        last_para = bullet_para
                anchor_para = last_para
                added_count += 1
        
        # Check and add PROJECTS if present
        projects = []
        for key in ['projects', 'project']:
            if key in sections:
                proj_data = sections[key]
                if isinstance(proj_data, str):
                    projects = [l.strip() for l in proj_data.split('\n') if l.strip()]
                elif isinstance(proj_data, list):
                    projects = [str(p).strip() for p in proj_data if str(p).strip()]
                if projects:
                    break
        
        if projects:
            print(f"  ➕ Adding PROJECTS section ({len(projects)} projects)")
            blank = self._insert_paragraph_after(anchor_para, '')
            if blank:
                anchor_para = blank
            heading = self._insert_paragraph_after(anchor_para, 'PROJECTS')
            if heading:
                for r in heading.runs:
                    r.bold = True
                    r.underline = True
                    r.font.size = Pt(11)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
                last_para = heading
                for proj in projects:
                    bullet_para = self._insert_paragraph_after(last_para, '')
                    if bullet_para:
                        run = bullet_para.add_run('• ' + proj.lstrip('•–—-*● '))
                        run.font.size = Pt(10)
                        bullet_para.paragraph_format.space_after = Pt(2)
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
                        last_para = bullet_para
                anchor_para = last_para
                added_count += 1
        
        print(f"  ✅ Added {added_count} missing sections")
        return added_count
    
    def _detect_table_type(self, table):
        """
        DYNAMICALLY detect table type by analyzing column headers.
        Returns: 'skills', 'experience', 'education', or None
        """
        if len(table.rows) < 1:
            return None
        
        # Get all text from first few rows (headers might span multiple rows)
        header_text = ''
        for row_idx in range(min(3, len(table.rows))):
            for cell in table.rows[row_idx].cells:
                header_text += ' ' + cell.text.lower()
        
        # Skills table indicators
        skills_indicators = ['skill', 'technology', 'competency', 'expertise', 'proficiency', 
                           'years used', 'last used', 'technical']
        
        # Experience/Employment table indicators  
        experience_indicators = ['employment', 'company', 'employer', 'position', 'role', 
                                'job title', 'work history', 'experience', 'responsibilities']
        
        # Education table indicators
        education_indicators = ['education', 'degree', 'institution', 'university', 'college', 
                              'school', 'graduation', 'qualification']
        
        # Count matches for each type
        skills_score = sum(1 for ind in skills_indicators if ind in header_text)
        exp_score = sum(1 for ind in experience_indicators if ind in header_text)
        edu_score = sum(1 for ind in education_indicators if ind in header_text)
        
        # Return type with highest score
        if skills_score > 0 and skills_score >= exp_score and skills_score >= edu_score:
            return 'skills'
        elif exp_score > 0 and exp_score >= edu_score:
            return 'experience'
        elif edu_score > 0:
            return 'education'
        
        return None
    
    def _fill_dynamic_table(self, table, table_type):
        """
        DYNAMICALLY fill table based on column headers and data type.
        Works with ANY column structure!
        """
        if len(table.rows) < 1:
            return 0
        
        print(f"     🔍 Analyzing table structure...")
        
        # Get column headers (from first row)
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        print(f"     📋 Column headers: {headers}")
        
        # Map columns to data fields intelligently
        column_mapping = self._map_columns_to_fields(headers, table_type)
        print(f"     🗺️  Column mapping: {column_mapping}")
        
        # Get data to fill
        if table_type == 'skills':
            data_items = self._get_skills_data()
        elif table_type == 'experience':
            data_items = self._get_experience_data()
        elif table_type == 'education':
            data_items = self._get_education_data()
        else:
            return 0
        
        if not data_items:
            print(f"     ⚠️  No {table_type} data available")
            return 0
        
        print(f"     📊 Found {len(data_items)} {table_type} items to fill")
        
        # Fill rows dynamically - insert ALL items from resume
        filled = 0
        for item in data_items:
            new_row = table.add_row()
            
            for col_idx, field_name in column_mapping.items():
                value = item.get(field_name, '')
                if col_idx < len(new_row.cells):
                    new_row.cells[col_idx].text = str(value)
            
            filled += 1
            if filled <= 3:
                print(f"        ✓ Row {filled}: {list(item.values())[:3]}")
        
        return filled
    
    def _map_columns_to_fields(self, headers, table_type):
        """
        INTELLIGENT column mapping - maps table columns to resume data fields
        based on semantic understanding of column names.
        """
        mapping = {}  # {column_index: field_name}
        
        if table_type == 'skills':
            # Define possible column patterns for skills
            patterns = {
                'skill': ['skill', 'technology', 'competency', 'expertise', 'tool', 'name'],
                'years': ['year', 'experience', 'exp', 'yrs', 'duration', 'used'],
                'last_used': ['last', 'recent', 'current', 'latest', 'when']
            }
        elif table_type == 'experience':
            patterns = {
                'company': ['company', 'employer', 'organization', 'firm'],
                'role': ['role', 'position', 'title', 'job'],
                'duration': ['date', 'year', 'period', 'duration', 'from', 'to', 'when'],
                'location': ['location', 'city', 'state', 'place'],
                'responsibilities': ['responsibilit', 'duties', 'description', 'summary']
            }
        elif table_type == 'education':
            patterns = {
                'degree': ['degree', 'qualification', 'certificate', 'program'],
                'institution': ['institution', 'university', 'college', 'school'],
                'year': ['year', 'date', 'graduation', 'completion'],
                'field': ['field', 'major', 'specialization', 'subject'],
                'gpa': ['gpa', 'grade', 'marks', 'score']
            }
        else:
            return mapping
        
        # Map each column to best matching field
        for col_idx, header in enumerate(headers):
            header_lower = header.lower()
            best_match = None
            best_score = 0
            
            for field_name, keywords in patterns.items():
                score = sum(1 for kw in keywords if kw in header_lower)
                if score > best_score:
                    best_score = score
                    best_match = field_name
            
            if best_match and best_score > 0:
                mapping[col_idx] = best_match
                print(f"        Column {col_idx} ('{header}') → {best_match}")
        
        return mapping
    
    def _get_skills_data(self):
        """Get skills data in standardized format"""
        skills_list = []
        raw_skills = self.resume_data.get('skills', [])
        
        for skill in raw_skills[:15]:
            skill_name = skill if isinstance(skill, str) else skill.get('name', '')
            skills_list.append({
                'skill': skill_name,
                'years': '2+ years',  # Default
                'last_used': 'Recent'
            })
        
        return skills_list
    
    def _get_experience_data(self):
        """Get experience data in standardized format"""
        exp_list = []
        experiences = self.resume_data.get('experience', [])
        
        for exp in experiences:  # Return ALL experience entries
            exp_list.append({
                'company': exp.get('company', ''),
                'role': exp.get('role', ''),
                'duration': exp.get('duration', ''),
                'location': '',  # Extract if available
                'responsibilities': '\n'.join(exp.get('details', [])[:3])
            })
        
        return exp_list
    
    def _get_education_data(self):
        """Get education data in standardized format"""
        edu_list = []
        education = self.resume_data.get('education', [])
        
        for edu in education:  # Return ALL education entries
            edu_list.append({
                'degree': edu.get('degree', ''),
                'institution': edu.get('institution', ''),
                'year': edu.get('year', ''),
                'field': '',  # Extract from degree if available
                'gpa': ''
            })
        
        return edu_list
    
    def _is_skills_table(self, table):
        """Check if table is a skills table by examining headers - FLEXIBLE detection"""
        if len(table.rows) < 1:  # Changed from 2 to 1 - just need header row
            print(f"       ⚠️  Table has <1 rows ({len(table.rows)}), skipping")
            return False
        
        # Get first row (header) text - check multiple rows in case header spans multiple
        header_texts = []
        rows_to_check = min(3, len(table.rows))  # Check first 3 rows for headers
        
        for row_idx in range(rows_to_check):
            row_texts = [cell.text.strip().lower() for cell in table.rows[row_idx].cells]
            # Skip completely empty rows
            if any(t for t in row_texts):
                header_texts.extend(row_texts)
        
        # Join all potential headers
        all_headers = ' '.join(header_texts)
        
        print(f"       🔍 Table has {len(table.rows)} rows, {len(table.columns)} columns")
        print(f"       🔍 First row cells: {[cell.text.strip() for cell in table.rows[0].cells]}")
        print(f"       🔍 All header candidates: {header_texts[:6]}")  # Show first 6
        print(f"       🔍 Combined text: '{all_headers[:100]}'")  # First 100 chars
        
        # Check for skills table indicators - VERY FLEXIBLE
        skills_keywords = ['skill', 'skills', 'technology', 'technologies', 'competency', 'competencies', 
                          'technical', 'proficiency', 'expertise', 'tool', 'tools', 'qualification']
        years_keywords = ['years', 'experience', 'years used', 'years of experience', 'exp', 'yrs', 
                         'year', 'duration']
        last_used_keywords = ['last used', 'last', 'recent', 'most recent', 'latest', 'when', 'current']
        
        has_skill_col = any(kw in all_headers for kw in skills_keywords)
        has_years_col = any(kw in all_headers for kw in years_keywords)
        has_last_used_col = any(kw in all_headers for kw in last_used_keywords)
        
        # Also check if table has exactly 3 columns (Skill, Years, Last Used pattern)
        has_three_cols = len(table.columns) == 3
        
        print(f"       📊 Detection results:")
        print(f"          - Has 3 columns: {has_three_cols} (actual: {len(table.columns)})")
        print(f"          - Has skill column: {has_skill_col}")
        print(f"          - Has years column: {has_years_col}")
        print(f"          - Has last_used column: {has_last_used_col}")
        
        # It's a skills table if:
        # 1. Has skill keyword AND (years OR last_used keyword)
        # 2. OR has 3 columns with years AND last_used (common pattern)
        is_skills = (has_skill_col and (has_years_col or has_last_used_col)) or \
                    (has_three_cols and has_years_col and has_last_used_col)
        
        print(f"       {'✅ IS SKILLS TABLE' if is_skills else '❌ NOT SKILLS TABLE'}")
        
        return is_skills
    
    def _fill_skills_table(self, table):
        """Fill skills table with candidate's skills data using exact 3-column format"""
        if len(table.rows) < 1:
            print(f"     ⚠️  Table has no rows, cannot fill")
            return 0
        
        # Get header row to identify columns
        header_row = table.rows[0]
        header_texts = [cell.text.strip() for cell in header_row.cells]
        header_texts_lower = [h.lower() for h in header_texts]
        
        print(f"     📋 Filling skills table...")
        print(f"     📋 Table headers: {header_texts}")
        print(f"     📋 Table has {len(table.rows)} rows initially")
        
        # CRITICAL: Check if table has standard 3-column format (SKILL_NAME, YEARS_USED, LAST_USED)
        # If headers don't match, set them to the standard format
        expected_headers = ['SKILL_NAME', 'YEARS_USED', 'LAST_USED']
        
        # Detect if we need to standardize headers
        needs_header_fix = False
        if len(table.columns) == 3:
            # Check if headers match our expected format
            skill_keywords = ['skill', 'technology', 'competency', 'technical', 'tool']
            years_keywords = ['years', 'experience', 'exp', 'yrs', 'used']
            last_keywords = ['last', 'recent', 'latest']
            
            col0_is_skill = any(kw in header_texts_lower[0] for kw in skill_keywords)
            col1_is_years = any(kw in header_texts_lower[1] for kw in years_keywords) if len(header_texts_lower) > 1 else False
            col2_is_last = any(kw in header_texts_lower[2] for kw in last_keywords) if len(header_texts_lower) > 2 else False
            
            # If headers don't match pattern or are placeholders, standardize them
            if not (col0_is_skill and col1_is_years and col2_is_last):
                needs_header_fix = True
                print(f"     🔧 Standardizing table headers to: {expected_headers}")
                for idx, expected in enumerate(expected_headers):
                    if idx < len(header_row.cells):
                        header_row.cells[idx].text = expected
        
        # Find column indices - FLEXIBLE matching
        skill_col = None
        years_col = None
        last_used_col = None
        
        skill_keywords = ['skill', 'technology', 'competency', 'technical', 'tool', 'expertise', 'proficiency']
        years_keywords = ['years', 'experience', 'exp', 'yrs', 'years used', 'used']
        last_keywords = ['last', 'recent', 'latest', 'last used', 'most recent']
        
        for idx, header in enumerate(header_texts_lower):
            # Match skill column
            if any(kw in header for kw in skill_keywords) and skill_col is None:
                skill_col = idx
                print(f"     ✓ Skill column: {idx} ('{header_texts[idx]}')")
            # Match years column
            elif any(kw in header for kw in years_keywords) and years_col is None:
                years_col = idx
                print(f"     ✓ Years column: {idx} ('{header_texts[idx]}')")
            # Match last used column
            elif any(kw in header for kw in last_keywords) and last_used_col is None:
                last_used_col = idx
                print(f"     ✓ Last Used column: {idx} ('{header_texts[idx]}')")
        
        # Fallback: If columns not detected by keywords, assume standard 3-column layout
        if skill_col is None and len(table.columns) >= 3:
            print(f"     ℹ️  Using default 3-column layout (columns 0,1,2)")
            skill_col = 0
            years_col = 1
            last_used_col = 2
        elif skill_col is None:
            print(f"     ⚠️  No skill column found and table has {len(table.columns)} columns")
            return 0
        
        # Get comprehensive skills with detailed experience mapping
        skills_data = self._extract_skills_with_details()
        
        print(f"     📊 Enhanced skills extraction complete: {len(skills_data) if skills_data else 0} skills with detailed experience data")
        
        if not skills_data or len(skills_data) == 0:
            print(f"     ⚠️  No comprehensive skills data found, trying fallback extraction...")
            
            # Enhanced fallback: try multiple extraction methods
            raw_skills = self.resume_data.get('skills', [])
            experience = self.resume_data.get('experience', [])
            
            print(f"     🔄 Fallback extraction from {len(raw_skills)} raw skills and {len(experience)} experience entries")
            
            if raw_skills or experience:
                # Parse individual skills from multiple sources
                parsed_skills = []
                
                # From skills section
                if raw_skills:
                    parsed_skills.extend(self._parse_individual_skills(raw_skills))
                
                # From experience bullets (quick extraction)
                if experience:
                    for exp in experience:
                        if isinstance(exp, dict):
                            details = exp.get('details', [])
                            for detail in details:
                                exp_skills = self._parse_individual_skills([str(detail)])
                                parsed_skills.extend(exp_skills)
                
                # Remove duplicates and limit
                unique_skills = list(set(parsed_skills))[:20]
                
                # Calculate total experience for years
                total_years = self._calculate_total_experience_years()
                
                # Create fallback skills data with better estimation
                skills_data = []
                for skill in unique_skills:
                    years_exp = self._estimate_skill_years(skill, total_years)
                    
                    skills_data.append({
                        'skill': skill, 
                        'years': years_exp, 
                        'last_used': 'Current',
                        'total_experience': min(total_years, 6)
                    })
                
                # Sort by estimated experience
                skills_data.sort(key=lambda x: x.get('total_experience', 0), reverse=True)
                
                print(f"     ✓ Fallback generated {len(skills_data)} skills with {total_years} years career experience")
            else:
                print(f"     ❌ No skills data available from any source")
                return 0
        
        # Clear existing data rows (keep header)
        rows_to_delete = []
        for i in range(1, len(table.rows)):
            rows_to_delete.append(i)
        
        # Delete from bottom to top to avoid index issues
        for i in reversed(rows_to_delete):
            table._element.remove(table.rows[i]._element)
        
        # Add skills rows
        filled_count = 0
        print(f"     🔄 Adding {len(skills_data)} skill rows to table...")
        
        for skill_info in skills_data:  # Insert ALL skills
            # Add new row
            new_row = table.add_row()
            
            skill_name = skill_info.get('skill', '')
            
            # Fill skill name
            if skill_col is not None:
                new_row.cells[skill_col].text = skill_name
            
            # Fill years
            if years_col is not None:
                new_row.cells[years_col].text = skill_info.get('years', '')
            
            # Fill last used
            if last_used_col is not None:
                new_row.cells[last_used_col].text = skill_info.get('last_used', '')
            
            filled_count += 1
            if filled_count <= 3:
                print(f"        ✓ Added: {skill_name}")
        
        print(f"     ✅ Successfully filled {filled_count} skill rows")
        return filled_count
    
    def _parse_individual_skills(self, skills_raw):
        """Parse individual skill names from long description strings.
        Extracts clean tool/technology names from descriptive text.
        
        Example:
        Input: "Skilled in updating fiber records, creating documentation using Excel, GIS software..."
        Output: ["Excel", "GIS Software", ...]
        """
        import re
        individual_skills = []
        
        # Common skill/technology patterns to extract
        known_patterns = [
            # Programming languages
            r'\b(Python|Java|JavaScript|TypeScript|C\+\+|C#|Ruby|PHP|Go|Rust|Swift|Kotlin|Scala)\b',
            # Cloud platforms
            r'\b(AWS|Azure|Google Cloud|GCP|Oracle Cloud)\b',
            r'\b(Amazon Web Services|Microsoft Azure)\b',
            # DevOps tools
            r'\b(Docker|Kubernetes|Jenkins|GitLab|GitHub|Terraform|Ansible|Chef|Puppet)\b',
            r'\b(CI/CD|Git|SVN|Mercurial)\b',
            # Databases
            r'\b(MySQL|PostgreSQL|MongoDB|Redis|Oracle|SQL Server|Cassandra|DynamoDB)\b',
            # Microsoft Office
            r'\b(Excel|Word|PowerPoint|Outlook|Access|Microsoft Office|MS Office)\b',
            # Fiber optic / Telecom tools
            r'\b(OTDR|CDD|OFCW|AOSS|GIS|Bluebeam|AutoCAD)\b',
            r'\b(Fiber Splicing|Fiber Records|Circuit Vision)\b',
            # Operating Systems
            r'\b(Windows|Linux|Unix|macOS|Ubuntu|CentOS|Red Hat)\b',
            # Web frameworks
            r'\b(React|Angular|Vue|Django|Flask|Spring|Express|Node\.js)\b',
            # Other common tools
            r'\b(Photoshop|Illustrator|Figma|Sketch|InVision)\b'
        ]
        
        for skill_line in skills_raw:
            skill_text = skill_line if isinstance(skill_line, str) else str(skill_line)
            skill_text = skill_text.strip()
            
            # Strategy 1: Extract known technologies/tools using patterns
            for pattern in known_patterns:
                matches = re.finditer(pattern, skill_text, re.IGNORECASE)
                for match in matches:
                    skill_name = match.group(0).strip()
                    if skill_name and len(skill_name) >= 2:
                        individual_skills.append(skill_name)
            
            # Strategy 2: Handle clean comma-separated lists (short lines)
            if len(skill_text) < 100 and ',' in skill_text:
                # Remove common prefixes first
                cleaned_text = re.sub(r'^(skilled in|proficient in|experience with|knowledge of|expertise in)\s+', '', skill_text, flags=re.IGNORECASE)
                parts = re.split(r',\s*(?:and\s+)?', cleaned_text)
                for part in parts:
                    # Clean each part
                    part = part.strip().lstrip('•–—-*● ')
                    # Remove action verbs at start
                    part = re.sub(r'^(using|creating|updating|managing|implementing|configuring|analyzing|monitoring|troubleshooting)\s+', '', part, flags=re.IGNORECASE)
                    # Remove trailing descriptive phrases
                    part = re.sub(r'\s+(for|to|with|in|on|at)\s+.*$', '', part)
                    part = part.strip()
                    
                    # Only keep if it looks like a clean skill name (2-40 chars, starts with uppercase or number)
                    if 2 <= len(part) <= 40 and not part.lower().startswith(('and ', 'or ', 'the ', 'a ')):
                        # Title case if all lowercase
                        if part.islower():
                            part = part.title()
                        individual_skills.append(part)
            
            # Strategy 3: Extract capitalized terms (likely proper nouns = tools/technologies)
            elif len(skill_text) > 100:
                # Look for capitalized words/acronyms that are likely tool names
                capitalized = re.findall(r'\b([A-Z][A-Za-z]*(?:\s+[A-Z][A-Za-z]*)*|[A-Z]{2,})\b', skill_text)
                for cap in capitalized:
                    # Filter out common words
                    if cap.lower() not in ['skilled', 'hands', 'experience', 'proficient', 'experienced', 'including', 'for', 'and', 'the', 'with']:
                        if 2 <= len(cap) <= 40:
                            individual_skills.append(cap)
            
            # Strategy 4: Extract from "like X, Y, and Z" patterns
            like_patterns = re.findall(r'(?:like|such as|including)\s+([A-Za-z0-9\s,\.]+?)(?:\s+for|\s+to|\s+and\s+[a-z]+ing|$)', skill_text, re.IGNORECASE)
            for pattern_match in like_patterns:
                items = re.split(r',\s*(?:and\s+)?', pattern_match)
                for item in items:
                    item = item.strip().strip('.')
                    if 2 <= len(item) <= 40 and not item.lower().startswith(('for ', 'to ', 'and ', 'or ')):
                        individual_skills.append(item)
        
        # Remove duplicates while preserving order, and clean up
        seen = set()
        unique_skills = []
        
        # Words to filter out (common/generic terms, action verbs, descriptive phrases)
        filter_words = {
            'network', 'software', 'tools', 'system', 'platform', 'technology',
            'experienced', 'skilled', 'proficient', 'hands', 'knowledge',
            'fiber records', 'cable preparation', 'fusion splicing infrastructure',
            'managing fiber cables', 'distribution boxes', 'network plans',
            'updating fiber records', 'creating documentation', 'monitoring networks',
            'analyzing fiber performance', 'including cable preparation',
            'and network plans', 'and gitlab ci/cd', 'devops tools'
        }
        
        for skill in individual_skills:
            skill = skill.strip()
            skill_lower = skill.lower()
            
            # Skip if in filter list
            if skill_lower in filter_words:
                continue
            
            # Skip if it's just a common action phrase
            if any(skill_lower.startswith(prefix) for prefix in ['updating ', 'creating ', 'managing ', 'including ', 'and ']):
                continue
            
            # Skip very long descriptive phrases (likely not a tool name)
            if len(skill) > 35:
                continue
            
            # Prefer longer versions (e.g., "Google Cloud Platform" over "Google Cloud")
            # Check if this is a substring of any existing skill
            is_substring = False
            for existing in unique_skills:
                if skill_lower in existing.lower() and skill_lower != existing.lower():
                    is_substring = True
                    break
            
            if is_substring:
                continue
            
            # Remove any existing skills that are substrings of this one
            unique_skills = [s for s in unique_skills if s.lower() not in skill_lower or s.lower() == skill_lower]
            
            # Add if not already seen
            if skill_lower not in seen and len(skill) >= 2:
                seen.add(skill_lower)
                unique_skills.append(skill)
        
        print(f"     ✂️  Parsed {len(unique_skills)} individual skills from {len(skills_raw)} raw entries")
        if unique_skills:
            for i, s in enumerate(unique_skills[:8]):
                print(f"        {i+1}. {s}")
        
        return unique_skills
    
    def _extract_skills_from_experience_bullets(self, experience_entries):
        """Extract technical skills and tools from job experience bullet points.
        Creates comprehensive skill descriptions directly from actual work experience bullets.
        """
        extracted_skills = []
        all_bullets = []
        
        # Collect all bullet points from all experience entries
        for exp in experience_entries:
            if isinstance(exp, dict):
                details = exp.get('details', [])
                for detail in details:
                    if isinstance(detail, str) and len(detail) > 20:
                        all_bullets.append(detail.strip())
        
        print(f"     🔫 Analyzing {len(all_bullets)} experience bullets for comprehensive skill extraction")
        
        # STRATEGY: Extract actual comprehensive skill statements from bullets
        # Look for bullets that describe broad technical competencies
        
        skill_patterns = []
        
        # Pattern 1: Equipment/Technology experience bullets
        equipment_bullets = []
        design_bullets = []
        troubleshooting_bullets = []
        fiber_bullets = []
        network_bullets = []
        
        for bullet in all_bullets:
            bullet_lower = bullet.lower().strip()
            
            # Look for comprehensive technology experience statements
            if any(term in bullet_lower for term in ['routers', 'switches', 'firewalls', 'enterprise', 'network equipment']):
                # Check if it mentions multiple technologies and actions
                action_words = ['configured', 'managed', 'installed', 'designed', 'maintained', 'troubleshot', 'implemented', 'deployed']
                if any(action in bullet_lower for action in action_words) and len(bullet) > 50:
                    equipment_bullets.append(bullet)
            
            # Design and installation experience
            if any(term in bullet_lower for term in ['design', 'install', 'configure', 'local-area', 'wide-area']):
                if len(bullet) > 40 and any(tech in bullet_lower for tech in ['network', 'lan', 'wan', 'enterprise']):
                    design_bullets.append(bullet)
            
            # Troubleshooting and maintenance
            if any(term in bullet_lower for term in ['troubleshoot', 'troubleshooting', 'maintain', 'upgrade', 'configure']):
                if len(bullet) > 40 and any(tech in bullet_lower for tech in ['router', 'switch', 'network', 'firewall']):
                    troubleshooting_bullets.append(bullet)
            
            # Fiber optic specific experience
            if any(term in bullet_lower for term in ['fiber', 'splicing', 'otdr', 'opgw', 'adss']):
                if len(bullet) > 40:
                    fiber_bullets.append(bullet)
            
            # Network architecture and engineering
            if any(term in bullet_lower for term in ['network architecture', 'engineering', 'implementation', 'scalable']):
                if len(bullet) > 40:
                    network_bullets.append(bullet)
        
        # CREATE COMPREHENSIVE SKILL DESCRIPTIONS based on actual experience
        
        # 1. Equipment and Infrastructure Experience
        if equipment_bullets:
            # Extract technologies mentioned across equipment bullets
            all_equipment_text = ' '.join(equipment_bullets).lower()
            technologies = []
            
            if 'router' in all_equipment_text:
                technologies.append('routers')
            if 'switch' in all_equipment_text:
                technologies.append('switches')
            if 'firewall' in all_equipment_text:
                technologies.append('firewalls')
            if 'vpn' in all_equipment_text or 'concentrator' in all_equipment_text:
                technologies.append('VPN concentrators')
            if 'wireless' in all_equipment_text or 'access point' in all_equipment_text:
                technologies.append('wireless access points')
            if 'wan' in all_equipment_text:
                technologies.append('Wide Area Networks')
            
            if technologies:
                tech_str = ', '.join(technologies)
                extracted_skills.append(f"Considerable knowledge and hands-on working experience with enterprise {tech_str}")
        
        # 2. Configuration and Maintenance Experience
        if troubleshooting_bullets:
            # Extract action verbs and create comprehensive description
            all_troubleshooting_text = ' '.join(troubleshooting_bullets).lower()
            actions = []
            
            if 'configur' in all_troubleshooting_text:
                actions.append('configuring')
            if 'upgrad' in all_troubleshooting_text:
                actions.append('upgrading')
            if 'manag' in all_troubleshooting_text:
                actions.append('managing')
            if 'maintain' in all_troubleshooting_text:
                actions.append('maintaining')
            if 'troubleshoot' in all_troubleshooting_text:
                actions.append('troubleshooting')
            if 'install' in all_troubleshooting_text or 'setup' in all_troubleshooting_text:
                actions.append('setting up')
            
            if actions and len(actions) >= 3:
                action_str = ', '.join(actions[:5])
                extracted_skills.append(f"Considerable hands-on working experience {action_str} routers/switches, and firewalls")
        
        # 3. Design and Implementation Experience
        if design_bullets:
            all_design_text = ' '.join(design_bullets).lower()
            if any(term in all_design_text for term in ['local-area', 'wide-area', 'enterprise', 'network']):
                extracted_skills.append("Demonstrated and hands-on ability to design, install and configure in local-area and wide-area enterprise networks")
        
        # 4. Fiber Optic Engineering Experience
        if fiber_bullets:
            all_fiber_text = ' '.join(fiber_bullets).lower()
            if len(fiber_bullets) >= 2:  # Multiple fiber-related bullets
                extracted_skills.append("Considerable hands-on experience engineering and design experience in fiber optic networking industry")
        
        # 5. Network Architecture Experience
        if network_bullets:
            all_network_text = ' '.join(network_bullets).lower()
            if any(term in all_network_text for term in ['architect', 'implement', 'scalable', 'fault-tolerant']):
                extracted_skills.append("In-depth experience designing installing and troubleshooting local-area and wide-area enterprise networks")
        
        # 6. Additional technical skills from specific bullets
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if 'monitoring' in bullet_lower and 'it' in bullet_lower and len(bullet) > 30:
                if not any('monitoring' in skill.lower() for skill in extracted_skills):
                    extracted_skills.append("Experience monitoring IT")
                break
        
        print(f"     ✅ Created {len(extracted_skills)} comprehensive skill descriptions from experience bullets")
        for i, skill in enumerate(extracted_skills):
            print(f"        {i+1}. {skill[:80]}...")
        
        return extracted_skills
    
    def _calculate_total_experience_years(self):
        """Calculate total years of experience from work history."""
        experience = self.resume_data.get('experience', [])
        if not experience:
            return 2  # Default minimum
        
        current_year = 2025
        earliest_year = current_year
        latest_year = 0
        
        import re
        for exp in experience:
            duration = exp.get('duration', '') if isinstance(exp, dict) else ''
            year_matches = re.findall(r'(20\d{2})', str(duration))
            
            if year_matches:
                start_year = int(year_matches[0])
                end_year = int(year_matches[-1]) if len(year_matches) > 1 else current_year
                
                # Check for present/current
                if 'present' in str(duration).lower() or 'current' in str(duration).lower():
                    end_year = current_year
                
                earliest_year = min(earliest_year, start_year)
                latest_year = max(latest_year, end_year)
        
        # Calculate total years
        if latest_year > 0:
            total_years = max(1, latest_year - earliest_year)
        else:
            total_years = 2  # Default
        
        print(f"     📅 Total experience: {total_years}+ years (from {earliest_year} to {latest_year or current_year})")
        return total_years
    
    def _group_skills_by_category(self, skills_raw):
        """Group related skills into professional categories with descriptive text.
        Follows industry standard format for managed staffing systems.
        """
        import re
        
        # Parse individual tools/technologies from raw text
        individual_skills = self._parse_individual_skills(skills_raw)
        
        # Define skill categories with professional descriptions
        skill_groups = {
            'Programming Languages': {
                'keywords': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'go', 'rust', 'swift', 'kotlin', 'scala'],
                'template': 'Considerable knowledge and hands-on experience with programming languages including {skills}',
                'skills': []
            },
            'Cloud Platforms & Services': {
                'keywords': ['aws', 'azure', 'google cloud', 'gcp', 'cloud'],
                'template': 'Experience with cloud platforms and services including {skills}',
                'skills': []
            },
            'DevOps & Containerization': {
                'keywords': ['docker', 'kubernetes', 'jenkins', 'gitlab', 'github', 'terraform', 'ansible', 'ci/cd', 'git'],
                'template': 'Proficient in DevOps tools and practices including {skills}',
                'skills': []
            },
            'Database Technologies': {
                'keywords': ['mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sql server', 'database'],
                'template': 'Experience working with database technologies including {skills}',
                'skills': []
            },
            'Microsoft Office Suite': {
                'keywords': ['excel', 'word', 'powerpoint', 'outlook', 'access', 'microsoft office', 'ms office'],
                'template': 'Proficient in Microsoft Office Suite including {skills}',
                'skills': []
            },
            'Network & Fiber Optic Tools': {
                'keywords': ['otdr', 'cdd', 'ofcw', 'aoss', 'fiber splicing', 'fiber optic', 'network'],
                'template': 'Considerable hands-on experience with fiber optic and network testing tools including {skills}',
                'skills': []
            },
            'Design & Documentation Software': {
                'keywords': ['gis', 'bluebeam', 'autocad', 'circuit vision', 'visio', 'cad'],
                'template': 'Experience with design and documentation software including {skills}',
                'skills': []
            },
            'Operating Systems': {
                'keywords': ['windows', 'linux', 'unix', 'macos', 'ubuntu', 'centos', 'red hat'],
                'template': 'Experience with operating systems including {skills}',
                'skills': []
            },
            'Web Frameworks & Libraries': {
                'keywords': ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'node.js', 'nodejs'],
                'template': 'Proficient in web frameworks and libraries including {skills}',
                'skills': []
            }
        }
        
        # Categorize each skill
        uncategorized = []
        
        for skill in individual_skills:
            skill_lower = skill.lower()
            categorized = False
            
            for category, config in skill_groups.items():
                for keyword in config['keywords']:
                    if keyword in skill_lower or skill_lower in keyword:
                        if skill not in config['skills']:
                            config['skills'].append(skill)
                        categorized = True
                        break
                if categorized:
                    break
            
            if not categorized and len(skill) >= 3:
                uncategorized.append(skill)
        
        # Build professional skill descriptions
        grouped_skills = []
        
        for category, config in skill_groups.items():
            if config['skills']:
                # Create professional description
                skill_list = ', '.join(config['skills'])
                description = config['template'].format(skills=skill_list)
                grouped_skills.append(description)
        
        # Add uncategorized skills as individual entries (if they look legitimate)
        for skill in uncategorized[:5]:  # Limit uncategorized
            if len(skill) >= 4 and not any(word in skill.lower() for word in ['and', 'or', 'the', 'with']):
                grouped_skills.append(f"Experience with {skill}")
        
        print(f"     📦 Grouped into {len(grouped_skills)} skill categories")
        for i, desc in enumerate(grouped_skills[:3]):
            print(f"        {i+1}. {desc[:80]}...")
        
        return grouped_skills
    
    def _extract_skills_with_details(self):
        """Extract comprehensive skills with accurate years and last used info.
        Following the detailed pseudocode logic:
        1. Extract full-sentence comprehensive skill statements from resume
        2. Calculate YEARS USED from actual job date ranges
        3. Calculate LAST USED from most recent job using that skill
        """
        import re
        from datetime import datetime
        
        skills_list = []
        
        # Get data from multiple sources
        experience = self.resume_data.get('experience', [])
        skills_raw = self.resume_data.get('skills', [])
        summary = self.resume_data.get('summary', '')
        
        current_year = datetime.now().year
        
        print(f"     🔍 Comprehensive skills extraction (following pseudocode logic):")
        print(f"        - Experience entries: {len(experience)} jobs")
        print(f"        - Raw skills: {len(skills_raw)} entries")
        print(f"        - Summary text: {len(str(summary))} chars")
        
        # STEP 1: Extract comprehensive skill statements
        # These are full-sentence descriptions, not keywords
        comprehensive_skills = self._extract_comprehensive_skills(experience, skills_raw, summary)
        
        print(f"     📊 Extracted {len(comprehensive_skills)} comprehensive skill statements")
        
        # STEP 2: For each comprehensive skill, calculate years and last used
        # Following the TRUE LOGIC: count only years where skill was actively used
        for summary_skill in comprehensive_skills:
            skill_text = summary_skill['text']
            skill_keywords = summary_skill['keywords']  # Keywords to match in job descriptions
            
            # Find all jobs where this skill is present
            active_years = set()  # Only years where skill was actually used
            last_used_year = None
            
            for job in experience:
                if not isinstance(job, dict):
                    continue
                    
                # Check if skill is present in this job
                if self._skill_is_present(skill_keywords, job):
                    # Extract year range for this job
                    start_year, end_year = self._extract_years_from_duration(job.get('duration', ''))
                    
                    if start_year and end_year:
                        # Add all years in this job's range (skill was used throughout)
                        for year in range(start_year, end_year + 1):
                            active_years.add(year)
                        
                        # Track most recent usage (last job where skill appeared)
                        if last_used_year is None or end_year > last_used_year:
                            last_used_year = end_year
            
            # Calculate YEARS USED and LAST USED per TRUE LOGIC
            if active_years:
                # TRUE LOGIC: Count unique years in the set (not span)
                # Example: skill used 2015-2017 and 2020-2023 = 7 years, not 9
                total_years = len(active_years)
                max_year = max(active_years)
                
                # Check if skill is ongoing (used in current year)
                ongoing = (max_year == current_year)
                
                # Format: count of unique years with "+" if ongoing
                if ongoing:
                    years_str = f"{total_years}+"
                    last_str = str(current_year)
                else:
                    years_str = str(total_years)
                    last_str = str(last_used_year) if last_used_year else str(max_year)
            else:
                # No specific jobs found, use defaults
                years_str = "2+"
                last_str = str(current_year)
            
            skills_list.append({
                'skill': skill_text,
                'years': years_str,
                'last_used': last_str,
                'keywords': skill_keywords
            })
        
        print(f"     ✅ Generated {len(skills_list)} detailed skill entries with calculated years")
        if skills_list:
            print(f"     📋 Sample skills:")
            for i, s in enumerate(skills_list[:3]):
                print(f"        {i+1}. {s['skill'][:60]}... | {s['years']} yrs | Last: {s['last_used']}")
        
        return skills_list
    
    def _extract_comprehensive_skills(self, experience, skills_raw, summary):
        """Extract comprehensive skill statements by synthesizing from all resume sections.
        Returns list of dicts with 'text' (full sentence) and 'keywords' (for matching).
        
        This follows the pseudocode logic: create maximal, non-redundant summaries
        that describe 'what' was done + 'how' or 'with what'.
        """
        import re
        comprehensive_skills = []
        
        # Collect all experience bullets
        all_bullets = []
        for exp in experience:
            if isinstance(exp, dict):
                details = exp.get('details', [])
                for detail in details:
                    if isinstance(detail, str) and len(detail) > 30:
                        all_bullets.append(detail.strip())
        
        print(f"     🔫 Analyzing {len(all_bullets)} experience bullets for comprehensive skills")
        
        # STRATEGY: Group bullets by technical domain and synthesize comprehensive statements
        
        # 1. NETWORKING INFRASTRUCTURE SKILL
        networking_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['router', 'switch', 'firewall', 'network', 'vpn', 'wireless', 'lan', 'wan', 'enterprise']):
                networking_bullets.append(bullet)
        
        if networking_bullets:
            # Extract all networking technologies mentioned
            all_net_text = ' '.join(networking_bullets).lower()
            techs = []
            if 'router' in all_net_text:
                techs.append('routers')
            if 'switch' in all_net_text:
                techs.append('switches')
            if 'vpn' in all_net_text or 'concentrator' in all_net_text:
                techs.append('VPN concentrators')
            if 'firewall' in all_net_text:
                techs.append('firewalls')
            if 'wireless' in all_net_text or 'access point' in all_net_text:
                techs.append('wireless access points')
            
            if techs:
                tech_str = ', '.join(techs)
                comprehensive_skills.append({
                    'text': f"Considerable knowledge and hands-on working experience with enterprise {tech_str}",
                    'keywords': ['network', 'router', 'switch', 'firewall', 'vpn', 'wireless', 'enterprise']
                })
        
        # 2. NETWORK DESIGN & CONFIGURATION SKILL
        design_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['design', 'install', 'configure', 'local-area', 'wide-area', 'lan', 'wan']):
                if any(tech in bullet_lower for tech in ['network', 'enterprise', 'infrastructure']):
                    design_bullets.append(bullet)
        
        if design_bullets:
            comprehensive_skills.append({
                'text': "Demonstrated and hands-on ability to design, install and configure in local-area and wide-area enterprise networks",
                'keywords': ['design', 'install', 'configure', 'local-area', 'wide-area', 'lan', 'wan', 'network']
            })
        
        # 3. TROUBLESHOOTING & MAINTENANCE SKILL
        troubleshoot_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['troubleshoot', 'maintain', 'upgrade', 'configure', 'manage']):
                if any(term in bullet_lower for term in ['router', 'switch', 'firewall', 'network']):
                    troubleshoot_bullets.append(bullet)
        
        if troubleshoot_bullets:
            comprehensive_skills.append({
                'text': "Considerable hands-on working experience configuring, upgrading, managing, maintaining, and troubleshooting routers/switches, and firewalls",
                'keywords': ['configure', 'upgrade', 'manage', 'maintain', 'troubleshoot', 'router', 'switch', 'firewall']
            })
        
        # 4. FIBER OPTIC SKILL
        fiber_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['fiber', 'splicing', 'otdr', 'opgw', 'adss', 'optical', 'cable']):
                fiber_bullets.append(bullet)
        
        if fiber_bullets:
            all_fiber_text = ' '.join(fiber_bullets).lower()
            
            # Check for specific fiber technologies
            fiber_techs = []
            if 'splicing' in all_fiber_text:
                fiber_techs.append('Splicing')
            if 'otdr' in all_fiber_text:
                fiber_techs.append('Otdr')
            if 'opgw' in all_fiber_text or 'adss' in all_fiber_text:
                fiber_techs.append('OPGW & ADSS')
            
            if fiber_techs:
                tech_str = ', '.join(fiber_techs)
                comprehensive_skills.append({
                    'text': f"Considerable knowledge of fiber optic systems and hands-on working experience with fiber installation, splicing, and testing equipment with Fiber, {tech_str}",
                    'keywords': ['fiber', 'splicing', 'otdr', 'opgw', 'adss', 'optical', 'cable']
                })
            else:
                comprehensive_skills.append({
                    'text': "Considerable hands-on experience engineering and design experience in fiber optic networking industry",
                    'keywords': ['fiber', 'optical', 'engineering', 'design']
                })
        
        # 5. NETWORK ARCHITECTURE SKILL
        architecture_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['architect', 'implement', 'scalable', 'fault-tolerant', 'design']):
                if 'network' in bullet_lower:
                    architecture_bullets.append(bullet)
        
        if architecture_bullets:
            comprehensive_skills.append({
                'text': "In-depth experience designing installing and troubleshooting local-area and wide-area enterprise networks",
                'keywords': ['design', 'install', 'troubleshoot', 'local-area', 'wide-area', 'enterprise', 'network']
            })
        
        # 6. MONITORING & PERFORMANCE SKILL
        monitoring_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['monitor', 'performance', 'metrics', 'analyze', 'statistics']):
                monitoring_bullets.append(bullet)
        
        if monitoring_bullets:
            comprehensive_skills.append({
                'text': "Experience performance tuning, monitoring and collecting statistics metrics collection, and disaster recovery",
                'keywords': ['performance', 'monitor', 'metrics', 'statistics', 'tuning', 'recovery']
            })
        
        # 7. DOCUMENTATION SKILL
        doc_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['document', 'record', 'report', 'manual', 'procedure', 'excel', 'gis']):
                doc_bullets.append(bullet)
        
        if doc_bullets:
            all_doc_text = ' '.join(doc_bullets).lower()
            tools = []
            if 'excel' in all_doc_text:
                tools.append('Excel')
            if 'gis' in all_doc_text:
                tools.append('GIS software')
            if 'bluebeam' in all_doc_text:
                tools.append('Bluebeam')
            
            if tools:
                tool_str = ', '.join(tools)
                comprehensive_skills.append({
                    'text': f"Skilled in updating fiber records, creating documentation using {tool_str}",
                    'keywords': ['document', 'record', 'excel', 'gis', 'update', 'create']
                })
        
        # 8. CLOUD/DEVOPS SKILL (if present)
        cloud_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['cloud', 'aws', 'azure', 'docker', 'kubernetes', 'ci/cd', 'devops']):
                cloud_bullets.append(bullet)
        
        if cloud_bullets:
            all_cloud_text = ' '.join(cloud_bullets).lower()
            platforms = []
            if 'aws' in all_cloud_text:
                platforms.append('AWS')
            if 'azure' in all_cloud_text:
                platforms.append('Azure')
            if 'docker' in all_cloud_text:
                platforms.append('Docker')
            if 'kubernetes' in all_cloud_text:
                platforms.append('Kubernetes')
            
            if platforms:
                platform_str = ', '.join(platforms)
                comprehensive_skills.append({
                    'text': f"Experience designing and implementing cloud infrastructure solutions using {platform_str}",
                    'keywords': ['cloud', 'aws', 'azure', 'docker', 'kubernetes', 'infrastructure']
                })
        
        # 9. DATABASE SKILL (if present)
        db_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['database', 'sql', 'mysql', 'postgresql', 'oracle', 'mongodb']):
                db_bullets.append(bullet)
        
        if db_bullets:
            comprehensive_skills.append({
                'text': "Experience designing, installing and configuring database systems with performance monitoring and optimization",
                'keywords': ['database', 'sql', 'mysql', 'postgresql', 'configure', 'performance']
            })
        
        # 10. SECURITY/COMPLIANCE SKILL (if present)
        security_bullets = []
        for bullet in all_bullets:
            bullet_lower = bullet.lower()
            if any(term in bullet_lower for term in ['security', 'compliance', 'policy', 'standard', 'authentication']):
                security_bullets.append(bullet)
        
        if security_bullets:
            comprehensive_skills.append({
                'text': "Experience creating environments for compliance with networking security architecture policies, and standards",
                'keywords': ['security', 'compliance', 'policy', 'standard', 'architecture']
            })
        
        # If we found very few skills, add some generic ones from summary or skills section
        if len(comprehensive_skills) < 3 and skills_raw:
            for skill in skills_raw[:5]:
                skill_text = str(skill) if not isinstance(skill, str) else skill
                if len(skill_text) > 5:
                    comprehensive_skills.append({
                        'text': f"Experience with {skill_text}",
                        'keywords': skill_text.lower().split()[:3]
                    })
        
        print(f"     ✅ Synthesized {len(comprehensive_skills)} comprehensive skill statements")
        for i, skill in enumerate(comprehensive_skills[:5]):
            print(f"        {i+1}. {skill['text'][:70]}...")
        
        return comprehensive_skills
    
    def _skill_is_present(self, skill_keywords, job):
        """Check if skill is present in job description using semantic matching.
        
        Args:
            skill_keywords: List of keywords that indicate this skill
            job: Job dict with 'role', 'company', 'details', etc.
        
        Returns:
            True if skill is clearly present in job context
        """
        # Combine all job text
        role = job.get('role', '')
        company = job.get('company', '')
        details = job.get('details', [])
        
        job_text = f"{role} {company} " + ' '.join([str(d) for d in details])
        job_text_lower = job_text.lower()
        
        # Synonym mapping for better semantic matching
        synonyms = {
            'network': ['network', 'networking', 'lan', 'wan', 'infrastructure'],
            'router': ['router', 'routers', 'routing'],
            'switch': ['switch', 'switches', 'switching'],
            'firewall': ['firewall', 'firewalls', 'security'],
            'configure': ['configure', 'configuration', 'configuring', 'setup', 'set up'],
            'troubleshoot': ['troubleshoot', 'troubleshooting', 'debug', 'diagnose', 'fix', 'resolve'],
            'maintain': ['maintain', 'maintenance', 'maintaining', 'support'],
            'monitor': ['monitor', 'monitoring', 'track', 'tracking', 'observe'],
            'design': ['design', 'designing', 'architect', 'architecture', 'plan', 'planning'],
            'install': ['install', 'installation', 'installing', 'deploy', 'deployment'],
            'fiber': ['fiber', 'fibre', 'optical', 'optic'],
            'splicing': ['splicing', 'splice', 'fusion'],
            'document': ['document', 'documentation', 'documenting', 'record', 'recording'],
            'manage': ['manage', 'managing', 'management', 'administer', 'administering'],
            'upgrade': ['upgrade', 'upgrading', 'update', 'updating']
        }
        
        # Check if any keyword or its synonyms appear in job text
        matches = 0
        matched_keywords = []
        
        for keyword in skill_keywords:
            keyword_lower = keyword.lower()
            
            # Direct match
            if keyword_lower in job_text_lower:
                matches += 1
                matched_keywords.append(keyword)
                continue
            
            # Check synonyms
            if keyword_lower in synonyms:
                for syn in synonyms[keyword_lower]:
                    if syn in job_text_lower:
                        matches += 1
                        matched_keywords.append(keyword)
                        break
        
        # Matching criteria:
        # - 2+ keyword matches = strong confidence
        # - 1 specific technical term = medium confidence
        # - 1 generic term = weak, need more context
        
        if matches >= 2:
            return True
        elif matches == 1:
            # Check if it's a specific technical term (not generic)
            specific_terms = ['router', 'switch', 'firewall', 'fiber', 'otdr', 'splicing', 
                            'aws', 'azure', 'docker', 'kubernetes', 'database', 'sql',
                            'cisco', 'juniper', 'vpn', 'wireless', 'gis', 'excel',
                            'cloud', 'security', 'optical']
            
            for keyword in matched_keywords:
                if keyword.lower() in specific_terms:
                    return True
        
        return False
    
    def _extract_comprehensive_skills_from_experience(self, experience):
        """Extract detailed skill descriptions from work experience to match professional format"""
        detailed_skills = []
        
        for exp in experience:
            if not isinstance(exp, dict):
                continue
                
            # Extract from job details/bullets
            details = exp.get('details', [])
            company = exp.get('company', '')
            role = exp.get('role', '')
            duration = exp.get('duration', '')
            
            # Extract years from this experience
            start_year, end_year = self._extract_years_from_duration(duration)
            years_in_role = (end_year - start_year) if start_year and end_year else 1
            
            for detail in details:
                detail_text = str(detail).strip()
                if len(detail_text) < 20:  # Skip very short details
                    continue
                    
                # Create detailed skill descriptions based on experience bullets
                skill_descriptions = self._create_detailed_skill_descriptions(detail_text, years_in_role, end_year or 2025)
                detailed_skills.extend(skill_descriptions)
        
        print(f"     🔍 Generated {len(detailed_skills)} detailed skill descriptions from work experience")
        return detailed_skills
    
    def _create_detailed_skill_descriptions(self, experience_text, years_in_role, end_year):
        """Create professional, detailed skill descriptions from experience bullets"""
        descriptions = []
        
        # Technology and tool mapping for detailed descriptions
        tech_patterns = {
            'networking': {
                'keywords': ['network', 'router', 'switch', 'firewall', 'wireless', 'lan', 'wan', 'vpn'],
                'template': 'Considerable knowledge of networking and hands-on working experience with enterprise networking infrastructure, routers, switches, and firewalls'
            },
            'fiber_optic': {
                'keywords': ['fiber', 'otdr', 'splicing', 'cable', 'optical'],
                'template': 'Considerable knowledge of fiber optic systems and hands-on working experience with fiber installation, splicing, and testing equipment'
            },
            'cloud': {
                'keywords': ['cloud', 'aws', 'azure', 'gcp', 'docker', 'kubernetes'],
                'template': 'Experience designing and implementing cloud infrastructure solutions using enterprise-grade platforms and containerization technologies'
            },
            'database': {
                'keywords': ['database', 'sql', 'mysql', 'postgresql', 'oracle', 'mongodb'],
                'template': 'Experience designing, installing and configuring database systems with performance monitoring and optimization'
            },
            'monitoring': {
                'keywords': ['monitor', 'performance', 'metrics', 'analytics', 'troubleshoot'],
                'template': 'Experience performance tuning, monitoring and collecting statistics metrics collection, and disaster recovery'
            },
            'security': {
                'keywords': ['security', 'compliance', 'policy', 'authentication', 'encryption'],
                'template': 'Experience creating environments for compliance with networking security architecture policies, and standards'
            },
            'programming': {
                'keywords': ['python', 'java', 'javascript', 'php', 'code', 'develop', 'programming'],
                'template': 'In-depth experience designing, installing and configuring software applications and development environments'
            },
            'documentation': {
                'keywords': ['document', 'report', 'record', 'manual', 'procedure'],
                'template': 'Experience creating detailed technical documentation, procedures, and maintaining accurate project records'
            },
            'project_management': {
                'keywords': ['project', 'manage', 'coordinate', 'plan', 'schedule'],
                'template': 'Experience managing complex technical projects with cross-functional teams and stakeholder coordination'
            },
            'troubleshooting': {
                'keywords': ['troubleshoot', 'diagnose', 'resolve', 'fix', 'repair', 'debug'],
                'template': 'Demonstrated ability to diagnose, troubleshoot and resolve complex technical issues in enterprise environments'
            }
        }
        
        text_lower = experience_text.lower()
        
        # Find matching technology categories
        for tech_category, tech_info in tech_patterns.items():
            if any(keyword in text_lower for keyword in tech_info['keywords']):
                # Customize the description based on specific technologies mentioned
                description = tech_info['template']
                
                # Add specific technologies found in the text
                specific_techs = []
                
                # Extract specific technology names
                tech_extracts = {
                    'AWS|Azure|GCP|Google Cloud': r'\b(AWS|Azure|GCP|Google Cloud|Amazon Web Services)\b',
                    'Docker|Kubernetes': r'\b(Docker|Kubernetes|K8s)\b',
                    'MySQL|PostgreSQL|Oracle': r'\b(MySQL|PostgreSQL|Oracle|SQL Server|MongoDB)\b',
                    'Python|Java|JavaScript': r'\b(Python|Java|JavaScript|PHP|C\+\+|C#)\b',
                    'Cisco|Juniper|HP': r'\b(Cisco|Juniper|HP|Dell|NetGear)\b',
                    'OTDR|Splicing|Fiber': r'\b(OTDR|Splicing|Fiber|Optical|Cable)\b'
                }
                
                for tech_group, pattern in tech_extracts.items():
                    matches = re.findall(pattern, experience_text, re.IGNORECASE)
                    if matches:
                        specific_techs.extend([m.title() for m in matches])
                
                # Enhance description with specific technologies
                if specific_techs:
                    tech_list = ', '.join(specific_techs[:4])  # Limit to 4 technologies
                    if 'enterprise' in description:
                        description += f' including {tech_list}'
                    else:
                        description += f' with {tech_list}'
                
                descriptions.append({
                    'description': description,
                    'category': tech_category,
                    'years': years_in_role,
                    'end_year': end_year
                })
        
        # If no specific patterns matched, create a general skill description
        if not descriptions and len(experience_text) > 30:
            # Extract action and object from the experience
            action_match = re.search(r'^([A-Z][a-z]+(?:ed|ing)?)\s+(.+)', experience_text)
            if action_match:
                action = action_match.group(1)
                object_part = action_match.group(2)
                
                if action.lower() in ['managed', 'coordinated', 'led', 'supervised']:
                    desc = f"Experience managing and coordinating {object_part[:60]}"
                elif action.lower() in ['developed', 'created', 'built', 'designed']:
                    desc = f"Experience developing and creating {object_part[:60]}"
                elif action.lower() in ['implemented', 'deployed', 'installed']:
                    desc = f"Experience implementing and deploying {object_part[:60]}"
                else:
                    desc = f"Experience with {object_part[:60]}"
                
                descriptions.append({
                    'description': desc,
                    'category': 'general',
                    'years': years_in_role,
                    'end_year': end_year
                })
        
        return descriptions
    
    def _convert_simple_skill_to_detailed(self, skill_text, total_years):
        """Convert simple skill names to detailed professional descriptions"""
        skill_lower = skill_text.lower()
        
        # Map simple skills to detailed templates
        skill_mappings = {
            'python': 'In-depth experience designing, installing and configuring Python applications and development environments',
            'java': 'In-depth experience designing, installing and configuring Java applications and enterprise development environments',
            'javascript': 'Experience developing and implementing JavaScript applications with modern frameworks and libraries',
            'networking': 'Considerable knowledge of networking and hands-on working experience with enterprise networking infrastructure',
            'aws': 'Experience designing and implementing AWS cloud infrastructure solutions using enterprise-grade services',
            'azure': 'Experience designing and implementing Microsoft Azure cloud infrastructure solutions and services',
            'docker': 'Experience designing and implementing containerized applications using Docker and orchestration platforms',
            'kubernetes': 'Experience managing and deploying containerized applications using Kubernetes orchestration platform',
            'mysql': 'Experience designing, installing and configuring MySQL database systems with performance optimization',
            'excel': 'Considerable hands-on experience with Microsoft Excel for data analysis, reporting and business intelligence',
            'office': 'Considerable hands-on experience with Microsoft Office suite for business productivity and documentation',
            'project management': 'Experience managing complex technical projects with cross-functional teams and stakeholder coordination',
            'troubleshooting': 'Demonstrated ability to diagnose, troubleshoot and resolve complex technical issues in enterprise environments'
        }
        
        # Check for direct matches or partial matches
        for key, template in skill_mappings.items():
            if key in skill_lower or any(word in skill_lower for word in key.split()):
                return {
                    'description': template,
                    'category': self._categorize_skill(key),
                    'years': min(total_years, 8),
                    'end_year': 2025
                }
        
        # Generic template for unrecognized skills
        if len(skill_text) > 2:
            return {
                'description': f'Experience working with {skill_text} technologies and related tools',
                'category': 'general',
                'years': min(total_years, 5),
                'end_year': 2025
            }
        
        return None
    
    def _extract_detailed_from_section_content(self, section_content, total_years):
        """Extract detailed descriptions from skills section content"""
        descriptions = []
        content_text = str(section_content)
        
        # If the content is already detailed (long text), use it directly
        if len(content_text) > 50:
            # This might already be a detailed description
            descriptions.append({
                'description': content_text[:120] + ('...' if len(content_text) > 120 else ''),
                'category': 'general',
                'years': total_years,
                'end_year': 2025
            })
        else:
            # Parse as simple skills and convert to detailed
            simple_skills = self._parse_individual_skills([content_text])
            for skill in simple_skills:
                detailed = self._convert_simple_skill_to_detailed(skill, total_years)
                if detailed:
                    descriptions.append(detailed)
        
        return descriptions
    
    def _deduplicate_skill_descriptions(self, all_descriptions):
        """Remove duplicate skill descriptions based on category and merge similar ones"""
        category_groups = {}
        
        # Group by category
        for desc in all_descriptions:
            category = desc.get('category', 'general')
            if category not in category_groups:
                category_groups[category] = []
            category_groups[category].append(desc)
        
        # Keep the most comprehensive description per category
        unique_descriptions = []
        for category, descriptions in category_groups.items():
            if not descriptions:
                continue
                
            # Sort by description length (more detailed first) and years of experience
            descriptions.sort(key=lambda x: (len(x.get('description', '')), x.get('years', 0)), reverse=True)
            
            # Take the most detailed description
            best_desc = descriptions[0]
            
            # Merge years from multiple entries in the same category
            total_category_years = max([d.get('years', 0) for d in descriptions])
            most_recent_year = max([d.get('end_year', 0) for d in descriptions])
            
            best_desc['years'] = total_category_years
            best_desc['end_year'] = most_recent_year
            
            unique_descriptions.append(best_desc)
        
        return unique_descriptions
    
    def _categorize_skill(self, skill_name):
        """Categorize a skill into appropriate category"""
        skill_lower = skill_name.lower()
        
        categories = {
            'networking': ['network', 'router', 'switch', 'firewall', 'wireless', 'lan', 'wan'],
            'fiber_optic': ['fiber', 'otdr', 'splicing', 'optical', 'cable'],
            'cloud': ['aws', 'azure', 'gcp', 'cloud', 'docker', 'kubernetes'],
            'database': ['mysql', 'postgresql', 'oracle', 'mongodb', 'database', 'sql'],
            'programming': ['python', 'java', 'javascript', 'php', 'programming', 'code'],
            'security': ['security', 'firewall', 'encryption', 'compliance'],
            'monitoring': ['monitor', 'performance', 'metrics', 'analytics'],
            'project_management': ['project', 'manage', 'coordinate', 'plan'],
            'documentation': ['document', 'report', 'record', 'manual']
        }
        
        for category, keywords in categories.items():
            if any(kw in skill_lower for kw in keywords):
                return category
        
        return 'general'
    
    def _map_skills_to_experience(self, skills, experience):
        """Map each skill to work experience periods to calculate accurate years and last used dates"""
        skill_map = {}
        current_year = 2025
        
        for skill in skills:
            skill_lower = skill.lower()
            experience_periods = []
            
            for exp in experience:
                if not isinstance(exp, dict):
                    continue
                    
                # Check if skill is mentioned in this experience
                company = exp.get('company', '')
                role = exp.get('role', '')
                details = exp.get('details', [])
                duration = exp.get('duration', '')
                
                # Combine all text for searching
                exp_text = f"{role} {company} " + ' '.join([str(d) for d in details])
                
                # Check if skill appears in this experience
                if (skill_lower in exp_text.lower() or 
                    any(word in skill_lower for word in exp_text.lower().split() if len(word) > 3)):
                    
                    # Extract years from duration
                    start_year, end_year = self._extract_years_from_duration(duration)
                    if start_year and end_year:
                        experience_periods.append((start_year, end_year))
            
            if experience_periods:
                # Calculate total years and date range
                experience_periods.sort()
                total_years = 0
                last_used = experience_periods[-1][1]  # End of most recent period
                first_used = experience_periods[0][0]   # Start of earliest period
                
                # Sum up non-overlapping periods
                merged_periods = []
                for start, end in experience_periods:
                    if not merged_periods or start > merged_periods[-1][1]:
                        merged_periods.append((start, end))
                    else:
                        merged_periods[-1] = (merged_periods[-1][0], max(merged_periods[-1][1], end))
                
                for start, end in merged_periods:
                    total_years += (end - start)
                
                skill_map[skill] = {
                    'years': max(1, total_years),  # Minimum 1 year if mentioned
                    'last_used': last_used,
                    'first_used': first_used
                }
        
        return skill_map
    
    def _extract_years_from_duration(self, duration):
        """Extract start and end years from duration string"""
        if not duration:
            return None, None
            
        # Common patterns: "2020-2023", "Jan 2020 - Present", "2020-Current"
        year_matches = re.findall(r'(20\d{2})', str(duration))
        
        if len(year_matches) >= 2:
            start_year = int(year_matches[0])
            end_year = int(year_matches[-1])
        elif len(year_matches) == 1:
            start_year = int(year_matches[0])
            # Check for "present", "current", etc.
            if any(word in duration.lower() for word in ['present', 'current', 'now']):
                end_year = 2025
            else:
                end_year = start_year + 1  # Assume 1 year if only start given
        else:
            return None, None
            
        return start_year, end_year
    
    def _estimate_skill_years(self, skill, total_career_years):
        """Estimate years of experience for a skill based on its type and total career length"""
        skill_lower = skill.lower()
        
        # Technology age estimation
        if any(term in skill_lower for term in ['cloud', 'docker', 'kubernetes', 'react', 'nodejs']):
            return f"{min(total_career_years, 6)} years"
        elif any(term in skill_lower for term in ['python', 'java', 'javascript', 'sql']):
            return f"{min(total_career_years, 8)} years"
        elif any(term in skill_lower for term in ['excel', 'word', 'office', 'windows', 'network']):
            return f"{total_career_years} years"
        else:
            return f"{min(total_career_years, 5)} years"
    
    def _find_matching_resume_section(self, section_key, resume_sections):
        """Find matching resume section with synonyms"""
        # Direct match
        if section_key in resume_sections:
            return resume_sections[section_key]

        synonyms = {
            'experience': ['experience', 'employment', 'work', 'professional', 'career', 'history', 'background'],
            'education': ['education', 'academic', 'qualification', 'academics', 'educational', 'certificates', 'certifications', 'credentials', 'degrees', 'training', 'schooling', 'learning'],
            'skills': ['skills', 'technical', 'competencies', 'expertise', 'abilities', 'proficiencies', 'capabilities'],
            'summary': ['summary', 'objective', 'profile', 'about', 'overview', 'statement', 'introduction'],
            'projects': ['projects', 'portfolio', 'work samples', 'personal projects'],
            'certifications': ['certifications', 'certificates', 'licenses', 'credentials', 'accreditations'],
            'awards': ['awards', 'achievements', 'honors', 'recognition', 'accomplishments'],
            'publications': ['publications', 'papers', 'articles', 'research'],
            'languages': ['languages', 'language skills', 'linguistic'],
            'references': ['references', 'recommendations', 'contacts']
        }

        patterns = synonyms.get(section_key, [section_key])
        for resume_key, content in resume_sections.items():
            key_lower = resume_key.lower()
            if any(p in key_lower for p in patterns):
                return content

        return []
    
    def _extract_all_candidate_sections(self):
        """Extract all sections from candidate resume including additional ones not in template"""
        candidate_sections = {}
        resume_sections = self.resume_data.get('sections', {})
        
        # Define comprehensive section mapping
        section_types = {
            'summary': ['summary', 'objective', 'profile', 'about', 'overview', 'statement', 'introduction'],
            'experience': ['experience', 'employment', 'work', 'professional', 'career', 'history', 'background'],
            'education': ['education', 'academic', 'qualification', 'academics', 'educational', 'certificates', 'certifications', 'credentials', 'degrees', 'training', 'schooling', 'learning'],
            'skills': ['skills', 'technical', 'competencies', 'expertise', 'abilities', 'proficiencies', 'capabilities'],
            'projects': ['projects', 'portfolio', 'work samples', 'personal projects'],
            'certifications': ['certifications', 'certificates', 'licenses', 'credentials', 'accreditations'],
            'awards': ['awards', 'achievements', 'honors', 'recognition', 'accomplishments'],
            'publications': ['publications', 'papers', 'articles', 'research'],
            'languages': ['languages', 'language skills', 'linguistic'],
            'references': ['references', 'recommendations', 'contacts']
        }
        
        # First, extract known section types
        for section_type, synonyms in section_types.items():
            content = self._find_matching_resume_section(section_type, resume_sections)
            if content:
                candidate_sections[section_type] = content
        
        # Then, identify any additional sections not covered above
        mapped_keys = set()
        for synonyms in section_types.values():
            for resume_key in resume_sections.keys():
                if any(syn in resume_key.lower() for syn in synonyms):
                    mapped_keys.add(resume_key)
        
        # Add unmapped sections as additional content
        for resume_key, content in resume_sections.items():
            if resume_key not in mapped_keys and content:
                # Clean up the section name
                clean_name = resume_key.replace('_', ' ').title()
                candidate_sections[f'additional_{resume_key}'] = {
                    'title': clean_name,
                    'content': content
                }
        
        return candidate_sections
        
    def _insert_additional_sections(self, doc, candidate_sections):
        """Insert additional sections from candidate resume that are not in template"""
        additional_inserted = 0
        
        # Get sections that exist in template
        template_sections = set(k.lower() for k in self._primary_anchors.keys() if self._primary_anchors[k] is not None)
        
        # Find insertion point (after last major section or at end)
        insertion_point = len(doc.paragraphs) - 1
        
        # Look for better insertion point (after EDUCATION if exists, or after EMPLOYMENT)
        if self._primary_anchors.get('EDUCATION'):
            edu_idx = self._primary_anchors['EDUCATION']
            # Find end of education section
            for i in range(edu_idx + 1, len(doc.paragraphs)):
                para_text = doc.paragraphs[i].text.strip().upper()
                if len(para_text) < 50 and any(h in para_text for h in ['CERTIFICATIONS', 'PROJECTS', 'AWARDS', 'REFERENCES']):
                    insertion_point = i
                    break
        elif self._primary_anchors.get('EMPLOYMENT'):
            emp_idx = self._primary_anchors['EMPLOYMENT']
            # Find end of employment section
            for i in range(emp_idx + 1, len(doc.paragraphs)):
                para_text = doc.paragraphs[i].text.strip().upper()
                if len(para_text) < 50 and any(h in para_text for h in ['EDUCATION', 'SKILLS', 'CERTIFICATIONS']):
                    insertion_point = i
                    break
        
        print(f"\n📋 Inserting additional sections at paragraph {insertion_point}")
        
        # Insert additional sections
        for section_key, section_data in candidate_sections.items():
            if section_key.startswith('additional_'):
                title = section_data['title']
                content = section_data['content']
                
                if content and len(content) > 0:
                    print(f"  📄 Adding section: {title} ({len(content)} items)")
                    
                    # Insert section heading
                    heading_para = doc.paragraphs[insertion_point]._element
                    new_heading = self._create_paragraph_after(heading_para)
                    new_heading.text = title.upper()
                    
                    # Apply heading formatting
                    for run in new_heading.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                    
                    # Insert content - preserve ALL items
                    for item in content:  # Insert ALL items
                        item_text = str(item).strip()
                        if item_text and len(item_text) > 0:
                            content_para = self._create_paragraph_after(new_heading._element)
                            content_para.text = f"• {item_text}"
                            
                    additional_inserted += 1
                    insertion_point += len(content) + 2  # Account for heading and content paragraphs
        
        return additional_inserted
    
    def _ensure_education_completeness(self):
        """Ensure all education-related content is extracted from candidate resume"""
        education_data = self.resume_data.get('education', []) or []
        
        # If no structured education data, try to extract from all education-related sections
        if not education_data:
            resume_sections = self.resume_data.get('sections', {})
            
            # Look for education content in multiple section types
            all_education_content = []
            
            # Check all education synonyms
            education_synonyms = ['education', 'academic', 'qualification', 'academics', 'educational', 
                                'certificates', 'certifications', 'credentials', 'degrees', 'training', 
                                'schooling', 'learning']
            
            for section_key, content in resume_sections.items():
                key_lower = section_key.lower()
                if any(syn in key_lower for syn in education_synonyms):
                    if isinstance(content, list):
                        all_education_content.extend(content)
                    else:
                        all_education_content.append(str(content))
            
            # Build education entries from all found content
            if all_education_content:
                education_data = self._build_education_from_bullets(all_education_content)
                print(f"  🎓 Enhanced education extraction: found {len(education_data)} entries from {len(all_education_content)} content lines")
        
        return education_data
    
    def _scan_existing_template_sections(self, doc):
        """Comprehensively scan template for existing sections and their exact positions"""
        existing_sections = {}
        
        # Scan all paragraphs for section headings
        print(f"    🔍 Scanning {len(doc.paragraphs)} paragraphs for existing sections...")
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip().upper()
            
            # Skip empty or very long paragraphs (likely content, not headings)
            if not para_text:
                continue
            if len(para_text) > 100:
                if para_idx < 25:  # Debug first 25 paragraphs
                    print(f"       Para {para_idx}: Too long ({len(para_text)} chars): '{para_text[:50]}...'")
                continue
            
            # Check for section headings with various patterns
            section_patterns = {
                'SUMMARY': ['SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE', 'CAREER SUMMARY', 'OVERVIEW'],
                'EMPLOYMENT': ['EMPLOYMENT HISTORY', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EMPLOYMENT', 'EXPERIENCE'],
                'EDUCATION': ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND', 'CERTIFICATES', 'CERTIFICATIONS', 'CREDENTIALS', 'ACADEMICS', 'EDUCATION/CERTIFICATES', 'EDUCATION / CERTIFICATES'],
                'SKILLS': ['SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES', 'EXPERTISE', 'ABILITIES'],
                'PROJECTS': ['PROJECTS', 'PORTFOLIO', 'PERSONAL PROJECTS', 'KEY PROJECTS'],
                'CERTIFICATIONS': ['CERTIFICATIONS', 'CERTIFICATES', 'LICENSES', 'PROFESSIONAL CERTIFICATIONS'],
                'AWARDS': ['AWARDS', 'ACHIEVEMENTS', 'HONORS', 'RECOGNITION', 'ACCOMPLISHMENTS'],
                'PUBLICATIONS': ['PUBLICATIONS', 'PAPERS', 'ARTICLES', 'RESEARCH'],
                'LANGUAGES': ['LANGUAGES', 'LANGUAGE SKILLS'],
                'REFERENCES': ['REFERENCES', 'RECOMMENDATIONS']
            }
            
            # Check if this paragraph matches any section pattern
            for section_key, patterns in section_patterns.items():
                if any(pattern in para_text for pattern in patterns):
                    # Don't overwrite if we already found this section (take first occurrence)
                    if section_key not in existing_sections:
                        existing_sections[section_key] = para_idx
                        print(f"    🔍 Found {section_key} section at paragraph {para_idx}: '{para_text[:50]}'")
                    break
        
        # Also scan tables for skills tables (special case)
        for table_idx, table in enumerate(doc.tables):
            if self._is_skills_table(table):
                # Find the paragraph position of this table
                table_position = self._find_table_paragraph_position(doc, table)
                if 'SKILLS_TABLE' not in existing_sections:
                    existing_sections['SKILLS_TABLE'] = table_position
                    print(f"    🔍 Found SKILLS_TABLE at position {table_position}")
        
        # CRITICAL FIX: Force-check for EDUCATION near EMPLOYMENT (common template pattern)
        if 'EDUCATION' not in existing_sections and 'EMPLOYMENT' in existing_sections:
            employment_idx = existing_sections['EMPLOYMENT']
            # Check 5 paragraphs after EMPLOYMENT for EDUCATION heading
            for offset in range(1, 6):
                check_idx = employment_idx + offset
                if check_idx < len(doc.paragraphs):
                    check_text = doc.paragraphs[check_idx].text.strip().upper()
                    if 'EDUCATION' in check_text and len(check_text) < 300:  # Allow longer text for templates
                        existing_sections['EDUCATION'] = check_idx
                        print(f"    🔍 Found EDUCATION section at paragraph {check_idx}: '{check_text[:50]}' (detected after EMPLOYMENT)")
                        break
        
        return existing_sections
    
    def _find_table_paragraph_position(self, doc, target_table):
        """Find the approximate paragraph position of a table"""
        # This is a best-effort approach since tables don't have direct paragraph indices
        # We'll estimate based on document structure
        total_tables = len(doc.tables)
        table_idx = None
        
        for idx, table in enumerate(doc.tables):
            if table == target_table:
                table_idx = idx
                break
        
        if table_idx is not None:
            # Estimate position based on table index and total paragraphs
            estimated_position = int((table_idx / max(total_tables, 1)) * len(doc.paragraphs))
            return estimated_position
        
        return 0
    
    def _mark_existing_template_sections(self):
        """Mark which sections already exist in template to prevent duplicates"""
        
        # Update the primary anchors with existing template sections
        for section_key, position in self._existing_template_sections.items():
            # Map to the anchor keys used in the rest of the system
            anchor_key = section_key
            if section_key == 'EMPLOYMENT':
                anchor_key = 'EMPLOYMENT' 
            elif section_key == 'SKILLS_TABLE':
                anchor_key = 'SKILLS'
            
            # Update primary anchors to reflect existing sections
            self._primary_anchors[anchor_key] = position
            
            # Mark these sections as found so we don't create duplicates
            if section_key == 'EMPLOYMENT':
                print(f"    ✅ Template has EMPLOYMENT HISTORY at position {position} - will fill existing section")
            elif section_key == 'EDUCATION':
                print(f"    ✅ Template has EDUCATION at position {position} - will fill existing section")
            elif section_key == 'SKILLS' or section_key == 'SKILLS_TABLE':
                print(f"    ✅ Template has SKILLS/SKILLS_TABLE at position {position} - will fill existing section")
            elif section_key == 'SUMMARY':
                print(f"    ✅ Template has SUMMARY at position {position} - will fill existing section")
    
    def _identify_truly_missing_sections(self):
        """Only identify sections that are completely missing from template AND candidate has data for"""
        truly_missing = []
        
        # Only add sections if they're completely missing AND candidate has data
        essential_sections = {
            'EDUCATION': lambda: bool(self.resume_data.get('education') or 
                                    self._find_matching_resume_section('education', self.resume_data.get('sections', {}))),
            'SKILLS': lambda: bool(self.resume_data.get('skills') or 
                                 self._find_matching_resume_section('skills', self.resume_data.get('sections', {}))),
            'EMPLOYMENT': lambda: bool(self.resume_data.get('experience'))
        }
        
        for section, has_data_func in essential_sections.items():
            # Check if section is completely missing from template
            section_exists = (section in self._existing_template_sections or 
                            (section == 'SKILLS' and 'SKILLS_TABLE' in self._existing_template_sections))
            
            if not section_exists and has_data_func():
                truly_missing.append(section)
        
        return truly_missing
    
    def _analyze_template_coverage(self):
        """Analyze what sections the template expects and covers"""
        template_coverage = {}
        
        # Standard sections that should be in most professional templates
        expected_sections = {
            'SUMMARY': ['summary', 'objective', 'profile', 'overview'],
            'EXPERIENCE': ['experience', 'employment', 'work history', 'professional experience'],
            'EDUCATION': ['education', 'academic background', 'qualifications'],
            'SKILLS': ['skills', 'technical skills', 'competencies'],
            'PROJECTS': ['projects', 'portfolio', 'key projects'],
            'CERTIFICATIONS': ['certifications', 'certificates', 'licenses'],
            'AWARDS': ['awards', 'achievements', 'honors', 'recognition'],
            'PUBLICATIONS': ['publications', 'papers', 'articles'],
            'LANGUAGES': ['languages', 'language skills'],
            'REFERENCES': ['references', 'recommendations']
        }
        
        # Check which sections are covered by template anchors
        for section_key in expected_sections:
            if self._primary_anchors.get(section_key) is not None:
                template_coverage[section_key] = 'covered_by_anchor'
            elif any(anchor for anchor in self._all_anchors.get(section_key, [])):
                template_coverage[section_key] = 'covered_by_secondary_anchor'
            else:
                template_coverage[section_key] = 'missing'
        
        return template_coverage
    
    def _identify_missing_template_sections(self):
        """Identify critical sections missing from template that candidate has data for"""
        missing_sections = []
        
        # Check for critical missing sections where candidate has data
        critical_sections = ['EDUCATION', 'EXPERIENCE', 'SKILLS']
        
        for section in critical_sections:
            if self._template_section_coverage.get(section) == 'missing':
                # Check if candidate has data for this section
                has_data = False
                if section == 'EDUCATION':
                    has_data = bool(self.resume_data.get('education') or 
                                   self._find_matching_resume_section('education', self.resume_data.get('sections', {})))
                elif section == 'EXPERIENCE':
                    has_data = bool(self.resume_data.get('experience'))
                elif section == 'SKILLS':
                    has_data = bool(self.resume_data.get('skills') or 
                                   self._find_matching_resume_section('skills', self.resume_data.get('sections', {})))
                
                if has_data:
                    missing_sections.append(section)
        
        return missing_sections
    
    def _add_missing_template_sections(self, doc, missing_sections):
        """Add missing critical sections to template"""
        print(f"  🔧 Adding missing template sections: {missing_sections}")
        
        # Find a good insertion point (after skills table if exists, or after last major section)
        insertion_point = self._find_optimal_insertion_point(doc)
        
        for section in missing_sections:
            if section == 'EDUCATION':
                self._insert_education_section_at_point(doc, insertion_point)
                insertion_point += 3  # Account for heading + content
            elif section == 'SKILLS':
                self._insert_skills_section_at_point(doc, insertion_point)
                insertion_point += 3
            elif section == 'EXPERIENCE':
                self._insert_experience_section_at_point(doc, insertion_point)
                insertion_point += 5  # Experience takes more space
    
    def _find_optimal_insertion_point(self, doc):
        """Find the best place to insert missing sections - AFTER Employment History"""
        # PRIORITY 1: Find EMPLOYMENT HISTORY section and insert after it
        employment_end = None
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            if any(keyword in text for keyword in ['EMPLOYMENT HISTORY', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT']):
                if len(text) < 50:  # Likely a heading
                    # Scan forward to find the end of employment section
                    for j in range(para_idx + 1, min(para_idx + 100, len(doc.paragraphs))):
                        next_text = doc.paragraphs[j].text.strip().upper()
                        # Stop at next major section
                        if any(kw in next_text for kw in ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECTS', 'CERTIFICATIONS']) and len(next_text) < 50:
                            employment_end = j
                            print(f"  📍 Found Employment end at paragraph {j}, will insert EDUCATION here")
                            break
                    # If no next section found, use some paragraphs after employment heading
                    if not employment_end:
                        employment_end = para_idx + 20
                    break
        
        if employment_end:
            return employment_end
        
        # PRIORITY 2: Look for skills tables
        skills_table_end = None
        for table_idx, table in enumerate(doc.tables):
            if self._is_skills_table(table):
                # Find the paragraph after this table
                for para_idx, para in enumerate(doc.paragraphs):
                    if para._element.getparent() == table._element.getparent():
                        skills_table_end = para_idx + 1
                        break
        
        if skills_table_end:
            return skills_table_end
        
        # PRIORITY 3: Find after the last major section
        last_major_section = 0
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            if any(keyword in text for keyword in ['EMPLOYMENT', 'EXPERIENCE', 'EDUCATION', 'SKILLS', 'SUMMARY']):
                if len(text) < 50:  # Likely a heading
                    last_major_section = para_idx
        
        return last_major_section + 5  # Give some space after content
    
    def _insert_education_section_at_point(self, doc, insertion_point):
        """Insert education section at specific point in document"""
        try:
            # Insert EDUCATION heading
            if insertion_point < len(doc.paragraphs):
                anchor_para = doc.paragraphs[insertion_point]
                heading = self._insert_paragraph_after(anchor_para, '')
            else:
                heading = doc.add_paragraph('')
            
            # CRITICAL: Ensure heading has content and formatting (BOLD + UNDERLINE + CAPS)
            if heading:
                # Clear any existing content
                heading.clear()
                # Add EDUCATION text with proper formatting
                run = heading.add_run('EDUCATION')
                run.bold = True
                run.underline = True
                run.font.size = Pt(12)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
                print(f"    📝 Created EDUCATION heading at insertion point {insertion_point}")
            
            # Insert education content
            education_data = self.resume_data.get('education', [])
            if not education_data:
                education_data = self._ensure_education_completeness()
            
            if education_data and heading:
                last_element = heading
                for edu in education_data:  # Insert ALL education entries
                    edu_block = self._insert_education_block(doc, last_element, edu)
                    if edu_block:
                        last_element = edu_block
                
                print(f"    ✅ Added EDUCATION section with {len(education_data)} entries at paragraph {insertion_point}")
                print(f"    📍 EDUCATION heading text: '{heading.text if heading else 'NO HEADING'}'")
                print(f"    📍 Total paragraphs in document now: {len(doc.paragraphs)}")
            
            # Update anchors to reflect new section
            self._primary_anchors['EDUCATION'] = insertion_point
            
            # Mark as inserted to prevent duplication
            self._education_inserted = True
            
            # CRITICAL: Add to protected sections to prevent cleanup deletion
            if not hasattr(self, '_protected_sections'):
                self._protected_sections = []
            self._protected_sections.append('EDUCATION')
            
            # Mark paragraph range as protected
            if not hasattr(self, '_protected_ranges'):
                self._protected_ranges = []
            
            # CRITICAL FIX: Get actual paragraph indices AFTER insertion
            # Find the EDUCATION heading we just inserted (use LAST occurrence if multiple)
            education_start_idx = None
            print(f"    🔍 Searching for EDUCATION heading in {len(doc.paragraphs)} paragraphs...")
            
            for idx, para in enumerate(doc.paragraphs):
                text = (para.text or '').strip().upper()
                if 'EDUCATION' == text or (text.startswith('EDUCATION') and len(text) < 50):
                    education_start_idx = idx  # Keep updating to get LAST occurrence
                    print(f"    📍 Found EDUCATION at paragraph {idx}: '{para.text[:50]}'")
            
            if education_start_idx is not None:
                education_end_idx = min(education_start_idx + 15, len(doc.paragraphs))
                self._protected_ranges.append((education_start_idx, education_end_idx))
                print(f"    🔒 EDUCATION section locked and protected (paras {education_start_idx}-{education_end_idx})")
            else:
                print(f"    ⚠️  Could not find EDUCATION heading to protect!")
                print(f"    🔍 Paragraph texts:")
                for idx, para in enumerate(doc.paragraphs[-10:]):  # Show last 10
                    print(f"        Para {len(doc.paragraphs) - 10 + idx}: '{para.text[:60]}'")
            
        except Exception as e:
            print(f"    ⚠️  Error adding EDUCATION section: {e}")
            import traceback
            traceback.print_exc()
    
    def _insert_skills_section_at_point(self, doc, insertion_point):
        """Insert skills section at specific point in document"""
        try:
            # Insert SKILLS heading
            if insertion_point < len(doc.paragraphs):
                anchor_para = doc.paragraphs[insertion_point]
                heading = self._insert_paragraph_after(anchor_para, 'SKILLS')
            else:
                heading = doc.add_paragraph('SKILLS')
            
            # Format heading (BOLD + UNDERLINE)
            for run in heading.runs:
                run.bold = True
                run.underline = True
                run.font.size = Pt(12)
            
            # Insert skills content
            skills_list = self.resume_data.get('skills', [])
            if skills_list:
                self._insert_skills_bullets(doc, heading, skills_list)
                print(f"    ✅ Added SKILLS section with {len(skills_list)} skills")
            
            # Update anchors
            self._primary_anchors['SKILLS'] = insertion_point
            
            # Mark as inserted to prevent duplication
            self._skills_inserted = True
            
        except Exception as e:
            print(f"    ⚠️  Error adding SKILLS section: {e}")
    
    def _insert_experience_section_at_point(self, doc, insertion_point):
        """Insert experience section at specific point in document"""
        try:
            # Insert EXPERIENCE heading
            if insertion_point < len(doc.paragraphs):
                anchor_para = doc.paragraphs[insertion_point]
                heading = self._insert_paragraph_after(anchor_para, 'EMPLOYMENT HISTORY')
            else:
                heading = doc.add_paragraph('EMPLOYMENT HISTORY')
            
            # Format heading
            for run in heading.runs:
                run.bold = True
                run.font.size = Pt(12)
            
            # Insert experience content
            experience_data = self.resume_data.get('experience', [])
            if experience_data:
                last_element = heading
                for exp in experience_data:  # Insert ALL employment entries
                    exp_block = self._insert_experience_block(doc, last_element, exp)
                    if exp_block:
                        last_element = exp_block
                
                print(f"    ✅ Added EMPLOYMENT HISTORY section with {len(experience_data)} entries")
            
            # Update anchors
            self._primary_anchors['EMPLOYMENT'] = insertion_point
            
            # Mark as inserted to prevent duplication
            self._experience_inserted = True
            
        except Exception as e:
            print(f"    ⚠️  Error adding EMPLOYMENT HISTORY section: {e}")
    
    def _get_template_covered_sections(self):
        """Get list of sections already covered by template processing"""
        covered = set()
        
        # Sections with anchors (primary template sections)
        for section_key, anchor_idx in self._primary_anchors.items():
            if anchor_idx is not None:
                covered.add(section_key.lower())
        
        # Add sections we know were processed
        if self._summary_inserted:
            covered.add('summary')
        if self._experience_inserted:
            covered.add('experience')
        if self._skills_inserted:
            covered.add('skills') 
        if self._education_inserted:
            covered.add('education')
        
        return covered
    
    def _get_uncovered_candidate_sections(self, covered_sections):
        """Get candidate sections not covered by template"""
        uncovered = {}
        
        # Check all candidate sections
        for section_key, content in self._candidate_sections.items():
            section_lower = section_key.lower()
            
            # Skip if already covered by template
            is_covered = False
            for covered in covered_sections:
                if (covered in section_lower or section_lower in covered or
                    any(syn in section_lower for syn in self._get_section_synonyms(covered))):
                    is_covered = True
                    break
            
            if not is_covered and content:
                uncovered[section_key] = content
        
        return uncovered
    
    def _get_section_synonyms(self, section_name):
        """Get synonyms for a section to check coverage"""
        synonyms_map = {
            'experience': ['employment', 'work', 'professional', 'career', 'history'],
            'education': ['academic', 'qualification', 'degrees', 'training', 'schooling'],
            'skills': ['technical', 'competencies', 'expertise', 'abilities'],
            'summary': ['objective', 'profile', 'overview', 'about'],
            'projects': ['portfolio', 'work samples'],
            'certifications': ['certificates', 'licenses', 'credentials'],
            'awards': ['achievements', 'honors', 'recognition']
        }
        return synonyms_map.get(section_name, [])
    
    def _insert_comprehensive_additional_sections(self, doc, uncovered_sections):
        """Insert all uncovered candidate sections in template format"""
        sections_added = 0
        
        # Find insertion point after all main content
        insertion_point = len(doc.paragraphs) - 1
        
        # Look for better insertion point (after education or skills if they exist)
        for section_name in ['EDUCATION', 'SKILLS', 'EMPLOYMENT']:
            section_idx = self._primary_anchors.get(section_name)
            if section_idx is not None:
                # Find end of this section's content
                for i in range(section_idx + 1, len(doc.paragraphs)):
                    para_text = doc.paragraphs[i].text.strip().upper()
                    if len(para_text) < 50 and any(h in para_text for h in ['CERTIFICATIONS', 'PROJECTS', 'AWARDS']):
                        insertion_point = i
                        break
        
        print(f"    📍 Inserting uncovered sections at paragraph {insertion_point}")
        
        # Add each uncovered section in professional format
        for section_key, content in uncovered_sections.items():
            if section_key.startswith('additional_'):
                section_name = section_key.replace('additional_', '').replace('_', ' ').title()
                actual_content = content.get('content', content) if isinstance(content, dict) else content
            else:
                section_name = section_key.replace('_', ' ').title()
                actual_content = content
            
            if actual_content and len(actual_content) > 0:
                # Insert section heading
                if insertion_point < len(doc.paragraphs):
                    anchor_para = doc.paragraphs[insertion_point]
                    heading_para = self._insert_paragraph_after(anchor_para, section_name.upper())
                else:
                    heading_para = doc.add_paragraph(section_name.upper())
                
                # Format heading
                for run in heading_para.runs:
                    run.bold = True
                    run.font.size = Pt(12)
                
                # Insert content
                content_items = actual_content if isinstance(actual_content, list) else [actual_content]
                last_para = heading_para
                
                for item in content_items:  # Insert ALL items
                    item_text = str(item).strip()
                    if item_text and len(item_text) > 0:
                        content_para = self._insert_paragraph_after(last_para, f"• {item_text}")
                        content_para.paragraph_format.left_indent = Inches(0.25)
                        for run in content_para.runs:
                            run.font.size = Pt(10)
                        last_para = content_para
                
                sections_added += 1
                insertion_point += len(content_items) + 2
                print(f"    ✅ Added {section_name} section with {len(content_items)} items")
        
        return sections_added
    
    def _get_existing_template_covered_sections(self):
        """Get sections that are covered by existing template structure"""
        covered = set()
        
        # Add all existing template sections
        for section_key in self._existing_template_sections.keys():
            covered.add(section_key.lower())
            
            # Add synonyms for comprehensive coverage detection
            if section_key == 'EMPLOYMENT':
                covered.update(['experience', 'work', 'professional', 'career', 'history'])
            elif section_key == 'EDUCATION':
                covered.update(['education', 'academic', 'qualification', 'degrees', 'training'])
            elif section_key == 'SKILLS' or section_key == 'SKILLS_TABLE':
                covered.update(['skills', 'technical', 'competencies', 'expertise', 'abilities'])
            elif section_key == 'SUMMARY':
                covered.update(['summary', 'objective', 'profile', 'overview'])
            elif section_key == 'PROJECTS':
                covered.update(['projects', 'portfolio'])
            elif section_key == 'CERTIFICATIONS':
                covered.update(['certifications', 'certificates', 'licenses'])
        
        # Also add sections that were processed during template filling
        if self._summary_inserted:
            covered.update(['summary', 'objective', 'profile'])
        if self._experience_inserted:
            covered.update(['experience', 'employment', 'work'])
        if self._skills_inserted:
            covered.update(['skills', 'technical'])
        if self._education_inserted:
            covered.update(['education', 'academic'])
        
        return covered
    
    def _get_truly_uncovered_candidate_sections(self, template_covered_sections):
        """Get candidate sections that are truly not covered by existing template"""
        uncovered = {}
        
        # STEP 1: Check if candidate has combined EDUCATION/CERTIFICATIONS section
        has_combined_edu_cert = False
        combined_section_key = None
        
        for section_key in self._candidate_sections.keys():
            section_normalized = section_key.lower().replace('_', ' ').replace('-', ' ').replace('/', ' ')
            # Check for combined section patterns
            if ('education' in section_normalized and 'certification' in section_normalized) or \
               ('education' in section_normalized and 'certificate' in section_normalized):
                has_combined_edu_cert = True
                combined_section_key = section_key
                print(f"    🔍 Detected combined EDUCATION/CERTIFICATIONS in candidate: '{section_key}'")
                break
        
        # STEP 2: Process each section
        for section_key, content in self._candidate_sections.items():
            section_lower = section_key.lower()
            
            # SPECIAL CASE: Handle combined EDUCATION/CERTIFICATIONS
            if has_combined_edu_cert:
                if section_key == combined_section_key:
                    # Add the combined section with proper name
                    uncovered['education_certifications'] = content
                    print(f"    📋 Adding combined section as 'EDUCATION/CERTIFICATIONS'")
                    continue
                elif 'education' in section_lower or 'certification' in section_lower or 'certificate' in section_lower:
                    # Skip separate education/cert sections - already in combined
                    print(f"    ⏭️  Skipping '{section_key}' - already included in combined section")
                    continue
            
            # Check if this section is covered by existing template
            is_covered = False
            
            for covered_section in template_covered_sections:
                # Direct match or substring match
                if (covered_section in section_lower or section_lower in covered_section):
                    is_covered = True
                    print(f"    ✓ Section '{section_key}' covered by template section '{covered_section}'")
                    break
                
                # Synonym matching
                if self._sections_are_synonymous(section_lower, covered_section):
                    is_covered = True
                    print(f"    ✓ Section '{section_key}' covered by template (synonym of '{covered_section}')")
                    break
            
            # Only add if truly not covered and has content
            if not is_covered and content:
                # Skip generic additional_ prefixed sections that are already covered
                if section_key.startswith('additional_'):
                    base_section = section_key.replace('additional_', '')
                    if not any(covered in base_section.lower() for covered in template_covered_sections):
                        uncovered[section_key] = content
                        print(f"    📋 Section '{section_key}' is truly uncovered - will add")
                else:
                    uncovered[section_key] = content
                    print(f"    📋 Section '{section_key}' is truly uncovered - will add")
        
        return uncovered
    
    def _sections_are_synonymous(self, section1, section2):
        """Check if two section names are synonymous"""
        synonym_groups = [
            ['experience', 'employment', 'work', 'professional', 'career', 'history'],
            ['education', 'academic', 'qualification', 'degrees', 'training', 'schooling'],
            ['skills', 'technical', 'competencies', 'expertise', 'abilities'],
            ['summary', 'objective', 'profile', 'overview', 'about'],
            ['projects', 'portfolio', 'work samples'],
            ['certifications', 'certificates', 'licenses', 'credentials'],
            ['awards', 'achievements', 'honors', 'recognition']
        ]
        
        for group in synonym_groups:
            if section1 in group and section2 in group:
                return True
        
        return False
    
    def _insert_non_duplicate_additional_sections(self, doc, uncovered_sections):
        """Insert uncovered sections while ensuring no duplicates with existing template"""
        sections_added = 0
        
        if not uncovered_sections:
            return 0
        
        # Find the best insertion point - after all existing template sections
        insertion_point = self._find_post_template_insertion_point(doc)
        print(f"    📍 Inserting additional sections at paragraph {insertion_point}")
        
        # Add each truly uncovered section
        for section_key, content in uncovered_sections.items():
            # Clean up section name
            if section_key.startswith('additional_'):
                section_name = section_key.replace('additional_', '').replace('_', ' ').title()
                actual_content = content.get('content', content) if isinstance(content, dict) else content
            elif section_key == 'education_certifications':
                # Special formatting for combined section
                section_name = 'EDUCATION/ CERTIFICATIONS'
                actual_content = content
            else:
                section_name = section_key.replace('_', ' ').title()
                actual_content = content
            
            # Double-check this section doesn't exist in template
            if self._section_already_exists_in_template(section_name, doc):
                print(f"    ⚠️  Skipping {section_name} - already exists in template")
                continue
            
            if actual_content and len(actual_content) > 0:
                # Insert section heading - always append at end for dynamic sections
                heading_para = doc.add_paragraph(section_name.upper())
                
                # Format heading
                for run in heading_para.runs:
                    run.bold = True
                    run.font.size = Pt(12)
                heading_para.paragraph_format.space_before = Pt(12)
                heading_para.paragraph_format.space_after = Pt(6)
                
                # Insert content - preserve ALL items from candidate resume
                content_items = actual_content if isinstance(actual_content, list) else [actual_content]
                
                for item in content_items:  # Insert ALL items
                    item_text = str(item).strip()
                    if item_text and len(item_text) > 0:
                        content_para = doc.add_paragraph(f"• {item_text}")
                        content_para.paragraph_format.left_indent = Pt(18)  # 0.25 inches
                        for run in content_para.runs:
                            run.font.size = Pt(10)
                        content_para.paragraph_format.space_after = Pt(3)
                
                sections_added += 1
                print(f"    ✅ Added {section_name} section with {len(content_items)} items")
        
        return sections_added
    
    def _find_post_template_insertion_point(self, doc):
        """Find the best point to insert additional sections after all template sections"""
        # Scan document BACKWARDS to find the last non-empty paragraph
        # This ensures we insert AFTER all existing content, not in the middle
        
        last_content_idx = len(doc.paragraphs) - 1
        
        # Scan backwards to find last paragraph with actual content
        for idx in range(len(doc.paragraphs) - 1, -1, -1):
            para = doc.paragraphs[idx]
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Found last content paragraph
            last_content_idx = idx
            break
        
        # Insert AFTER the last content (add 1 to insert after, not replace)
        insertion_point = last_content_idx + 1
        
        print(f"    📍 DEBUG: Last content at para {last_content_idx}, inserting at {insertion_point}")
        
        return insertion_point
    
    def _section_already_exists_in_template(self, section_name, doc):
        """Check if a section with this name already exists in the template"""
        section_upper = section_name.upper()
        
        # Check existing template sections
        for existing_section in self._existing_template_sections.keys():
            if existing_section in section_upper or section_upper in existing_section:
                return True
        
        # Also scan document paragraphs for similar headings
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip().upper()
            if len(para_text) < 100 and para_text:  # Likely a heading
                if section_upper in para_text or para_text in section_upper:
                    return True
        
        return False
    
    def _cleanup_empty_paragraphs(self, doc):
        """Remove excessive empty paragraphs between sections and normalize spacing"""
        removed_count = 0
        section_headings = ['SUMMARY', 'EMPLOYMENT', 'WORK HISTORY', 'EXPERIENCE', 'EDUCATION', 
                          'SKILLS', 'TECHNICAL SKILLS', 'CERTIFICATIONS', 'PROJECTS', 'LANGUAGES']
        
        try:
            # Build set of protected paragraph indices
            protected_indices = set()
            
            # Protect sections by name
            if hasattr(self, '_protected_sections'):
                for idx, para in enumerate(doc.paragraphs):
                    text = (para.text or '').strip().upper()
                    for section in self._protected_sections:
                        if section.upper() in text and len(text) < 50:
                            # Protect this paragraph and next 10
                            for j in range(idx, min(idx + 10, len(doc.paragraphs))):
                                protected_indices.add(j)
                            break
            
            # Protect specific ranges
            if hasattr(self, '_protected_ranges'):
                for start, end in self._protected_ranges:
                    for j in range(start, min(end, len(doc.paragraphs))):
                        protected_indices.add(j)
            
            paragraphs_to_remove = []
            prev_was_empty = False
            prev_was_section = False
            
            for idx, para in enumerate(doc.paragraphs):
                # Skip protected paragraphs
                if idx in protected_indices:
                    prev_was_empty = False
                    prev_was_section = False
                    continue
                
                text = (para.text or '').strip()
                is_empty = not text
                is_section = any(h in text.upper() for h in section_headings) and len(text) < 50
                
                # Remove multiple consecutive empty paragraphs (keep only one for spacing)
                if is_empty:
                    if prev_was_empty or prev_was_section:
                        paragraphs_to_remove.append(para)
                        removed_count += 1
                    prev_was_empty = True
                else:
                    prev_was_empty = False
                
                # Track section headings
                if is_section:
                    prev_was_section = True
                    # Normalize section heading spacing
                    para.paragraph_format.space_before = Pt(12)
                    para.paragraph_format.space_after = Pt(6)
                else:
                    prev_was_section = False
            
            # Remove marked paragraphs
            for para in paragraphs_to_remove:
                self._delete_paragraph(para)
            
            if protected_indices:
                print(f"  Protected {len(protected_indices)} paragraphs from cleanup")
            print(f"  Removed {removed_count} excessive empty paragraphs")
            
        except Exception as e:
            print(f"  Cleanup error: {e}")
    
    def _verify_complete_content_preservation(self):
        """Verify that all candidate resume content has been preserved"""
        print(f"    Verifying complete content preservation...")
        
        # Count original content lines
        original_lines = 0
        resume_sections = self.resume_data.get('sections', {})
        
        for section_content in resume_sections.values():
            if isinstance(section_content, list):
                original_lines += len(section_content)
            else:
                original_lines += 1
        
        # Add structured data lines
        original_lines += len(self.resume_data.get('experience', []))
        original_lines += len(self.resume_data.get('education', []))
        original_lines += len(self.resume_data.get('skills', []))
        
        print(f"    Original resume content lines: {original_lines}")
        print(f"    Content preservation verification complete")
    
    def _convert_to_pdf(self, docx_path, pdf_path):
        """Convert DOCX to PDF"""
        try:
            if HAS_WIN32:
                # Use Word COM to convert
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF format
                doc.Close()
                word.Quit()
                
                return True
            else:
                print("⚠️  PDF conversion requires Microsoft Word")
                return False
                
        except Exception as e:
            print(f"❌ PDF conversion error: {e}")
            return False


def format_word_document(resume_data, template_analysis, output_path):
    """Main function for Word document formatting"""
    formatter = WordFormatter(resume_data, template_analysis, output_path)
    return formatter.format()
