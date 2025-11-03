"""
Ordered Section Renderer
Dynamically aligns resume sections to template order while preserving formatting
"""

from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy


class OrderedSectionRenderer:
    """
    Renders resume sections in template order with consistent formatting
    """
    
    # Standard section order (fallback if template doesn't specify)
    DEFAULT_SECTION_ORDER = [
        "SUMMARY",
        "PROFESSIONAL SUMMARY",
        "PROFILE",
        "EMPLOYMENT HISTORY",
        "WORK EXPERIENCE",
        "EXPERIENCE",
        "EDUCATION",
        "SKILLS",
        "TECHNICAL SKILLS",
        "PROJECTS",
        "CERTIFICATIONS",
        "CERTIFICATES",
        "AWARDS",
        "ACHIEVEMENTS",
        "LANGUAGES",
        "REFERENCES",
        "OTHER"
    ]
    
    def __init__(self, template_path: str):
        """
        Initialize renderer with template
        
        Args:
            template_path: Path to template DOCX file
        """
        self.template_path = template_path
        self.template_doc = Document(template_path)
        self.section_order = self._extract_section_order()
        self.section_styles = self._extract_section_styles()
    
    def _extract_section_order(self) -> List[str]:
        """
        Extract section order from template
        
        Returns:
            List of section names in template order
        """
        sections = []
        
        for para in self.template_doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings
            is_heading = (
                para.style.name.startswith('Heading') or
                (para.runs and para.runs[0].bold and len(text.split()) <= 5) or
                (text.isupper() and len(text.split()) <= 4)
            )
            
            if is_heading:
                sections.append(text)
        
        # Fallback to default order if no sections found
        if not sections:
            sections = self.DEFAULT_SECTION_ORDER
        
        return sections
    
    def _extract_section_styles(self) -> Dict[str, Dict]:
        """
        Extract formatting styles for each section from template
        
        Returns:
            Dict mapping section names to style properties
        """
        styles = {}
        current_section = None
        
        for para in self.template_doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            is_heading = (
                para.style.name.startswith('Heading') or
                (para.runs and para.runs[0].bold and len(text.split()) <= 5) or
                (text.isupper() and len(text.split()) <= 4)
            )
            
            if is_heading:
                current_section = text
                
                # Extract heading style
                style_info = {
                    'style_name': para.style.name,
                    'alignment': para.alignment,
                    'bold': para.runs[0].bold if para.runs else False,
                    'font_size': para.runs[0].font.size if para.runs and para.runs[0].font.size else Pt(12),
                    'font_name': para.runs[0].font.name if para.runs and para.runs[0].font.name else 'Calibri',
                    'color': para.runs[0].font.color.rgb if para.runs and para.runs[0].font.color else None,
                    'uppercase': text.isupper(),
                    'space_before': para.paragraph_format.space_before,
                    'space_after': para.paragraph_format.space_after
                }
                
                styles[current_section] = {
                    'heading': style_info,
                    'content': None
                }
            elif current_section and current_section in styles:
                # Extract content style (first content paragraph under this heading)
                if styles[current_section]['content'] is None:
                    content_style = {
                        'style_name': para.style.name,
                        'alignment': para.alignment,
                        'font_size': para.runs[0].font.size if para.runs and para.runs[0].font.size else Pt(11),
                        'font_name': para.runs[0].font.name if para.runs and para.runs[0].font.name else 'Calibri',
                        'space_before': para.paragraph_format.space_before,
                        'space_after': para.paragraph_format.space_after
                    }
                    styles[current_section]['content'] = content_style
        
        return styles
    
    def _apply_heading_style(self, paragraph, section_name: str):
        """
        Apply heading style to a paragraph
        
        Args:
            paragraph: Paragraph object to style
            section_name: Section name to get style from
        """
        if section_name not in self.section_styles:
            # Use default heading style
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.size = Pt(12)
            return
        
        style_info = self.section_styles[section_name]['heading']
        
        if paragraph.runs:
            run = paragraph.runs[0]
            run.bold = style_info.get('bold', True)
            run.font.size = style_info.get('font_size', Pt(12))
            run.font.name = style_info.get('font_name', 'Calibri')
            
            if style_info.get('color'):
                run.font.color.rgb = style_info['color']
        
        paragraph.alignment = style_info.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        if style_info.get('space_before'):
            paragraph.paragraph_format.space_before = style_info['space_before']
        if style_info.get('space_after'):
            paragraph.paragraph_format.space_after = style_info['space_after']
    
    def _apply_content_style(self, paragraph, section_name: str):
        """
        Apply content style to a paragraph
        
        Args:
            paragraph: Paragraph object to style
            section_name: Section name to get style from
        """
        if section_name not in self.section_styles or not self.section_styles[section_name]['content']:
            # Use default content style
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(11)
            return
        
        style_info = self.section_styles[section_name]['content']
        
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.size = style_info.get('font_size', Pt(11))
            run.font.name = style_info.get('font_name', 'Calibri')
        
        paragraph.alignment = style_info.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        
        if style_info.get('space_before'):
            paragraph.paragraph_format.space_before = style_info['space_before']
        if style_info.get('space_after'):
            paragraph.paragraph_format.space_after = style_info['space_after']
    
    def _find_matching_template_section(self, section_name: str) -> Optional[str]:
        """
        Find matching template section for a given section name
        
        Args:
            section_name: Section name to match
            
        Returns:
            Matching template section name or None
        """
        section_lower = section_name.lower()
        
        # Exact match
        for template_section in self.section_order:
            if section_lower == template_section.lower():
                return template_section
        
        # Partial match
        for template_section in self.section_order:
            if section_lower in template_section.lower() or template_section.lower() in section_lower:
                return template_section
        
        return None
    
    def render(self, mapped_sections: Dict[str, str], output_path: str, 
               contact_info: Optional[Dict] = None):
        """
        Render formatted resume with sections in template order
        
        Args:
            mapped_sections: Dict mapping section names to content
            output_path: Path to save output DOCX
            contact_info: Optional dict with name, email, phone, address
        """
        output_doc = Document()
        
        # Copy styles from template
        try:
            output_doc.styles._element = self.template_doc.styles._element
        except:
            pass
        
        print(f"\n{'='*70}")
        print(f"📝 RENDERING RESUME")
        print(f"{'='*70}\n")
        
        # Add contact information if provided
        if contact_info:
            self._add_contact_info(output_doc, contact_info)
        
        # Track which sections have been added
        added_sections = set()
        
        # Add sections in template order
        for template_section in self.section_order:
            # Find matching section in mapped data
            matched_content = None
            matched_key = None
            
            # Check exact match
            if template_section in mapped_sections:
                matched_content = mapped_sections[template_section]
                matched_key = template_section
            else:
                # Check case-insensitive match
                for section_key, content in mapped_sections.items():
                    if section_key.lower() == template_section.lower():
                        matched_content = content
                        matched_key = section_key
                        break
            
            if matched_content and matched_key not in added_sections:
                # Add section heading
                heading_para = output_doc.add_paragraph(template_section)
                self._apply_heading_style(heading_para, template_section)
                
                # Add section content
                content_lines = matched_content.split('\n')
                for line in content_lines:
                    if line.strip():
                        content_para = output_doc.add_paragraph(line.strip())
                        self._apply_content_style(content_para, template_section)
                
                added_sections.add(matched_key)
                print(f"  ✓ Added section: {template_section}")
        
        # Add any remaining sections not in template order
        remaining_sections = set(mapped_sections.keys()) - added_sections
        if remaining_sections:
            print(f"\n  📌 Adding {len(remaining_sections)} additional sections:")
            
            for section_key in remaining_sections:
                if section_key == '_uncertain':
                    # Skip uncertain sections or add them at the end
                    continue
                
                content = mapped_sections[section_key]
                
                # Add heading
                heading_para = output_doc.add_paragraph(section_key)
                heading_para.runs[0].bold = True
                heading_para.runs[0].font.size = Pt(12)
                
                # Add content
                content_lines = content.split('\n')
                for line in content_lines:
                    if line.strip():
                        output_doc.add_paragraph(line.strip())
                
                print(f"    ✓ {section_key}")
        
        # Save output
        output_doc.save(output_path)
        print(f"\n✅ Resume saved to: {output_path}")
        print(f"{'='*70}\n")
    
    def _add_contact_info(self, doc: Document, contact_info: Dict):
        """
        Add contact information at the top of the document
        
        Args:
            doc: Document object
            contact_info: Dict with name, email, phone, address
        """
        # Name (centered, bold, larger font)
        if contact_info.get('name'):
            name_para = doc.add_paragraph(contact_info['name'])
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if name_para.runs:
                name_para.runs[0].bold = True
                name_para.runs[0].font.size = Pt(16)
        
        # Contact details (centered, smaller font)
        contact_parts = []
        if contact_info.get('email'):
            contact_parts.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_parts.append(contact_info['phone'])
        if contact_info.get('address'):
            contact_parts.append(contact_info['address'])
        
        if contact_parts:
            contact_para = doc.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if contact_para.runs:
                contact_para.runs[0].font.size = Pt(10)
        
        # Add spacing
        doc.add_paragraph()


# Helper function
def render_ordered_resume(template_path: str, mapped_sections: Dict[str, str],
                         output_path: str, contact_info: Optional[Dict] = None):
    """
    Convenience function to render a resume
    
    Args:
        template_path: Path to template DOCX
        mapped_sections: Dict mapping section names to content
        output_path: Path to save output DOCX
        contact_info: Optional contact information
    """
    renderer = OrderedSectionRenderer(template_path)
    renderer.render(mapped_sections, output_path, contact_info)


# Example usage
if __name__ == "__main__":
    # Test data
    mapped_sections = {
        "SUMMARY": "Experienced software engineer with 5 years of expertise.",
        "EMPLOYMENT HISTORY": "Software Engineer at Google (2020-2024).",
        "EDUCATION": "B.Tech in Computer Science from JNTUH.",
        "SKILLS": "Python, Java, AWS, React"
    }
    
    contact_info = {
        "name": "John Doe",
        "email": "john.doe@example.com",
        "phone": "+1-234-567-8900",
        "address": "San Francisco, CA"
    }
    
    render_ordered_resume(
        template_path="template.docx",
        mapped_sections=mapped_sections,
        output_path="formatted_resume.docx",
        contact_info=contact_info
    )
