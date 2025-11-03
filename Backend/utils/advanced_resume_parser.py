"""
Advanced Resume Parser
Extracts comprehensive information from resumes:
- Personal Details: Name, DOB, Email, Phone, Address, LinkedIn
- Experience: Company, Role, Duration, Responsibilities
- Education: Degree, Institution, Year, Grade
- Skills: Technical, Soft, Languages
- Projects, Certifications, Awards
"""

import pdfplumber
from docx import Document
import re
from datetime import datetime
from collections import defaultdict
import os
from functools import lru_cache

# Import intelligent parser for smart section mapping
try:
    from utils.intelligent_resume_parser import get_intelligent_parser
    INTELLIGENT_PARSER_AVAILABLE = True
except ImportError:
    INTELLIGENT_PARSER_AVAILABLE = False
    print("⚠️  Intelligent parser not available, using basic matching")

class ResumeParser:
    """Comprehensive resume parsing"""
    
    def __init__(self, file_path, file_type):
        self.file_path = file_path
        self.file_type = file_type
        self.raw_text = ""
        self.lines = []
        
        # Initialize intelligent parser if available
        self.intelligent_parser = None
        if INTELLIGENT_PARSER_AVAILABLE:
            try:
                self.intelligent_parser = get_intelligent_parser()
                print("✅ Using intelligent section mapper")
            except Exception as e:
                print(f"⚠️  Failed to load intelligent parser: {e}")
                self.intelligent_parser = None
        
    def parse(self):
        """Main parsing method"""
        print(f"\n{'='*70}")
        print(f"📋 PARSING RESUME: {self.file_path.split('/')[-1]}")
        print(f"{'='*70}\n")
        
        # Extract text
        if self.file_type == 'pdf':
            self.raw_text = self._extract_pdf_text()
        else:
            self.raw_text = self._extract_docx_text()
        
        self.lines = [line.strip() for line in self.raw_text.split('\n') if line.strip()]
        
        # Extract all information
        resume_data = {
            'name': self._extract_name(),
            'email': self._extract_email(),
            'phone': self._extract_phone(),
            'address': self._extract_address(),
            'linkedin': self._extract_linkedin(),
            'dob': self._extract_dob(),
            'summary': self._extract_summary(),
            'experience': self._extract_experience(),
            'education': self._extract_education(),
            'skills': self._extract_skills(),
            'projects': self._extract_projects(),
            'certifications': self._extract_certifications(),
            'awards': self._extract_awards(),
            'languages': self._extract_languages(),
            'sections': self._extract_sections(),
            'raw_text': self.raw_text
        }
        
        self._print_parsing_summary(resume_data)
        return resume_data
    
    def _extract_pdf_text(self):
        """Extract text from PDF"""
        try:
            with pdfplumber.open(self.file_path) as pdf:
                return '\n'.join([page.extract_text() or '' for page in pdf.pages])
        except Exception as e:
            print(f"❌ Error extracting PDF text: {e}")
            return ""
    
    def _extract_docx_text(self):
        """Extract text from DOCX (paragraphs, tables, headers/footers)"""
        try:
            doc = Document(self.file_path)
            texts = []

            # Main body paragraphs
            for para in doc.paragraphs:
                if para.text is not None:
                    texts.append(para.text)

            # Tables content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text is not None:
                                texts.append(p.text)

            # Headers and footers
            try:
                for section in doc.sections:
                    for p in section.header.paragraphs:
                        if p.text is not None:
                            texts.append(p.text)
                    for p in section.footer.paragraphs:
                        if p.text is not None:
                            texts.append(p.text)
            except Exception:
                pass

            # Join and normalize lines
            return '\n'.join([t for t in texts if t and t.strip()])
        except Exception as e:
            print(f"❌ Error extracting DOCX text: {e}")
            return ""
    
    def _extract_name(self):
        """Extract candidate name robustly using context around contact lines"""
        lines = self.lines

        # Helper: find first line index that contains email or phone
        email_idx = phone_idx = None
        email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b')
        phone_patterns = [
            r'(?:\+?\d{1,3}[-\s]?)?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}',  # US formats
            r'(?:\+?91[-\s]?)?\d{10}',  # India +91XXXXXXXXXX or 10 digits
            r'\+?\d{1,3}[-\s]?\d{2,5}[-\s]?\d{5,10}',  # Intl generic
        ]
        phone_regexes = [re.compile(p) for p in phone_patterns]

        for idx, line in enumerate(lines[:30]):
            if email_idx is None and email_pattern.search(line):
                email_idx = idx
            if phone_idx is None and any(r.search(line) for r in phone_regexes):
                phone_idx = idx
            if email_idx is not None and phone_idx is not None:
                break

        # Try lines above the earliest contact line
        pivot = min([i for i in [email_idx, phone_idx] if i is not None], default=None)
        if pivot is not None:
            for i in range(max(0, pivot - 3), pivot + 1):
                line = lines[i].strip()
                if self._has_contact_info(line):
                    continue
                # Name heuristics: 2-4 words, mostly capitalized
                words = line.split()
                if 1 < len(words) <= 5 and sum(w[:1].isupper() for w in words) >= max(2, len(words) - 1):
                    return line

        # Fallback: scan first 10 non-contact lines for a likely name
        for line in lines[:10]:
            if self._has_contact_info(line):
                continue
            if len(line) < 3 or len(line) > 60:
                continue
            if re.match(r'^[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z\-\.]+){1,3}$', line.strip()):
                return line.strip()

        # Fallback 2: derive from file name
        try:
            base = os.path.basename(self.file_path)
            stem = os.path.splitext(base)[0]
            print(f"  📄 Name extraction from filename: '{base}' → '{stem}'")
            
            # Remove UUID prefix if present (e.g., "60ee09b2-c949-490f-aafe-7995a2a71be8_Calvin_McGuire...")
            if '_' in stem and len(stem.split('_')[0]) > 30:
                stem = '_'.join(stem.split('_')[1:])  # Remove UUID part
                print(f"  🔄 Removed UUID prefix: '{stem}'")
            
            # Remove common words and separators
            tokens = re.split(r'[\W_]+', stem)
            blacklist = {"resume", "cv", "profile", "updated", "final", "copy", "doc", "docx", "pdf", "state", "of", "va", "original"}
            tokens = [t for t in tokens if t and t.lower() not in blacklist]
            print(f"  🔍 Name tokens after filtering: {tokens}")
            
            if 1 <= len(tokens) <= 4:
                name_guess = ' '.join(tokens)
                # Capitalize words
                final_name = ' '.join(w[:1].upper() + w[1:] for w in name_guess.split())
                print(f"  ✅ Extracted name from filename: '{final_name}'")
                return final_name
        except Exception as e:
            print(f"  ❌ Error extracting name from filename: {e}")
            pass

        return "Unknown Candidate"
    
    def _extract_email(self):
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, self.raw_text)
        return match.group(0) if match else ""
    
    def _extract_phone(self):
        """Extract phone number (supports international formats)"""
        phone_patterns = [
            r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US formats
            r'(?:\+?91[-.\s]?)?\d{10}',  # India
            r'\+?\d{1,3}[-.\s]?\d{2,5}[-.\s]?\d{5,10}',  # Intl generic
        ]

        for pattern in phone_patterns:
            match = re.search(pattern, self.raw_text)
            if match:
                return match.group(0)
        return ""
    
    def _extract_address(self):
        """Extract address"""
        # Look for lines with address keywords
        address_keywords = ['street', 'avenue', 'road', 'city', 'state', 'zip', 'postal']
        for line in self.lines:
            if any(keyword in line.lower() for keyword in address_keywords):
                return line
        return ""
    
    def _extract_linkedin(self):
        """Extract LinkedIn profile"""
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        match = re.search(linkedin_pattern, self.raw_text, re.IGNORECASE)
        return match.group(0) if match else ""
    
    def _extract_dob(self):
        """Extract date of birth"""
        dob_patterns = [
            r'(?:DOB|Date of Birth|Birth Date)[\s:]+(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            r'(?:DOB|Date of Birth|Birth Date)[\s:]+(\w+ \d{1,2},? \d{4})',
        ]
        
        for pattern in dob_patterns:
            match = re.search(pattern, self.raw_text, re.IGNORECASE)
            if match:
                return match.group(1)
        return ""
    
    def _extract_summary(self):
        """Extract professional summary/objective - with or without explicit heading."""
        summary_keywords = ['summary', 'objective', 'profile', 'about']
        
        # METHOD 1: Detect certification line + summary paragraph pattern (PRIORITY)
        # This should be checked FIRST because it's more specific and reliable
        # Pattern: "PMP | Certified Scrum Master | Agile Practitioner" followed by summary paragraph
        for i, line in enumerate(self.lines[:15]):
            line_lower = line.lower()
            
            # Check if this line looks like a certifications line (short, has cert keywords, has |)
            cert_keywords = ['pmp', 'certified', 'scrum', 'agile', 'practitioner', 'professional']
            has_cert_keywords = sum(1 for kw in cert_keywords if kw in line_lower) >= 2
            has_pipes = line.count('|') >= 1
            is_short = len(line) < 150
            
            if has_cert_keywords and has_pipes and is_short:
                # This looks like a certification line
                # Check if next line is a summary paragraph
                if i + 1 < len(self.lines):
                    next_line = self.lines[i + 1].strip()
                    
                    # Check if next line is a summary paragraph
                    summary_indicators = [
                        'highly accomplished', 'experienced', 'professional', 'expertise',
                        'years of experience', 'proven track record', 'skilled in',
                        'knowledge of', 'background in', 'proficient', 'demonstrated'
                    ]
                    
                    if len(next_line) > 100 and any(ind in next_line.lower() for ind in summary_indicators):
                        # Found certification line + summary pattern!
                        summary_lines = [next_line]
                        
                        # Collect continuation lines
                        j = i + 2
                        while j < min(i + 7, len(self.lines)):
                            cont_line = self.lines[j].strip()
                            
                            # Stop at section header
                            if self._is_section_header(cont_line):
                                break
                            
                            # Stop at short lines
                            if len(cont_line) < 40:
                                break
                            
                            # Continue if substantial text
                            if cont_line and len(cont_line) > 60:
                                summary_lines.append(cont_line)
                            
                            j += 1
                        
                        result = ' '.join(summary_lines)
                        print(f"  ✅ Found summary after certification line: {result[:100]}...")
                        return result
        
        # METHOD 2: Look for explicit summary section heading (but only in header area, first 15 lines)
        # This avoids picking up "Professional Profile" headings in employment section
        for i, line in enumerate(self.lines[:15]):  # Only search header area
            if any(keyword in line.lower() for keyword in summary_keywords) and len(line) < 50:
                # Get next few lines as summary
                summary_lines = []
                for j in range(i+1, min(i+10, len(self.lines))):
                    if self._is_section_header(self.lines[j]):
                        break
                    if self.lines[j].strip():
                        summary_lines.append(self.lines[j])
                if summary_lines:
                    result = ' '.join(summary_lines)
                    print(f"  ✅ Found summary with heading in header area: {result[:100]}...")
                    return result
        
        # METHOD 3: Detect implicit summary - paragraph after name/contact, before first section
        # Find contact info end
        contact_end = 0
        for i, line in enumerate(self.lines[:10]):
            if self._has_contact_info(line):
                contact_end = i + 1
        
        # Find first major section start (experience, education, skills)
        first_section = 999
        for i, line in enumerate(self.lines):
            if self._is_section_header(line):
                first_section = i
                break
        
        print(f"  🔍 Searching for implicit summary between lines {contact_end} and {first_section}")
        
        # Look for substantial paragraphs in header area
        for i in range(contact_end, min(first_section, 15)):
            line = self.lines[i].strip()
            
            # Skip section headers and short lines
            if self._is_section_header(line) or len(line) < 80:
                continue
            
            # Skip contact info
            if self._has_contact_info(line):
                continue
            
            # Check if this looks like summary text (long, descriptive paragraph)
            line_lower = line.lower()
            summary_indicators = [
                'accomplished', 'experienced', 'professional', 'expertise',
                'years of experience', 'proven track record', 'skilled in',
                'knowledge of', 'background in', 'certified', 'proficient',
                'highly', 'dedicated', 'motivated', 'seeking', 'passionate'
            ]
            
            if any(indicator in line_lower for indicator in summary_indicators):
                # Found implicit summary - collect full paragraph
                summary_lines = [line]
                
                # Collect continuation lines
                j = i + 1
                while j < min(i + 5, first_section):
                    next_line = self.lines[j].strip()
                    
                    # Stop at section header
                    if self._is_section_header(next_line):
                        break
                    
                    # Stop at short lines (likely heading of next section)
                    if len(next_line) < 40:
                        break
                    
                    # Continue if it's substantial text
                    if next_line and len(next_line) > 60:
                        summary_lines.append(next_line)
                    
                    j += 1
                
                result = ' '.join(summary_lines)
                print(f"  ✅ Found implicit summary (no heading): {result[:100]}...")
                return result
        
        print(f"  ❌ No summary found")
        return ""
    
    def _extract_experience(self):
        """Extract work experience details with robust pairing (dates + role/company)."""
        experiences = []
        
        # CRITICAL: Search ENTIRE document for employment entries, not just experience section
        # This handles resumes where education/skills sections appear in the middle of employment
        section = self._find_section(['experience', 'work history', 'employment', 'professional experience', 'work experience', 'career history'])
        
        if not section:
            print("  ⚠️  No experience section found, searching entire document")
            # Fallback: use all lines BUT STOP at SKILLS/EDUCATION sections
            section = []
            for idx, line in enumerate(self.lines):
                line_lower = line.lower().strip()
                # CRITICAL: Stop if we hit SKILLS, EDUCATION, or other non-experience sections
                if any(kw in line_lower for kw in ['skills', 'technical skills', 'certifications', 'projects', 'awards']) and len(line) < 50:
                    if idx > 10:  # Only stop if we've collected some lines (avoid stopping too early)
                        print(f"  🛑 Stopped at section: '{line[:40]}'")
                        break
                section.append(line)
        else:
            print(f"  📋 Found experience section with {len(section)} lines")
            # ENHANCEMENT: Also scan lines AFTER the experience section for additional jobs
            # (in case education section interrupted the employment section)
            # BUT: Stop at SKILLS or final EDUCATION section to avoid mixing content
            exp_section_end = -1
            for idx, line in enumerate(self.lines):
                if any(keyword in line.lower() for keyword in ['experience', 'work history', 'employment']):
                    exp_section_end = idx
                    break
            
            if exp_section_end > 0:
                # Add remaining lines but STOP at SKILLS/EDUCATION sections
                remaining_lines = []
                for idx in range(exp_section_end + len(section) + 1, len(self.lines)):
                    line = self.lines[idx]
                    line_lower = line.lower().strip()
                    
                    # Stop if we hit SKILLS or EDUCATION section (these are separate)
                    if any(kw in line_lower for kw in ['skills', 'education', 'certifications', 'projects']) and len(line) < 50:
                        print(f"  🛑 Stopped extended search at: '{line[:40]}'")
                        break
                    
                    remaining_lines.append(line)
                
                if remaining_lines:
                    section.extend(remaining_lines)
                    print(f"  📋 Extended search to include {len(remaining_lines)} additional lines")

        # Normalize lines and strip bullets
        lines = [self._strip_bullet(self._normalize_text(l)) for l in section if l and l.strip()]
        
        print(f"  📝 First 10 lines of experience section:")
        for idx, l in enumerate(lines[:10]):
            print(f"     {idx}: {l[:80]}")
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Skip very short or empty lines
            if len(line) < 3:
                i += 1
                continue

            # Case A: Line with dates - intelligently detect company/role/dates in any format
            if self._contains_date_range(line):
                duration = self._clean_years(line)
                
                # Extract text by removing dates from current line
                text_without_dates = self._extract_role_from_dated_line(line)
                
                company = ''
                role = ''
                
                # Strategy: Check previous line, current line, and next line to find company and role
                
                # Check PREVIOUS line (might be company)
                prev_line = ''
                if i > 0 and not self._is_section_header(lines[i-1]) and not self._contains_date_range(lines[i-1]) and not self._looks_like_detail_bullet(lines[i-1]):
                    prev_line = lines[i-1].strip()
                
                # Check NEXT line (might be role or company)
                next_line = ''
                j = i + 1
                while j < len(lines) and len(lines[j].strip()) < 3:
                    j += 1
                if j < len(lines) and not self._is_section_header(lines[j]) and not self._contains_date_range(lines[j]):
                    next_line = lines[j].strip()
                
                # DECISION LOGIC: Determine what each line contains
                
                # Track if we consumed the next line as role (to skip it when collecting details)
                role_consumed_next_line = False
                
                # If current line has location markers (-, ,), it's likely company+location
                has_location = any(marker in text_without_dates for marker in [' - ', ', ', ' – ', ' — '])
                
                # If current line has role keywords, it might be role
                has_role_keyword = any(kw in text_without_dates.lower() for kw in ['manager', 'coordinator', 'director', 'specialist', 'analyst', 'engineer', 'developer', 'consultant', 'officer', 'assistant', 'supervisor', 'accountant', 'administrator'])
                
                if has_location and not has_role_keyword:
                    # Current line: Company+Location+Dates
                    # Next line: Role
                    company = self._strip_location(text_without_dates).strip()
                    role = next_line
                    print(f"    📌 Format: Company+Location+Dates | Role on next line")
                    print(f"       Company: '{company}' | Role: '{role}'")
                    # CRITICAL: Mark that we consumed the next line (role line)
                    role_consumed_next_line = True
                    
                elif has_location and has_role_keyword:
                    # Current line might have: Company+Location+Role+Dates (all in one)
                    # Try to split by role keyword
                    parts = text_without_dates.split()
                    # Simple heuristic: everything before role keyword is company
                    for idx, word in enumerate(parts):
                        if any(kw in word.lower() for kw in ['manager', 'coordinator', 'director', 'specialist', 'analyst', 'engineer', 'developer', 'consultant', 'officer', 'assistant', 'supervisor', 'accountant']):
                            company = ' '.join(parts[:idx]).strip()
                            role = ' '.join(parts[idx:]).strip()
                            break
                    
                    if not company:
                        # Fallback: strip location and use as company
                        company = self._strip_location(text_without_dates).strip()
                        role = next_line if next_line else text_without_dates
                    
                    print(f"    📌 Format: Company+Location+Role+Dates (combined)")
                    print(f"       Company: '{company}' | Role: '{role}'")
                    
                elif prev_line:
                    # Previous line: Company+Location
                    # Current line: Role+Dates
                    company = self._strip_location(prev_line).strip()
                    role = text_without_dates.strip()
                    print(f"    📌 Format: Company on previous | Role+Dates on current")
                    print(f"       Company: '{company}' | Role: '{role}'")
                    
                else:
                    # Fallback: Current line is role+dates, look for company elsewhere
                    role = text_without_dates.strip()
                    if next_line:
                        company = self._strip_location(next_line).strip()
                        print(f"    📌 Format: Role+Dates on current | Company on next")
                    else:
                        company = ''
                        print(f"    📌 Format: Role+Dates only (no company found)")
                    print(f"       Company: '{company}' | Role: '{role}'")
                
                role = role.strip()
                company = company.strip()

                exp = {
                    'company': company,
                    'role': role,
                    'title': f"{role} - {company}",
                    'duration': duration,
                    'details': []
                }

                # Collect following detail bullets until next date or header
                # CRITICAL: If we consumed the next line as role, start from i+2, not i+1
                k = i + 2 if role_consumed_next_line else i + 1
                while k < len(lines):
                    # CRITICAL FIX: Only break if it's a NEW job entry (not just a detail with dates)
                    # A new job entry has dates AND looks like a role/company line (short, title case)
                    # A detail bullet with dates is long and starts with action verb
                    if self._is_section_header(lines[k]):
                        break
                    
                    # Check if this is a new job entry (has dates AND looks like header, not detail)
                    if self._contains_date_range(lines[k]):
                        # CRITICAL: Check if this looks like a job entry (company + role + dates)
                        # Job entries have: location markers, title case words, or role keywords
                        line_k = lines[k]
                        has_location = any(marker in line_k for marker in [' - ', ', ', ' – ', '-'])
                        has_role_keywords = any(kw in line_k.lower() for kw in ['manager', 'coordinator', 'director', 'specialist', 'analyst', 'engineer', 'developer', 'consultant', 'officer', 'assistant', 'supervisor'])
                        
                        # If it has location markers or role keywords, it's likely a job entry
                        if has_location or has_role_keywords:
                            print(f"    🔄 Found new job entry while collecting details: '{line_k[:60]}'")
                            # It's a new job entry, stop collecting details
                            break
                        
                        # If it looks like a detail bullet, keep it as detail
                        if self._looks_like_detail_bullet(line_k):
                            detail = line_k.strip()
                            if detail and len(detail) > 10:
                                exp['details'].append(detail)
                            k += 1
                            continue
                        else:
                            # It's a new job entry, stop collecting details
                            break
                    
                    detail = self._strip_bullet(lines[k]).strip()
                    # Merge tiny fragments like "Word", "Access)" or "reporting"
                    if exp['details'] and self._should_merge_fragment(exp['details'][-1], detail):
                        exp['details'][-1] = self._merge_fragment(exp['details'][-1], detail)
                        k += 1
                        continue
                    # Skip very short lines and lines that look like section summaries
                    if detail and len(detail) > 10 and not self._looks_like_summary_text(detail):
                        exp['details'].append(detail)
                    k += 1
                
                experiences.append(exp)
                print(f"    ✓ Entry {len(experiences)}: {company or '(no company)'} - {role or '(no role)'} ({duration}) [{len(exp['details'])} details]")
                i = k
                continue

            # CRITICAL: Skip if this is a section heading (now properly filtered)
            if self._is_section_header(line):
                print(f"    ⏭️  Skipping section header: '{line[:40]}'")
                i += 1
                continue
            
            # Case B: Role/company line followed by date line
            if self._looks_like_company_or_role(line):
                company, role = self._parse_company_role_line(line)
                company, role = self._strip_location(company), self._strip_location(role)
                # Look ahead for date
                j = i + 1
                duration = ''
                if j < len(lines) and self._contains_date_range(lines[j]):
                    duration = self._clean_years(lines[j])
                    j += 1
                exp = {
                    'company': company,
                    'role': role,
                    'title': line,
                    'duration': duration,
                    'details': []
                }
                # Collect details
                k = j
                while k < len(lines):
                    if self._contains_date_range(lines[k]) or self._looks_like_company_or_role(lines[k]) or self._is_section_header(lines[k]):
                        break
                    detail = self._strip_bullet(lines[k]).strip()
                    if exp['details'] and self._should_merge_fragment(exp['details'][-1], detail):
                        exp['details'][-1] = self._merge_fragment(exp['details'][-1], detail)
                        k += 1
                        continue
                    # Skip lines that look like summary text
                    if detail and not self._looks_like_summary_text(detail):
                        exp['details'].append(detail)
                    k += 1
                
                # CRITICAL: Only append if we have company OR role (not empty)
                if company or role:
                    experiences.append(exp)
                    print(f"    ✓ Parsed experience: {company} - {role} ({duration})")
                else:
                    print(f"    ⏭️  Skipping invalid entry (no company/role): '{line[:40]}'")
                
                i = k
                continue

            i += 1
        
        print(f"  ✅ Total experiences extracted: {len(experiences)}")
        return experiences
    
    def _parse_company_role_line(self, line):
        """Parse a line to extract company and role"""
        # Common patterns:
        # "Company Name - Role Title"
        # "Role Title at Company Name"
        # "Company Name, Role Title"
        # "Role Title | Company Name"
        
        if ' - ' in line or ' – ' in line:
            # Split by dash
            parts = re.split(r'\s+[-–]\s+', line, maxsplit=1)
            if len(parts) == 2:
                # Usually "Company - Role" or "Role - Company"
                # Heuristic: if first part has common company indicators, it's company
                if any(word in parts[0].lower() for word in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'systems']):
                    return parts[0].strip(), parts[1].strip()
                else:
                    # Assume first is role, second is company
                    return parts[1].strip(), parts[0].strip()
        
        elif ' at ' in line.lower():
            # "Role at Company"
            parts = re.split(r'\s+at\s+', line, flags=re.IGNORECASE, maxsplit=1)
            if len(parts) == 2:
                return parts[1].strip(), parts[0].strip()
        
        elif '|' in line:
            # "Role | Company"
            parts = line.split('|', maxsplit=1)
            if len(parts) == 2:
                return parts[1].strip(), parts[0].strip()
        
        elif ',' in line:
            # "Company, Role" or "Role, Company"
            parts = line.split(',', maxsplit=1)
            if len(parts) == 2:
                # Heuristic: if first part has company indicators
                if any(word in parts[0].lower() for word in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'systems']):
                    return parts[0].strip(), parts[1].strip()
                else:
                    return parts[1].strip(), parts[0].strip()
        
        # If no pattern matched, assume entire line is company or role
        # Check if it looks more like a company name
        if any(word in line.lower() for word in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'systems', 'solutions']):
            return line.strip(), ''
        else:
            # Assume it's a role
            return '', line.strip()
        
        return line.strip(), ''
    
    def _extract_education(self):
        """Extract education - combine degree + institution properly.
        Enhanced to detect High School and similar qualifications and to clean degree text.
        """
        education = []
        section = self._find_section(['education', 'academic', 'qualification', 'academics'])
        if not section:
            print("  ⚠️  No education section found; attempting global scan")
            # Global scan fallback across entire resume lines
            education = self._extract_education_global()
            print(f"  ✅ Global scan produced {len(education)} education entries")
            return education
        
        print(f"  🎓 Found education section with {len(section)} lines")
        
        # Filter out lines that clearly look like work experience (have company/role patterns)
        filtered_section = []
        for line in section:
            line_lower = line.lower()
            
            # Skip lines that look like job titles or company names with employment indicators
            # BUT: Keep "Goal:" lines as they are education-related
            if any(word in line_lower for word in ['coordinator', 'manager', 'director', 'assistant', 'specialist', 'analyst']):
                if not any(edu_word in line_lower for edu_word in ['university', 'college', 'school', 'degree', 'bachelor', 'master', 'goal:']):
                    print(f"    ⚠️  Skipping potential experience line in education: {line[:60]}")
                    continue
            
            # Skip lines that are clearly experience bullets (start with action verbs typical of work)
            line_stripped = line.strip().lstrip('•–—-*● ')
            first_word = line_stripped.split()[0].lower() if line_stripped.split() else ''
            
            work_action_verbs = [
                'managed', 'oversaw', 'coordinated', 'conducted', 'led', 'supervised',
                'maintained', 'provided', 'facilitated', 'assisted', 'processed',
                'tracked', 'monitored', 'collaborated', 'acted', 'demonstrated',
                'proficient', 'applied', 'experienced', 'tested', 'collected'
            ]
            
            # If line starts with work action verb and doesn't contain education keywords, skip it
            if first_word in work_action_verbs:
                if not any(edu_word in line_lower for edu_word in ['university', 'college', 'school', 'degree', 'bachelor', 'master', 'graduation', 'gpa', 'graduated']):
                    print(f"    ⚠️  Skipping work-related bullet in education: {line[:60]}")
                    continue
            
            filtered_section.append(line)
        
        section = filtered_section
        print(f"  🎓 After filtering: {len(section)} lines")

        lines = [self._normalize_text(l) for l in section if l and l.strip()]
        i = 0
        while i < len(lines):
            line = self._strip_bullet(lines[i])
            
            # Skip empty or very short lines
            if len(line) < 5:
                i += 1
                continue

            # Check if this line contains degree keywords
            degree_keywords = [
                'bachelor', 'master', 'phd', 'mba', 'bsc', 'msc', 'b.sc', 'm.sc', 'degree', 'diploma',
                'high school', 'secondary school', 'higher secondary', 'matriculation', 'ssc', 'hsc', 'ged'
            ]
            has_degree = any(kw in line.lower() for kw in degree_keywords)
            has_institution = any(k in line.lower() for k in ['university', 'college', 'school', 'institute', 'academy'])
            
            # Only process if it looks like a degree line
            if has_degree or (has_institution and ':' in line):
                degree, institution = self._parse_degree_institution_line(line)
                year = self._clean_years(line)

                # Clean year tokens and generic suffixes from degree text
                if degree:
                    degree = re.sub(r'\b(?:19|20)\d{2}\b', '', degree)
                    degree = re.sub(r'\b(class of|graduation|graduated|passed out)\b', '', degree, flags=re.IGNORECASE)
                    degree = re.sub(r'\s+', ' ', degree).strip(' ,.;:-')
                
                # If institution is empty, check next line
                if not institution and i + 1 < len(lines):
                    next_line = self._strip_bullet(lines[i + 1])
                    if any(k in next_line.lower() for k in ['university', 'college', 'school', 'institute', 'academy']):
                        institution = self._strip_location(next_line).strip()
                        # Also get year from next line if not already found
                        if not year:
                            year = self._clean_years(next_line)
                        i += 1  # Skip the institution line
                
                # Collect details (but skip lines that look like other education entries or work experience)
                details = []
                j = i + 1
                while j < len(lines) and j < i + 4:
                    detail_line = self._strip_bullet(lines[j])
                    # Stop if we hit another degree entry
                    if any(kw in detail_line.lower() for kw in degree_keywords):
                        break
                    if self._is_section_header(detail_line):
                        break
                    
                    # Skip lines that start with work action verbs (these are work experience, not education)
                    detail_stripped = detail_line.strip().lstrip('•–—-*● ')
                    first_word_detail = detail_stripped.split()[0].lower() if detail_stripped.split() else ''
                    if first_word_detail in ['managed', 'oversaw', 'coordinated', 'conducted', 'maintained', 'provided', 'facilitated']:
                        print(f"      ⚠️  Skipping work bullet in education details: {detail_line[:50]}")
                        j += 1
                        continue
                    
                    # Add as detail if substantial and not an institution-only line
                    is_institution_only = any(k in detail_line.lower() for k in ['university', 'college', 'school']) and len(detail_line.split()) < 6
                    if len(detail_line) > 15 and not is_institution_only:
                        details.append(detail_line)
                    j += 1
                
                # Only add if we have at least a degree or institution AND they look like education content
                # Skip if degree/institution contain work-related terms
                degree_looks_valid = degree and not any(word in degree.lower() for word in ['coordinator', 'managed', 'oversaw', 'conducted', 'maintained'])
                institution_looks_valid = institution or degree_looks_valid  # Institution can be empty if degree is valid
                
                if (degree or institution) and (degree_looks_valid or institution):
                    edu = {
                        'degree': degree.strip(),
                        'institution': institution.strip(),
                        'year': year,
                        'details': details
                    }
                    education.append(edu)
                    print(f"    ✓ Parsed edu: {degree or '(no degree)'} - {institution or '(no inst)'} ({year})")
                else:
                    print(f"    ⚠️  Skipped invalid education entry: degree='{degree[:40] if degree else ''}'")
                
                i = j
                continue

            i += 1
        
        print(f"  ✅ Total education entries extracted: {len(education)}")
        return education

    def _extract_education_global(self):
        """Fallback: scan all lines for education-like entries even if no explicit section header exists."""
        results = []
        seen = set()
        degree_keywords = [
            'bachelor', 'master', 'phd', 'mba', 'bsc', 'msc', 'b.sc', 'm.sc', 'degree', 'diploma',
            'high school', 'secondary school', 'higher secondary', 'matriculation', 'ssc', 'hsc', 'ged'
        ]
        institution_keywords = ['university', 'college', 'school', 'institute', 'academy']
        
        # Expanded work-related terms to filter out job descriptions
        work_terms = ['coordinator', 'manager', 'director', 'assistant', 'specialist', 'analyst',
                     'managed', 'led', 'developed', 'implemented', 'coordinated', 'facilitated',
                     'leveraged', 'proficiently', 'spearheaded', 'collaborated', 'established',
                     'administered', 'directed', 'oversaw', 'supervised', 'executed']
        
        # Action verbs that indicate work experience, not education
        work_action_verbs = ['managed', 'led', 'developed', 'implemented', 'coordinated', 
                            'facilitated', 'leveraged', 'proficiently', 'spearheaded',
                            'engaged', 'directed', 'formulated', 'collaborated', 'configured',
                            'designed', 'created', 'built', 'established', 'conducted']

        for idx, raw in enumerate(self.lines):
            line = self._normalize_text(raw)
            if not line or len(line) < 5 or self._is_section_header(line):
                continue
            low = line.lower()
            
            # CRITICAL: Skip lines that start with work action verbs (job descriptions)
            # These are NOT education entries!
            if any(low.startswith(verb) for verb in work_action_verbs):
                continue
            
            # Skip obvious work lines
            if any(w in low for w in work_terms) and not any(k in low for k in institution_keywords + degree_keywords):
                continue
            
            # Skip long descriptive lines (likely job descriptions, not education)
            if len(line) > 150 and not any(k in low for k in degree_keywords):
                continue
                
            if any(k in low for k in degree_keywords + institution_keywords):
                degree, institution = self._parse_degree_institution_line(line)
                year = self._clean_years(line)
                degree = re.sub(r'\b(?:19|20)\d{2}\b', '', degree or '').strip(' ,.;:-')
                if not (degree or institution):
                    continue
                key = (degree.lower(), institution.lower(), year)
                if key in seen:
                    continue
                seen.add(key)
                results.append({'degree': degree, 'institution': institution, 'year': year, 'details': []})

                # Also peek at the next line if it likely contains the institution/year
                if not institution and idx + 1 < len(self.lines):
                    nxt = self._normalize_text(self.lines[idx + 1])
                    if any(k in nxt.lower() for k in institution_keywords):
                        inst2 = self._strip_location(nxt)
                        results[-1]['institution'] = inst2
                        if not results[-1]['year']:
                            results[-1]['year'] = self._clean_years(nxt)

        # Prefer most recent entries: sort by year descending when available
        def year_key(e):
            y = e.get('year') or ''
            m = re.search(r'((?:19|20)\d{2})$', y)
            return int(m.group(1)) if m else -1
        results.sort(key=year_key, reverse=True)
        # Limit to 5
        return results[:5]

    # ---------------------- Helpers ----------------------
    def _normalize_text(self, s):
        # Remove odd unicode like 'ï¼' and zero-width, collapse spaces
        if not s:
            return ''
        # Remove zero-width and BOM characters
        s = s.replace('\u200b', ' ').replace('\ufeff', ' ')
        # Remove common garbage unicode
        s = s.replace('ï¼', ' ').replace('â€"', '-').replace('–', '-').replace('—', '-')
        # Normalize quotes and apostrophes
        s = s.replace(''', "'").replace(''', "'").replace('"', '"').replace('"', '"')
        # Collapse multiple spaces
        s = re.sub(r'\s+', ' ', s).strip()
        return s

    def _strip_location(self, s):
        if not s:
            return s
        # Remove trailing city/state fragments after comma
        s = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', s, flags=re.IGNORECASE)
        return s.strip()

    def _clean_years(self, s):
        """Normalize date ranges.
        Prefer 'Mon YYYY-Mon YYYY' if months present; else 'YYYY-YYYY' or 'YYYY'.
        """
        if not s:
            return ''
        s = self._normalize_text(s)
        s = re.sub(r'[–—]', '-', s)
        s = re.sub(r'\s*(to|–|—|-)\s*', '-', s, flags=re.IGNORECASE)

        present = bool(re.search(r'\b(current|present)\b', s, re.IGNORECASE))

        month_map = {
            'january': 'Jan', 'jan': 'Jan',
            'february': 'Feb', 'feb': 'Feb',
            'march': 'Mar', 'mar': 'Mar',
            'april': 'Apr', 'apr': 'Apr',
            'may': 'May',
            'june': 'Jun', 'jun': 'Jun',
            'july': 'Jul', 'jul': 'Jul',
            'august': 'Aug', 'aug': 'Aug',
            'september': 'Sept', 'sept': 'Sept', 'sep': 'Sept',
            'october': 'Oct', 'oct': 'Oct',
            'november': 'Nov', 'nov': 'Nov',
            'december': 'Dec', 'dec': 'Dec',
        }
        def abbr(m):
            return month_map.get(m.lower(), m[:3].title())

        my = [m.groups() for m in re.finditer(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s*(?:,\s*)?((?:19|20)\d{2})\b', s, flags=re.IGNORECASE)]
        if my:
            sm, sy = my[0]
            sm = abbr(sm)
            if len(my) >= 2:
                em, ey = my[-1]
                em = abbr(em)
                return f"{sm} {sy}-{em} {ey}"
            if present:
                return f"{sm} {sy}-Present"
            yrs = re.findall(r'\b((?:19|20)\d{2})\b', s)
            if len(yrs) >= 2:
                return f"{sm} {yrs[0]}-{yrs[-1]}"
            return f"{sm} {sy}"

        years = re.findall(r'\b(?:19|20)\d{2}\b', s)
        if len(years) >= 2:
            return f"{years[0]}-{years[-1]}"
        elif len(years) == 1:
            return years[0]
        return ''
    
    def _parse_degree_institution_line(self, line):
        """Parse a line to extract degree and institution"""
        # Common patterns:
        # "Bachelor of Science, MIT"
        # "Master's in Computer Science - Stanford University"
        # "MBA | Harvard Business School"
        s = self._normalize_text(line)
        s = self._strip_bullet(s)
        # Prefer splitting at an institution keyword boundary
        m = re.search(r'(university|college|school|institute|academy)\b.*', s, flags=re.IGNORECASE)
        if m:
            degree = s[:m.start()].strip(' ,;:-')
            institution = s[m.start():].strip()
            # Remove trailing location fragments
            institution = self._strip_location(institution)
            return degree, institution
        
        if ',' in s:
            parts = s.split(',', maxsplit=1)
            return parts[0].strip(), parts[1].strip()
        
        elif ' - ' in s or ' – ' in s:
            parts = re.split(r'\s+[-–]\s+', s, maxsplit=1)
            if len(parts) == 2:
                return parts[0].strip(), parts[1].strip()
        
        elif '|' in s:
            parts = s.split('|', maxsplit=1)
            if len(parts) == 2:
                return parts[0].strip(), parts[1].strip()
        
        elif ' from ' in s.lower():
            parts = re.split(r'\s+from\s+', s, flags=re.IGNORECASE, maxsplit=1)
            if len(parts) == 2:
                return parts[0].strip(), parts[1].strip()
        
        # If no pattern, return entire line as degree
        return s.strip(), ''

    def _strip_bullet(self, s):
        if not s:
            return s
        return re.sub(r'^[\s•\-–—*●]+', '', s)

    def _extract_role_from_dated_line(self, s):
        """Remove date expressions from a dated line to recover the job title/role text."""
        if not s:
            return ''
        t = self._strip_bullet(self._normalize_text(s))
        # Remove month names + year e.g., 'Aug 2007', 'July 2025'
        t = re.sub(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-zA-Z]*\s+(?:19|20)\d{2}\b', '', t, flags=re.IGNORECASE)
        # Remove standalone years and connectors
        t = re.sub(r'\b(?:19|20)\d{2}\b', '', t)
        t = re.sub(r'\b(to|–|—|-)\b', '', t, flags=re.IGNORECASE)
        t = re.sub(r'\b(current|present)\b', '', t, flags=re.IGNORECASE)
        # Collapse spaces and trim
        t = re.sub(r'\s+', ' ', t).strip(' ,;:- ')
        return t
    
    def _extract_skills(self):
        """Extract skills - preserve all content from SKILLS section as-is.
        Each bullet/line is treated as one skill entry. We avoid splitting on '•' or commas
        to prevent breaking phrases like "Microsoft Office Suite (Excel, Word, Access)".
        Also stitch short continuation lines to the previous bullet when needed.
        """
        skills = []
        
        # DEBUG: Check what sections are available
        print(f"  🔍 Looking for SKILLS section...")
        for idx, line in enumerate(self.lines):
            line_clean = line.lower().strip()
            if any(kw in line_clean for kw in ['skills', 'technical skills']):
                print(f"     Found at line {idx}: '{line[:50]}'")
        
        skills_section = self._find_section(['skills', 'technical skills', 'competencies', 'expertise'])
        
        print(f"  📋 SKILLS section found: {len(skills_section) if skills_section else 0} lines")
        if skills_section:
            for idx, line in enumerate(skills_section[:5]):
                print(f"     {idx+1}. {line[:60]}")
        
        # FALLBACK: If no dedicated skills section, look in Education/Certifications sections
        if not skills_section:
            print(f"  🔍 No dedicated SKILLS section, checking Education/Certifications...")
            education_section = self._find_section(['education', 'certifications', 'education/ certifications'])
            if education_section:
                print(f"  📚 Found education section with {len(education_section)} lines, scanning for skills...")
                for idx, line in enumerate(education_section[:10]):
                    print(f"     Edu {idx+1}. '{line[:60]}'")
                
                # Extract skills from education section
                for raw in education_section:
                    line = raw.strip()
                    if not line:
                        continue
                        
                    # Skip degree/institution lines
                    if any(word in line.lower() for word in ['bachelor', 'master', 'degree', 'university', 'college', 'osha', 'confined', 'space']):
                        continue
                        
                    # Skip lines with dates  
                    if self._contains_date_range(line):
                        continue
                    
                    # Look for technical skill lines
                    skill_text = line.lstrip('•-* ').strip()
                    
                    # DEBUG: Show what we're checking
                    print(f"      🔍 Checking: '{skill_text}' (length: {len(skill_text)})")
                    
                    # Keep if it looks like a skill/technology name (be more lenient)
                    if (2 <= len(skill_text) <= 80 and  # More lenient length
                        skill_text not in ['', ' '] and  # Not empty
                        not skill_text.lower().startswith('bachelor')):  # Not education degree
                        
                        skills_section.append(raw)
                        print(f"    ✅ ADDED skill: '{skill_text}'")
                
                print(f"  📋 Extracted {len(skills_section)} skills from education section")
        
        if skills_section:
            print(f"  🔍 Processing {len(skills_section)} skills section lines:")
            for idx, line in enumerate(skills_section[:5]):
                print(f"     {idx+1}. '{line[:60]}'")
            
            for raw in skills_section:
                t = self._strip_bullet(self._normalize_text(raw)).strip()
                if not t or self._is_section_header(t):
                    continue
                
                # Skip date ranges and long paragraphs
                if self._contains_date_range(t) or len(t) > 300:
                    continue
                
                # CRITICAL: Split individual skills from combined lines
                # Handle cases like "Communication & Customer Service Appointment Scheduling (Athena Program)"
                individual_skills = self._parse_individual_skills_from_line(t)
                
                for skill in individual_skills:
                    if skill and len(skill) > 2:
                        skills.append(skill.strip())
                        print(f"    ✓ Skill: '{skill}'")
                    
            # Remove duplicates while preserving order
            skills = list(dict.fromkeys(skills))
        
        # Debug output: show extracted skills
        if skills:
            print(f"  🛠️  FINAL EXTRACTED {len(skills)} skills:")
            for idx, skill in enumerate(skills[:10]):
                print(f"     {idx+1}. {skill[:80]}")
            if len(skills) > 10:
                print(f"     ... and {len(skills) - 10} more")
        else:
            print(f"  ❌ NO SKILLS EXTRACTED - Check filtering logic!")
        
        print(f"  📊 RETURNING {len(skills)} skills to formatter")
        return skills
    
    def _parse_individual_skills_from_line(self, line):
        """Parse individual skills from a line that may contain multiple skills"""
        if not line:
            return []
        
        # Clean the line
        line_clean = line.strip()
        
        # For single skill lines (common case), just return the line as-is
        # This handles cases like "OPGW & ADDS", "Fiber Splicing", "AutoCAD", etc.
        if len(line_clean) < 80 and not any(sep in line_clean for sep in [',', ';', '|', ' and ', ' & ']):
            return [line_clean]
        
        # For lines with multiple skills separated by delimiters
        individual_skills = []
        
        # Smart split function that respects parentheses
        def smart_split(text, delimiter):
            """Split on delimiter but not inside parentheses"""
            parts = []
            current = []
            paren_depth = 0
            
            for char in text:
                if char == '(':
                    paren_depth += 1
                    current.append(char)
                elif char == ')':
                    paren_depth -= 1
                    current.append(char)
                elif char == delimiter and paren_depth == 0:
                    # Split here
                    parts.append(''.join(current).strip())
                    current = []
                else:
                    current.append(char)
            
            # Add last part
            if current:
                parts.append(''.join(current).strip())
            
            return parts
        
        # Split on common separators (respecting parentheses)
        if ',' in line_clean and '(' in line_clean:
            # Has commas AND parentheses - use smart split
            parts = smart_split(line_clean, ',')
            individual_skills.extend(p for p in parts if p and len(p) > 2)
        elif ',' in line_clean:
            # Simple comma split (no parentheses)
            parts = [p.strip() for p in line_clean.split(',')]
            individual_skills.extend(p for p in parts if p and len(p) > 2)
        elif ';' in line_clean:
            # Handle semicolon-separated skills
            parts = [p.strip() for p in line_clean.split(';')]
            individual_skills.extend(p for p in parts if p and len(p) > 2)
        elif '|' in line_clean:
            # Handle pipe-separated skills
            parts = [p.strip() for p in line_clean.split('|')]
            individual_skills.extend(p for p in parts if p and len(p) > 2)
        else:
            # Single skill line
            individual_skills.append(line_clean)
        
        # Clean up each skill
        cleaned_skills = []
        for skill in individual_skills:
            skill = skill.strip()
            # Remove common prefixes/suffixes
            skill = re.sub(r'^[-•*\s]+', '', skill)  # Remove bullets
            skill = re.sub(r'\s*[-•*\s]+$', '', skill)  # Remove trailing bullets
            skill = skill.strip()
            
            if skill and len(skill) > 1:
                cleaned_skills.append(skill)
        
        return cleaned_skills if cleaned_skills else [line_clean]
    
    def _extract_projects(self):
        """Extract projects"""
        projects = []
        projects_section = self._find_section(['projects', 'key projects', 'project work'])
        
        if projects_section:
            current_project = {}
            for line in projects_section:
                if line and not line.startswith('•') and not self._is_section_header(line):
                    if current_project:
                        projects.append(current_project)
                    current_project = {'name': line, 'details': []}
                elif line:
                    current_project.setdefault('details', []).append(line)
            
            if current_project:
                projects.append(current_project)
        
        return projects
    
    def _extract_certifications(self):
        """Extract certifications - only actual cert names, not long paragraphs"""
        certifications = []
        cert_section = self._find_section(['certifications', 'certificates', 'licenses'])
        
        if cert_section:
            for line in cert_section:
                if not line or self._is_section_header(line):
                    continue
                
                # CRITICAL: Skip long paragraphs (likely summary text, not certifications)
                # Certifications are typically short: "PMP", "AWS Certified", etc.
                # Summary paragraphs are long (>200 chars) and descriptive
                if len(line) > 200:
                    print(f"    ⏭️  Skipping long paragraph (likely summary, not cert): {line[:80]}...")
                    continue
                
                # Skip if it looks like summary text (has summary indicators)
                line_lower = line.lower()
                summary_indicators = [
                    'highly accomplished', 'experienced professional', 'proven track record',
                    'expertise extends', 'committed to', 'years of experience',
                    'demonstrated', 'proficient in', 'adept at'
                ]
                
                if any(indicator in line_lower for indicator in summary_indicators):
                    print(f"    ⏭️  Skipping summary-like text: {line[:80]}...")
                    continue
                
                # This looks like an actual certification
                certifications.append(line)
        
        return certifications
    
    def _extract_awards(self):
        """Extract awards and achievements"""
        awards = []
        awards_section = self._find_section(['awards', 'achievements', 'honors', 'recognition'])
        
        if awards_section:
            for line in awards_section:
                if line and not self._is_section_header(line):
                    awards.append(line)
        
        return awards
    
    def _extract_languages(self):
        """Extract languages"""
        languages = []
        lang_section = self._find_section(['languages', 'language proficiency'])
        
        if lang_section:
            for line in lang_section:
                if line and not self._is_section_header(line):
                    languages.append(line)
        
        return languages
    
    def _extract_sections(self):
        """Extract all sections for flexible formatting"""
        sections = defaultdict(list)
        current_section = None
        
        for line in self.lines:
            if self._is_section_header(line):
                current_section = line.lower().strip()
            elif current_section and line:
                sections[current_section].append(line)
        
        return dict(sections)
    
    def _find_section(self, keywords):
        """Find section by keywords - with INTELLIGENT semantic matching"""
        section_lines = []
        in_section = False
        section_start_idx = -1
        
        # ENHANCEMENT: Expand keywords with common synonyms
        expanded_keywords = list(keywords)
        
        # Add synonyms based on primary keyword
        primary = keywords[0].lower()
        
        if 'experience' in primary or 'employment' in primary:
            expanded_keywords.extend([
                'professional experience', 'work experience', 'work history',
                'employment history', 'career history', 'professional background',
                'relevant employment history'
                # NOTE: Removed 'professional summary' - it's a summary section, not experience!
            ])
        elif 'education' in primary:
            expanded_keywords.extend([
                'educational background', 'academic background', 'academic qualifications',
                'qualifications', 'education background', 'academics'
            ])
        elif 'skills' in primary:
            expanded_keywords.extend([
                'technical skills', 'core competencies', 'key skills',
                'professional skills', 'areas of expertise', 'competencies',
                'skill set', 'expertise', 'technical competencies'
            ])
        elif 'summary' in primary:
            expanded_keywords.extend([
                'professional summary', 'career summary', 'profile',
                'professional profile', 'career objective', 'objective',
                'executive summary', 'career overview', 'about me'
            ])
        
        print(f"  🔍 Searching for section with keywords: {keywords[0]}")
        print(f"     Expanded to {len(expanded_keywords)} variants")
        
        # First pass: find the section start
        for idx, line in enumerate(self.lines):
            line_lower = line.lower().strip()
            
            # Skip very long lines (not section headers)
            if len(line) > 100:
                continue
            
            # METHOD 1: Try intelligent matching if available
            if self.intelligent_parser and len(line) < 50:
                try:
                    matched = self.intelligent_parser._match_heading(
                        line,
                        expanded_keywords
                    )
                    if matched:
                        in_section = True
                        section_start_idx = idx
                        print(f"  ✅ Found '{keywords[0]}' at line {idx}: '{line[:50]}' (AI match → '{matched}')")
                        break
                except Exception as e:
                    print(f"  ⚠️  AI matching failed: {e}")
            
            # METHOD 2: Exact and partial string matching (fallback)
            for keyword in expanded_keywords:
                keyword_lower = keyword.lower()
                
                # Exact match
                if keyword_lower == line_lower:
                    in_section = True
                    section_start_idx = idx
                    print(f"  ✅ Found '{keywords[0]}' at line {idx}: '{line[:50]}' (exact match)")
                    break
                
                # Partial match (for short headers)
                if len(line_lower) < 50 and keyword_lower in line_lower:
                    # Ensure it's not part of a larger word
                    if re.search(rf'\b{re.escape(keyword_lower)}\b', line_lower):
                        in_section = True
                        section_start_idx = idx
                        print(f"  ✅ Found '{keywords[0]}' at line {idx}: '{line[:50]}' (partial match)")
                        break
            
            if in_section:
                break
        
        if not in_section:
            print(f"  ❌ Section '{keywords[0]}' not found")
            return section_lines
        
        # Second pass: collect lines until next section header
        collected_count = 0
        for idx in range(section_start_idx + 1, len(self.lines)):
            line = self.lines[idx].strip()
            
            # Skip empty lines
            if not line:
                continue
            
            # Stop at next major section header
            if self._is_section_header(line):
                # Don't stop if this is just a subsection or continuation
                line_lower = line.lower()
                major_sections = [
                    'experience', 'education', 'summary', 'certifications',
                    'projects', 'awards', 'skills', 'languages', 'profile'
                ]
                
                # Check if this is TRULY a new major section
                is_major_section = any(sect in line_lower for sect in major_sections)
                
                # CRITICAL: Don't stop at the same section type we're collecting
                # (e.g., if collecting experience, don't stop at "Professional Experience")
                is_same_section_type = any(kw.lower() in line_lower for kw in expanded_keywords)
                
                # SPECIAL CASE: If collecting certifications and hit "Professional Profile", STOP
                # "Professional Profile" is employment history, not certifications
                if 'certifications' in primary and 'profile' in line_lower:
                    print(f"    🛑 Stopped at next section: '{line[:40]}'")
                    break
                
                if is_major_section and not is_same_section_type:
                    print(f"    🛑 Stopped at next section: '{line[:40]}'")
                    break
            
            # CRITICAL: For EDUCATION section, stop if we hit skill-like content
            # Skills are often listed without a clear "SKILLS" header
            # Pattern: Multiple short lines (certifications, tools, technologies)
            if 'education' in primary:
                # Check if we're hitting a cluster of short skill-like lines
                # Education entries are typically 2-4 lines per degree
                # Skills are many short single-line items
                
                # If we've collected some education and now hit many short lines, likely skills
                if collected_count >= 2:  # Already have some education
                    # Look ahead: are the next 3 lines also short and skill-like?
                    short_lines_ahead = 0
                    for look_idx in range(idx, min(idx + 3, len(self.lines))):
                        look_line = self.lines[look_idx].strip()
                        if look_line and len(look_line) < 50 and not self._is_section_header(look_line):
                            # Check if it looks like a skill (tool, cert, technology)
                            skill_indicators = [
                                'osha', 'fiber', 'testing', 'software', 'autocad', 'excel',
                                'cisco', 'routing', 'switch', 'configuration', 'voip',
                                'lan', 'wan', 'troubleshooting', 'monitoring', 'cloud',
                                'microsoft', 'office', 'word', 'outlook', 'proficient'
                            ]
                            if any(ind in look_line.lower() for ind in skill_indicators):
                                short_lines_ahead += 1
                    
                    # If we see 2+ skill-like lines ahead, we've hit the skills section
                    if short_lines_ahead >= 2:
                        print(f"    🛑 Stopped at skills content (no header): '{line[:40]}'")
                        break
            
            section_lines.append(self.lines[idx])
            collected_count += 1
            
            # Safety: don't collect more than 100 lines per section
            if collected_count > 100:
                print(f"    ⚠️  Hit safety limit of 100 lines")
                break
        
        print(f"  📋 Collected {len(section_lines)} lines for '{keywords[0]}' section")
        return section_lines
    
    def _is_section_header(self, line):
        """Check if line is a section header"""
        if len(line) > 50:
            return False
        
        section_keywords = [
            'experience', 'education', 'skills', 'summary', 'objective',
            'projects', 'certifications', 'awards', 'languages', 'profile',
            'work history', 'employment', 'qualifications', 'achievements',
            'work experience', 'professional experience', 'career history'
        ]
        
        line_lower = line.lower().strip()
        return any(keyword == line_lower or line_lower.startswith(keyword) for keyword in section_keywords)
    
    def _looks_like_summary_text(self, text):
        """Check if text looks like summary/objective content rather than experience details"""
        if not text:
            return False
        
        # Summary text characteristics:
        # - Often starts with personal pronouns or adjectives
        # - Contains general career statements
        # - No action verbs typical of bullet points
        
        text_lower = text.lower().strip()
        
        # Skip if starts with typical summary phrases
        summary_starters = [
            'detail-oriented', 'results-oriented', 'highly motivated',
            'experienced professional', 'proven track record', 'dedicated',
            'proficient in', 'recognized for', 'skilled in',
            'strong background', 'extensive experience', 'adept at'
        ]
        
        # Check if line starts with these phrases (more likely to be summary)
        for starter in summary_starters:
            if text_lower.startswith(starter):
                return True
        
        # If the line is very long (>150 chars) and doesn't start with action verb, might be summary
        if len(text) > 150:
            action_verbs = ['managed', 'developed', 'created', 'led', 'coordinated', 
                          'implemented', 'executed', 'analyzed', 'designed', 'built',
                          'oversaw', 'directed', 'supervised', 'maintained', 'conducted',
                          'established', 'improved', 'increased', 'reduced', 'achieved',
                          'tracked', 'monitored', 'facilitated', 'collaborated', 'assisted',
                          'provided', 'processed', 'scheduled', 'organized']
            
            # Check if starts with action verb
            first_word = text_lower.split()[0] if text_lower.split() else ''
            if first_word not in action_verbs:
                return True
        
        return False
    
    def _has_contact_info(self, text):
        """Check if text contains contact information"""
        return bool(re.search(r'@|http|linkedin|\d{3}[-.\s]\d{3}', text, re.IGNORECASE))
    
    def _looks_like_company_or_role(self, line):
        """Check if line looks like company name or job title"""
        line_clean = line.strip()
        line_lower = line_clean.lower()
        
        # CRITICAL: Reject section headings
        section_keywords = [
            'experience', 'work history', 'employment', 'professional experience',
            'education', 'academic background', 'skills', 'technical skills',
            'summary', 'objective', 'profile', 'certifications', 'projects',
            'awards', 'references', 'languages', 'publications'
        ]
        if any(kw == line_lower for kw in section_keywords):
            return False
        
        # CRITICAL: Reject candidate names (usually 2-3 words, all caps, no location/company markers)
        # Names don't have dashes, commas with locations, or company indicators
        if line_clean.isupper() and len(line_clean.split()) <= 3:
            has_company_markers = any(marker in line_clean for marker in [' - ', ', ', ' Inc', ' LLC', ' Corp'])
            if not has_company_markers:
                return False  # Likely a name, not company
        
        # CRITICAL: Reject contact info
        if '@' in line_clean or re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', line_clean):
            return False
        
        # CRITICAL: Reject goal/summary statements
        if any(kw in line_lower for kw in ['goal:', 'seeking', 'motivated', 'objective:', 'summary:']):
            return False
        
        # CRITICAL: Reject single words (except known company names)
        words = line_clean.split()
        if len(words) == 1 and not any(ind in line_lower for ind in ['inc', 'corp', 'llc', 'ltd']):
            return False
        
        # Usually title case or upper, not too long, no bullets
        return (line_clean.istitle() or line_clean.isupper()) and len(line_clean) < 100 and not line_clean.startswith('•')
    
    def _looks_like_detail_bullet(self, text):
        """Check if text looks like a detail bullet (starts with action verb or is descriptive)"""
        if not text:
            return False
        
        text_lower = text.lower().strip()
        
        # Common action verbs used in resume bullets
        action_verbs = [
            'managed', 'developed', 'created', 'led', 'coordinated', 
            'implemented', 'executed', 'analyzed', 'designed', 'built',
            'oversaw', 'directed', 'supervised', 'maintained', 'conducted',
            'established', 'improved', 'increased', 'reduced', 'achieved',
            'tracked', 'monitored', 'facilitated', 'collaborated', 'assisted',
            'provided', 'processed', 'scheduled', 'organized', 'acted',
            'demonstrated', 'tested', 'collected', 'ensured', 'handled',
            'performed', 'supported', 'delivered', 'prepared', 'reviewed',
            'updated', 'resolved', 'trained', 'documented', 'communicated',
            'participated', 'contributed', 'streamlined', 'optimized'
        ]
        
        # Check if starts with action verb
        first_word = text_lower.split()[0] if text_lower.split() else ''
        if first_word in action_verbs:
            return True
        
        # Check if line is long and descriptive (likely a detail)
        # Detail bullets are usually longer than company/role names
        if len(text) > 60 and not (text.istitle() or text.isupper()):
            return True
        
        return False
    
    def _contains_date_range(self, line):
        """Check if line contains date range"""
        date_patterns = [
            r'\b(19|20)\d{2}\b.*\b(19|20)\d{2}\b',  # 2020 - 2023
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}',  # Jan 2020
            r'\d{1,2}/\d{4}',  # 01/2020
        ]
        return any(re.search(pattern, line, re.IGNORECASE) for pattern in date_patterns)

    def _should_merge_fragment(self, prev, curr):
        """Heuristic to decide if 'curr' is a continuation of 'prev' (fragment stitching).
        Merges tiny lines like 'Word', 'Access) and other', 'reporting', 'and documentation.' into the previous bullet.
        """
        if not curr:
            return False
        # Do not merge if curr looks like a section header or explicit new sentence with leading verb
        if self._is_section_header(curr) or self._looks_like_detail_bullet(curr):
            return False
        # Merge if curr is a very short fragment (<= 25 chars or <= 5 words)
        words = curr.split()
        if len(curr) <= 25 or len(words) <= 5:
            return True
        # Merge if previous line has open parenthesis not yet closed
        if prev and prev.count('(') > prev.count(')'):
            return True
        return False

    def _merge_fragment(self, prev, curr):
        """Join fragment 'curr' to 'prev' with an appropriate delimiter."""
        if not prev:
            return curr
        # If prev has an open parenthesis, separate with comma unless prev already ends with a space or comma
        if prev.count('(') > prev.count(')'):
            sep = ', ' if not prev.endswith((' ', ',')) else ''
            return (prev + sep + curr).strip()
        # If prev ends without terminal punctuation, add space
        if not prev.endswith(('.', '!', '?', ')', ']')):
            return (prev + ' ' + curr).strip()
        # Default: space join
        return (prev + ' ' + curr).strip()
    
    def _print_parsing_summary(self, data):
        """Print parsing summary"""
        print(f"👤 Name: {data['name']}")
        print(f"📧 Email: {data['email']}")
        print(f"📱 Phone: {data['phone']}")
        print(f"🔗 LinkedIn: {data['linkedin'][:50] if data['linkedin'] else 'Not found'}")
        print(f"📅 DOB: {data['dob']}")
        print(f"💼 Experience Entries: {len(data['experience'])}")
        print(f"🎓 Education Entries: {len(data['education'])}")
        print(f"🛠️  Skills: {len(data['skills'])}")
        print(f"📂 Projects: {len(data['projects'])}")
        print(f"🏆 Certifications: {len(data['certifications'])}")
        print(f"🏅 Awards: {len(data['awards'])}")
        print(f"🌐 Languages: {len(data['languages'])}")
        print(f"\n{'='*70}\n")


def parse_resume(file_path, file_type):
    """Main function to parse resume"""
    parser = ResumeParser(file_path, file_type)
    return parser.parse()
