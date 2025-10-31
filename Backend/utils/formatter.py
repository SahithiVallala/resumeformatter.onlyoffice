


from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import os
import io
import traceback

def format_resume(resume_data, template_format, output_path):
    """Generate formatted resume by cloning template and replacing content"""
    try:
        template_path = template_format.get('template_path')
        template_type = template_format.get('template_type')
        
        print(f"[Formatter] Template path: {template_path}")
        print(f"[Formatter] Template type: {template_type}")
        print(f"[Formatter] Output path: {output_path}")
        
        if not template_path:
            print(f"[Formatter] ERROR: Template path is None or empty")
            return False
            
        if not os.path.exists(template_path):
            print(f"[Formatter] ERROR: Template file not found at: {template_path}")
            return False
        
        print(f"[Formatter] Template file exists, proceeding with {template_type} formatting")
        
        if template_type == 'pdf':
            return format_pdf_resume(resume_data, template_path, template_format, output_path)
        elif template_type in ['docx', 'doc']:
            return format_word_resume(resume_data, template_path, template_format, output_path)
        else:
            print(f"[Formatter] ERROR: Unsupported template type: {template_type}")
            return False
            
    except Exception as e:
        print(f"[Formatter] ERROR: Exception in format_resume: {e}")
        traceback.print_exc()
        return False

def format_pdf_resume(resume_data, template_path, template_format, output_path):
    """Format resume using PDF template - preserves all visual elements"""
    try:
        print(f"[PDF Formatter] Starting PDF formatting")
        print(f"[PDF Formatter] Resume name: {resume_data.get('name', 'N/A')}")
        
        # Read the template PDF
        reader = PdfReader(template_path)
        writer = PdfWriter()
        
        print(f"[PDF Formatter] Template has {len(reader.pages)} pages")
        
        # Get the first page (template)
        template_page = reader.pages[0]
        
        # Create overlay with resume content
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=(float(template_page.mediabox.width), 
                                              float(template_page.mediabox.height)))
        
        # Get positioning from template format
        page_width = float(template_page.mediabox.width)
        page_height = float(template_page.mediabox.height)
        margins = template_format['page']['margins']
        
        print(f"[PDF Formatter] Page size: {page_width} x {page_height}")
        print(f"[PDF Formatter] Margins: {margins}")
        
        # Position for content (start below letterhead/header area)
        y_position = page_height - margins['top'] - 100  # Leave space for letterhead
        x_position = margins['left']
        
        # Add name if not in template
        if resume_data.get('name') and not has_name_in_template(template_format):
            can.setFont(template_format['name']['font'], template_format['name']['size'])
            if template_format['name']['alignment'] == 'center':
                can.drawCentredString(page_width / 2, y_position, resume_data['name'].upper())
            else:
                can.drawString(x_position, y_position, resume_data['name'].upper())
            y_position -= 30
        
        # Add contact info
        contact_items = [x for x in [resume_data.get('email'), resume_data.get('phone'), 
                                     resume_data.get('linkedin')] if x]
        if contact_items:
            can.setFont(template_format['body']['font'], template_format['body']['size'])
            contact_text = ' | '.join(contact_items)
            can.drawCentredString(page_width / 2, y_position, contact_text)
            y_position -= 25
        
        # Add sections
        print(f"[PDF Formatter] Template has {len(template_format['sections'])} sections")
        print(f"[PDF Formatter] Resume has {len(resume_data['sections'])} sections: {list(resume_data['sections'].keys())}")
        
        for section in template_format['sections']:
            section_content = find_matching_section(section['heading'], resume_data['sections'])
            
            print(f"[PDF Formatter] Section '{section['heading']}': found {len(section_content) if section_content else 0} items")
            
            if not section_content:
                continue
            
            # Section heading
            can.setFont(section['font'], section['size'])
            can.drawString(x_position, y_position, section['heading'])
            y_position -= 15
            
            # Section underline if needed
            if section.get('has_underline'):
                can.line(x_position, y_position + 5, page_width - margins['right'], y_position + 5)
                y_position -= 5
            
            # Section content
            can.setFont(template_format['body']['font'], template_format['body']['size'])
            for item in section_content[:15]:
                if item.strip() and y_position > margins['bottom'] + 50:
                    clean_text = item.strip().lstrip('•').strip()
                    # Wrap text if too long
                    max_width = page_width - margins['left'] - margins['right'] - 20
                    wrapped_lines = wrap_text(clean_text, can, max_width)
                    
                    for line in wrapped_lines:
                        if y_position > margins['bottom'] + 50:
                            can.drawString(x_position + 10, y_position, f'• {line}')
                            y_position -= template_format['body']['line_spacing']
            
            y_position -= 10  # Space between sections
        
        can.save()
        
        print(f"[PDF Formatter] Canvas saved, merging with template")
        
        # Merge overlay with template
        packet.seek(0)
        overlay_pdf = PdfReader(packet)
        overlay_page = overlay_pdf.pages[0]
        
        # Merge template (background with letterhead) with content (overlay)
        template_page.merge_page(overlay_page)
        writer.add_page(template_page)
        
        print(f"[PDF Formatter] Writing output to: {output_path}")
        
        # Write output
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        print(f"[PDF Formatter] Successfully created formatted resume")
        return True
        
    except Exception as e:
        print(f"[PDF Formatter] ERROR: {e}")
        traceback.print_exc()
        return False

def format_word_resume(resume_data, template_path, template_format, output_path):
    """Format resume using Word template - preserves all visual elements"""
    try:
        print(f"[Word Formatter] Starting Word formatting")
        print(f"[Word Formatter] Template path: {template_path}")
        print(f"[Word Formatter] Resume name: {resume_data.get('name', 'N/A')}")
        
        # Check if it's a .doc file (old format)
        if template_path.lower().endswith('.doc'):
            print(f"[Word Formatter] ERROR: .doc files (old Word format) are not supported")
            print(f"[Word Formatter] Please convert the template to .docx format")
            print(f"[Word Formatter] Falling back to PDF generation...")
            # Fall back to creating a PDF from scratch
            return create_pdf_from_scratch(resume_data, template_format, output_path)
        
        # Open template document
        doc = Document(template_path)
        
        # Find and replace placeholders
        replacements = {
            '[NAME]': resume_data.get('name', ''),
            '[Email]': resume_data.get('email', ''),
            '[Phone]': resume_data.get('phone', ''),
            '[LinkedIn]': resume_data.get('linkedin', ''),
            'Your Name': resume_data.get('name', ''),
            'your.email@example.com': resume_data.get('email', ''),
            '(123) 456-7890': resume_data.get('phone', ''),
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key.lower() in paragraph.text.lower():
                    # Preserve formatting while replacing
                    for run in paragraph.runs:
                        if key.lower() in run.text.lower():
                            run.text = run.text.replace(key, value)
        
        # Replace in tables (common in resume templates)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key.lower() in paragraph.text.lower():
                                for run in paragraph.runs:
                                    if key.lower() in run.text.lower():
                                        run.text = run.text.replace(key, value)
        
        # Add resume sections to template
        # Find section markers and replace with actual content
        section_markers = {
            'experience': ['WORK EXPERIENCE', 'EXPERIENCE', 'EMPLOYMENT HISTORY'],
            'education': ['EDUCATION', 'ACADEMIC BACKGROUND'],
            'skills': ['SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES'],
            'summary': ['SUMMARY', 'PROFESSIONAL SUMMARY', 'OBJECTIVE'],
            'projects': ['PROJECTS', 'KEY PROJECTS']
        }
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text_upper = paragraph.text.upper().strip()
            
            for section_key, markers in section_markers.items():
                if any(marker in para_text_upper for marker in markers):
                    # Found a section, add content after it
                    section_content = find_matching_section(paragraph.text, resume_data['sections'])
                    
                    if section_content:
                        # Add content after section heading
                        insert_idx = para_idx + 1
                        for item in section_content[:10]:
                            if item.strip():
                                new_para = doc.paragraphs[insert_idx]._element
                                new_para = doc.add_paragraph(f'• {item.strip().lstrip("•").strip()}')
                                insert_idx += 1
        
        # Convert to PDF (Word output)
        # For now, save as .docx, later we can add conversion to PDF
        doc.save(output_path.replace('.pdf', '.docx'))
        
        return True
        
    except Exception as e:
        print(f"Error formatting Word resume: {e}")
        traceback.print_exc()
        return False

def wrap_text(text, canvas_obj, max_width):
    """Wrap text to fit within max_width"""
    words = text.split()
    lines = []
    current_line = []
    
    for word in words:
        test_line = ' '.join(current_line + [word])
        if canvas_obj.stringWidth(test_line) <= max_width:
            current_line.append(word)
        else:
            if current_line:
                lines.append(' '.join(current_line))
            current_line = [word]
    
    if current_line:
        lines.append(' '.join(current_line))
    
    return lines if lines else [text]

def has_name_in_template(template_format):
    """Check if template already has a name placeholder"""
    # This would check if template has [NAME] or similar
    return False

def create_pdf_from_scratch(resume_data, template_format, output_path):
    """Create a PDF resume from scratch when template can't be used"""
    try:
        print(f"[PDF Creator] Creating PDF from scratch")
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        page_width = template_format['page']['width']
        page_height = template_format['page']['height']
        margins = template_format['page']['margins']
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=(page_width, page_height),
            topMargin=margins['top'],
            bottomMargin=margins['bottom'],
            leftMargin=margins['left'],
            rightMargin=margins['right']
        )
        
        name_style = ParagraphStyle(
            'Name',
            fontName=template_format['name']['font'],
            fontSize=template_format['name']['size'],
            alignment=TA_CENTER if template_format['name']['alignment'] == 'center' else TA_LEFT,
            spaceAfter=12
        )
        
        body_style = ParagraphStyle(
            'Body',
            fontName=template_format['body']['font'],
            fontSize=template_format['body']['size'],
            leading=template_format['body']['line_spacing'],
            spaceAfter=6
        )
        
        content = []
        
        # Add name
        if resume_data.get('name'):
            content.append(Paragraph(resume_data['name'].upper(), name_style))
        
        # Add contact info
        contact_items = [x for x in [resume_data.get('email'), resume_data.get('phone'), 
                                     resume_data.get('linkedin')] if x]
        if contact_items:
            contact_style = ParagraphStyle('Contact', parent=body_style, alignment=TA_CENTER)
            content.append(Paragraph(' | '.join(contact_items), contact_style))
        
        content.append(Spacer(1, 12))
        
        # Add sections
        for section in template_format.get('sections', []):
            section_content = find_matching_section(section['heading'], resume_data['sections'])
            
            if not section_content:
                continue
            
            section_style = ParagraphStyle(
                'Section',
                fontName=section['font'],
                fontSize=section['size'],
                spaceAfter=6,
                spaceBefore=12
            )
            
            content.append(Paragraph(section['heading'], section_style))
            
            if section.get('has_underline'):
                content.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=6))
            
            for item in section_content[:15]:
                if item.strip():
                    clean = item.strip().lstrip('•').strip()
                    content.append(Paragraph(f'• {clean}', body_style))
        
        doc.build(content)
        print(f"[PDF Creator] Successfully created PDF")
        return True
        
    except Exception as e:
        print(f"[PDF Creator] ERROR: {e}")
        traceback.print_exc()
        return False

def find_matching_section(template_heading, resume_sections):
    """Match template section heading to resume section content"""
    template_lower = template_heading.lower().strip()
    
    # Direct match
    if template_lower in resume_sections:
        return resume_sections[template_lower]
    
    # Keyword matching
    keywords = {
        'experience': ['experience', 'employment', 'work', 'professional'],
        'education': ['education', 'academic', 'qualification'],
        'skills': ['skills', 'technical', 'competencies', 'expertise'],
        'summary': ['summary', 'objective', 'profile', 'about'],
        'projects': ['projects', 'portfolio'],
        'certifications': ['certifications', 'certificates', 'licenses'],
        'awards': ['awards', 'achievements', 'honors']
    }
    
    for key, patterns in keywords.items():
        if any(p in template_lower for p in patterns):
            for resume_key, content in resume_sections.items():
                if any(p in resume_key for p in patterns):
                    return content
    
    return []
