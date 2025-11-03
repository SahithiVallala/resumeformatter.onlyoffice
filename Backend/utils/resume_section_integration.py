"""
Resume Section Integration Module
Integrates enhanced section classifier with ordered renderer
Provides a simple API for the main application
"""

from typing import Dict, List, Optional
from docx import Document
import os

try:
    from utils.enhanced_section_classifier import get_section_classifier
    CLASSIFIER_AVAILABLE = True
except ImportError:
    CLASSIFIER_AVAILABLE = False
    print("⚠️  Enhanced classifier not available")

try:
    from utils.ordered_section_renderer import OrderedSectionRenderer
    RENDERER_AVAILABLE = True
except ImportError:
    RENDERER_AVAILABLE = False
    print("⚠️  Ordered renderer not available")


class ResumeFormatter:
    """
    Main integration class for resume formatting with intelligent section mapping
    """
    
    def __init__(self, confidence_threshold: float = 0.6):
        """
        Initialize formatter
        
        Args:
            confidence_threshold: Minimum confidence for section classification (0-1)
        """
        self.confidence_threshold = confidence_threshold
        self.classifier = None
        
        if CLASSIFIER_AVAILABLE:
            self.classifier = get_section_classifier(confidence_threshold)
    
    def extract_sections_from_docx(self, docx_path: str) -> List[Dict]:
        """
        Extract sections from a DOCX resume
        
        Args:
            docx_path: Path to candidate's resume DOCX
            
        Returns:
            List of dicts with 'heading', 'content', 'position'
        """
        doc = Document(docx_path)
        sections = []
        current_section = None
        position_index = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Skip contact info (first few lines)
            if position_index < 3 and any(c in text for c in ['@', 'phone', 'email', '+91', '(']):
                position_index += 1
                continue
            
            # Detect if this is a heading
            is_heading = self._is_heading(para)
            
            if is_heading:
                # Save previous section
                if current_section:
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    'heading': text,
                    'content': [],
                    'position': position_index,
                    'has_heading': True
                }
            else:
                # Content paragraph
                if current_section is None:
                    # Paragraph without heading (e.g., summary at top)
                    current_section = {
                        'heading': None,
                        'content': [text],
                        'position': position_index,
                        'has_heading': False
                    }
                else:
                    current_section['content'].append(text)
            
            position_index += 1
        
        # Save last section
        if current_section:
            sections.append(current_section)
        
        # Convert content lists to strings
        for section in sections:
            section['content'] = '\n'.join(section['content'])
        
        return sections
    
    def _is_heading(self, para) -> bool:
        """
        Detect if paragraph is a heading
        
        Args:
            para: Paragraph object
            
        Returns:
            True if heading, False otherwise
        """
        text = para.text.strip()
        
        # Check style
        if para.style.name.startswith('Heading'):
            return True
        
        # Check formatting
        if para.runs:
            first_run = para.runs[0]
            if first_run.bold and len(text.split()) <= 5:
                return True
        
        # Check ALL CAPS short phrases
        if text.isupper() and len(text.split()) <= 4:
            return True
        
        # Check common section keywords
        section_keywords = [
            'employment', 'education', 'skills', 'experience', 'summary',
            'projects', 'certifications', 'awards', 'languages', 'profile',
            'objective', 'qualifications', 'background', 'history'
        ]
        text_lower = text.lower()
        if any(keyword in text_lower for keyword in section_keywords):
            if len(text.split()) <= 5:
                return True
        
        return False
    
    def extract_template_sections(self, template_path: str) -> List[str]:
        """
        Extract section names from template
        
        Args:
            template_path: Path to template DOCX
            
        Returns:
            List of section names in template order
        """
        doc = Document(template_path)
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings
            is_heading = (
                para.style.name.startswith('Heading') or
                (para.runs and para.runs[0].bold and len(text.split()) <= 5) or
                (text.isupper() and len(text.split()) <= 4)
            )
            
            if is_heading:
                sections.append(text)
        
        return sections
    
    def format_resume(self, candidate_docx: str, template_docx: str, 
                     output_path: str, contact_info: Optional[Dict] = None) -> Dict:
        """
        Main method: Format a resume using intelligent section mapping
        
        Args:
            candidate_docx: Path to candidate's resume DOCX
            template_docx: Path to template DOCX
            output_path: Path to save formatted resume
            contact_info: Optional dict with name, email, phone, address
            
        Returns:
            Dict with processing results and statistics
        """
        print(f"\n{'='*70}")
        print(f"🚀 INTELLIGENT RESUME FORMATTING")
        print(f"{'='*70}\n")
        
        # Step 1: Extract sections from candidate resume
        print("📄 Extracting sections from candidate resume...")
        candidate_sections = self.extract_sections_from_docx(candidate_docx)
        print(f"   Found {len(candidate_sections)} sections\n")
        
        # Step 2: Extract template structure
        print("📋 Analyzing template structure...")
        template_sections = self.extract_template_sections(template_docx)
        print(f"   Template has {len(template_sections)} sections: {', '.join(template_sections)}\n")
        
        # Step 3: Classify and map sections
        if self.classifier and CLASSIFIER_AVAILABLE:
            print("🧠 Classifying sections with ML...")
            mapped_sections = self.classifier.batch_classify(candidate_sections, template_sections)
        else:
            print("⚠️  Using basic section mapping (ML not available)...")
            mapped_sections = self._basic_section_mapping(candidate_sections, template_sections)
        
        # Step 4: Render formatted resume
        if RENDERER_AVAILABLE:
            print("\n📝 Rendering formatted resume...")
            renderer = OrderedSectionRenderer(template_docx)
            renderer.render(mapped_sections, output_path, contact_info)
        else:
            print("⚠️  Ordered renderer not available, using basic output...")
            self._basic_render(mapped_sections, output_path, contact_info)
        
        # Return statistics
        stats = {
            'success': True,
            'candidate_sections': len(candidate_sections),
            'template_sections': len(template_sections),
            'mapped_sections': len(mapped_sections),
            'output_path': output_path
        }
        
        print(f"\n✅ FORMATTING COMPLETE")
        print(f"   Mapped {stats['mapped_sections']}/{stats['candidate_sections']} sections")
        print(f"{'='*70}\n")
        
        return stats
    
    def _basic_section_mapping(self, candidate_sections: List[Dict], 
                               template_sections: List[str]) -> Dict[str, str]:
        """
        Basic section mapping fallback (when ML is not available)
        
        Args:
            candidate_sections: List of candidate sections
            template_sections: List of template section names
            
        Returns:
            Dict mapping template sections to content
        """
        mapped = {}
        
        for section in candidate_sections:
            heading = section.get('heading', '').lower()
            content = section.get('content', '')
            
            # Try to match with template sections
            matched = False
            for template_section in template_sections:
                template_lower = template_section.lower()
                
                if heading and (heading in template_lower or template_lower in heading):
                    mapped[template_section] = content
                    matched = True
                    break
            
            if not matched and heading:
                # Add with original heading
                mapped[section.get('heading', 'Other')] = content
        
        return mapped
    
    def _basic_render(self, mapped_sections: Dict[str, str], 
                     output_path: str, contact_info: Optional[Dict] = None):
        """
        Basic rendering fallback (when renderer is not available)
        
        Args:
            mapped_sections: Mapped sections
            output_path: Output path
            contact_info: Contact information
        """
        doc = Document()
        
        # Add contact info
        if contact_info:
            if contact_info.get('name'):
                doc.add_paragraph(contact_info['name']).runs[0].bold = True
            contact_parts = []
            if contact_info.get('email'):
                contact_parts.append(contact_info['email'])
            if contact_info.get('phone'):
                contact_parts.append(contact_info['phone'])
            if contact_parts:
                doc.add_paragraph(' | '.join(contact_parts))
            doc.add_paragraph()
        
        # Add sections
        for section_name, content in mapped_sections.items():
            # Add heading
            heading_para = doc.add_paragraph(section_name)
            heading_para.runs[0].bold = True
            
            # Add content
            doc.add_paragraph(content)
            doc.add_paragraph()
        
        doc.save(output_path)


# Convenience function for easy integration
def format_resume_with_intelligent_mapping(candidate_docx: str, template_docx: str,
                                          output_path: str, contact_info: Optional[Dict] = None,
                                          confidence_threshold: float = 0.6) -> Dict:
    """
    Format a resume with intelligent section mapping
    
    Args:
        candidate_docx: Path to candidate's resume DOCX
        template_docx: Path to template DOCX
        output_path: Path to save formatted resume
        contact_info: Optional contact information dict
        confidence_threshold: Minimum confidence for classification (0-1)
        
    Returns:
        Dict with processing statistics
    """
    formatter = ResumeFormatter(confidence_threshold)
    return formatter.format_resume(candidate_docx, template_docx, output_path, contact_info)


# Example usage
if __name__ == "__main__":
    result = format_resume_with_intelligent_mapping(
        candidate_docx="candidate_resume.docx",
        template_docx="template.docx",
        output_path="formatted_resume.docx",
        contact_info={
            "name": "John Doe",
            "email": "john@example.com",
            "phone": "+1-234-567-8900"
        },
        confidence_threshold=0.6
    )
    
    print(f"\nResults: {result}")
