"""
Quick test to verify clean output with NO duplication
"""
import os
import sys
from docx import Document

test_data = {
    'name': 'Test User',
    'email': 'test@email.com',
    'phone': '555-1234',
    'experience': [
        {
            'company': 'Test Company',
            'role': 'Software Engineer',
            'duration': '2020-2023',
            'details': ['Built applications', 'Led team']
        }
    ],
    'education': [
        {
            'degree': 'Bachelor of Science',
            'institution': 'Test University',
            'year': '2018',
            'details': []
        }
    ],
    'skills': ['Python', 'Java'],
    'sections': {}
}

print("\n" + "="*70)
print("CLEAN ATS RESUME TEST")
print("="*70)

from models.database import TemplateDB
from config import Config
from utils.word_formatter import format_word_document

db = TemplateDB()
templates = db.get_all_templates()

if not templates:
    print("\n❌ No templates!")
    sys.exit(1)

template = db.get_template(templates[0]['id'])
template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])

print(f"\nTemplate: {template['name']}")

template_analysis = {
    'template_path': template_path,
    'template_type': 'docx'
}

output_path = os.path.join(Config.OUTPUT_FOLDER, 'test_clean_ats.docx')

print(f"\n🎯 Generating clean resume...\n")

success = format_word_document(test_data, template_analysis, output_path)

if success:
    print("\n" + "="*70)
    print("✅ SUCCESS")
    print("="*70)
    
    doc = Document(output_path)
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    # Count occurrences
    test_company_count = full_text.count('Test Company')
    software_engineer_count = full_text.count('Software Engineer')
    
    print(f"\n📊 Verification:")
    print(f"  • Total paragraphs: {len(doc.paragraphs)}")
    print(f"  • Total tables: {len(doc.tables)}")
    print(f"  • 'Test Company' appears: {test_company_count} time(s)")
    print(f"  • 'Software Engineer' appears: {software_engineer_count} time(s)")
    
    if test_company_count == 1 and software_engineer_count == 1:
        print(f"\n✅ PERFECT! No duplication detected!")
    else:
        print(f"\n⚠️  WARNING: Content may be duplicated")
        print(f"     Expected each to appear 1 time")
    
    print(f"\n📁 Open file:")
    print(f"   {output_path}")
    print(f"\n💡 Check: Should see ONLY table format, NO raw bullets below")
    
else:
    print("\n❌ FAILED")

print("\n")
