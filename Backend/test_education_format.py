"""Test the new education format with both colon and 'in' separators"""
import os
from docx import Document
from utils.word_formatter import WordDocumentFormatter
from models.database import TemplateDB
from config import Config

# Test data with both formats
test_data = {
    'name': 'Test User',
    'email': 'test@email.com',
    'phone': '555-1234',
    'experience': [],
    'education': [
        {
            'degree': 'Master of Science : Leadership',  # Colon format
            'institution': 'Walden University',
            'year': '2018',
            'details': []
        },
        {
            'degree': 'Master of Science in Data Science',  # "in" format
            'institution': 'PS University',
            'year': '2020-2022',
            'details': []
        },
        {
            'degree': 'Bachelor of Technology in Computer Science',  # "in" format
            'institution': 'State Technical University',
            'year': '2014-2018',
            'details': []
        }
    ],
    'skills': [],
    'sections': {}
}

print("\n" + "="*70)
print("TESTING NEW EDUCATION FORMAT")
print("="*70)

db = TemplateDB()
templates = db.get_all_templates()

if not templates:
    print("\n❌ No templates!")
    exit(1)

template = db.get_template(templates[0]['id'])
template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])

print(f"\nTemplate: {template['name']}")

template_analysis = {
    'template_path': template_path,
    'template_type': 'docx'
}

output_path = os.path.join(Config.OUTPUT_FOLDER, 'test_education_format.docx')

print(f"\n🎓 Testing education formats:\n")
for i, edu in enumerate(test_data['education'], 1):
    print(f"  {i}. Degree: '{edu['degree']}'")
    print(f"     Institution: '{edu['institution']}'")
    print(f"     Year: '{edu['year']}'")
    print()

from utils.word_formatter import format_word_document

print(f"🎯 Generating formatted resume...\n")

success = format_word_document(test_data, template_analysis, output_path)

if success:
    print("\n" + "="*70)
    print("✅ SUCCESS - Check the console above for split details!")
    print("="*70)
    
    print(f"\n📄 Output: {output_path}")
    print(f"\n💡 Look for these lines in the output above:")
    print(f"   ✂️  Split at colon: LEFT='Master of Science' | Field='Leadership'")
    print(f"   ✂️  Split at 'in': LEFT='Master of Science' | Field='Data Science'")
    print(f"   📐 Format: LEFT='...' | RIGHT='...'")
    
    # Verify the document
    doc = Document(output_path)
    print(f"\n📊 Generated document has {len(doc.tables)} tables")
    
    print(f"\n🔍 Checking education tables:")
    for i, table in enumerate(doc.tables):
        if table.rows and table.columns:
            first_cell_text = table.rows[0].cells[0].text.strip()
            if any(degree_word in first_cell_text.lower() for degree_word in ['master', 'bachelor', 'science', 'technology']):
                left_text = table.rows[0].cells[0].text.strip()
                right_text = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ''
                print(f"\n  Table {i+1}:")
                print(f"    LEFT:  '{left_text}'")
                print(f"    RIGHT: '{right_text}'")
    
    print(f"\n✅ OPEN FILE TO VERIFY:")
    print(f"   {output_path}")
    print()
    
else:
    print("\n❌ FAILED")

print()
